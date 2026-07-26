#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Chapter 4: GraphPad Prism 9.0 Data Analysis Tutorial for IC50 Calculation
Appends to existing v2.0 document
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_PATH = "final/\u82cf\u6728\u5bf9CNE-2\u7ec6\u80de\u589e\u6b96\u6291\u5236\u7387\u53caIC50\u6d4b\u5b9a\u5168\u6d41\u7a0b\u64cd\u4f5c\u624b\u518c_v2.0.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_spacing(para, before=0, after=0, line_rule=None, line_val=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line_rule and line_val:
        spacing.set(qn('w:lineRule'), line_rule)
        spacing.set(qn('w:line'), str(line_val))
    pPr.append(spacing)

def add_chapter_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '\u9ed1\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        set_para_spacing(p, before=240, after=120)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = '\u9ed1\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        set_para_spacing(p, before=180, after=60)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '\u9ed1\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        set_para_spacing(p, before=120, after=60)
    return p

def add_body_para(doc, text, indent=False, bold_prefix=None):
    p = doc.add_paragraph()
    if indent:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '480')
        pPr.append(ind)
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.name = '\u5b8b\u4f53'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=60, line_rule='auto', line_val=360)
    return p

def add_numbered_step(doc, num, title, content_lines):
    p = doc.add_paragraph()
    run = p.add_run('\u6b65\u9aa4' + str(num) + '\uff1a' + title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '\u9ed1\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    set_para_spacing(p, before=120, after=30)
    for line in content_lines:
        p2 = doc.add_paragraph()
        pPr2 = p2._p.get_or_add_pPr()
        ind2 = OxmlElement('w:ind')
        ind2.set(qn('w:left'), '480')
        ind2.set(qn('w:firstLine'), '0')
        pPr2.append(ind2)
        run2 = p2.add_run('\u2022 ' + line)
        run2.font.size = Pt(11)
        run2.font.name = '\u5b8b\u4f53'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p2, before=0, after=30, line_rule='auto', line_val=360)

def add_warning_box(doc, title, lines, color='FF0000', bg='FFF2CC'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)
    p_title = cell.paragraphs[0]
    r = p_title.add_run('\u26a0 ' + title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
    set_para_spacing(p_title, before=60, after=30)
    for line in lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run('\u25b6 ' + line)
        r2.font.size = Pt(10.5)
        r2.font.name = '\u5b8b\u4f53'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p2, before=0, after=30, line_rule='auto', line_val=340)
    doc.add_paragraph()

def add_tip_box(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'E8F4FD')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2E74B5')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    p_title = cell.paragraphs[0]
    r = p_title.add_run('\U0001f4a1 ' + title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    set_para_spacing(p_title, before=60, after=30)
    for line in lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run('  ' + line)
        r2.font.size = Pt(10.5)
        r2.font.name = '\u5b8b\u4f53'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p2, before=0, after=30, line_rule='auto', line_val=340)
    doc.add_paragraph()

def add_three_line_table(doc, headers, rows, caption=None):
    if caption:
        p_cap = doc.add_paragraph()
        r_cap = p_cap.add_run(caption)
        r_cap.bold = True
        r_cap.font.size = Pt(10.5)
        r_cap.font.name = '\u5b8b\u4f53'
        r_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_cap, before=60, after=30)
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name, sz in [('top', '12'), ('bottom', '12')]:
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    for border_name in ['insideH', 'insideV']:
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_bg(cell, 'D6E4F7')
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bot_b = OxmlElement('w:bottom')
        bot_b.set(qn('w:val'), 'single')
        bot_b.set(qn('w:sz'), '6')
        bot_b.set(qn('w:space'), '0')
        bot_b.set(qn('w:color'), '000000')
        tcBorders.append(bot_b)
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '\u5b8b\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F5F9FF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '\u5b8b\u4f53'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    doc.add_paragraph()

def add_formula_box(doc, formula_title, formula_lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F0F7E6')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '548235')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    p_title = cell.paragraphs[0]
    r = p_title.add_run('\U0001f4d0 ' + formula_title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(0x54, 0x82, 0x35)
    set_para_spacing(p_title, before=60, after=30)
    for line in formula_lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(line)
        r2.font.size = Pt(11)
        r2.font.name = 'Courier New'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p2, before=0, after=30, line_rule='auto', line_val=360)
    doc.add_paragraph()


# ============================================================
# MAIN: Load document and add Chapter 4
# ============================================================
print("Loading document for Chapter 4 addition...")
doc = Document(DOC_PATH)

doc.add_page_break()
add_chapter_heading(doc, '\u7b2c4\u7ae0  \u6570\u636e\u5904\u7406\u4e0eIC\u2085\u2080\u8ba1\u7b97\u64cd\u4f5c\u6307\u5357', level=1)

add_body_para(doc,
    '\u672c\u7ae0\u63d0\u4f9b\u4ece\u539f\u59cb\u9176\u6807\u4eea OD450 \u6570\u636e\u5230\u6700\u7ec8 IC\u2085\u2080 \u62a5\u544a\u503c\u7684\u5b8c\u6574\u6570\u636e\u5904\u7406\u6d41\u7a0b\uff0c'
    '\u5305\u62ec Excel \u9884\u5904\u7406\u3001GraphPad Prism 9.0 \u56db\u53c2\u6570 Logistic\uff084PL\uff09\u975e\u7ebf\u6027\u56de\u5f52\u62df\u5408\uff0c'
    '\u4ee5\u53ca\u7edf\u8ba1\u7ed3\u679c\u7684\u89e3\u8bfb\u4e0e\u62a5\u544a\u89c4\u8303\u3002\u672c\u7ae0\u5185\u5bb9\u9002\u5408\u96f6\u57fa\u7840\u5b66\u751f\u6309\u6b65\u64cd\u4f5c\uff0c\u6bcf\u4e2a\u754c\u9762\u64cd\u4f5c\u5747\u6709\u8be6\u7ec6\u8bf4\u660e\u3002',
    indent=True)

# Data flow overview using colored text
add_body_para(doc, '\u6570\u636e\u5206\u6790\u6d41\u7a0b\u5206\u4e3a\u4e09\u4e2a\u9636\u6bb5\uff1a', indent=True)

for stage_text, color in [
    ('\u3010\u9636\u6bb5\u2460\u3011 Excel\u6570\u636e\u9884\u5904\u7406\uff08\u7a7a\u767d\u6821\u6b63\u3001\u6291\u5236\u7387\u8ba1\u7b97\u3001\u5f02\u5e38\u5024\u68c0\u9a8c\uff09', RGBColor(0x1F, 0x49, 0x7D)),
    ('\u3010\u9636\u6bb5\u2461\u3011 GraphPad Prism 9.0\u975e\u7ebf\u6027\u56de\u5f52\u62df\u5408\uff084PL\u66f2\u7ebf\uff0c\u83b7\u53d6 IC\u2085\u2080\uff09', RGBColor(0x2E, 0x74, 0xB5)),
    ('\u3010\u9636\u6bb5\u2462\u3011 \u591a\u6b21\u72ec\u7acb\u91cd\u590d\u7684\u7edf\u8ba1\u6c47\u603b\uff08IC\u2085\u2080 mean \u00b1 SD\uff0c95% CI \u8ba1\u7b97\uff09', RGBColor(0x54, 0x82, 0x35)),
]:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    run = p.add_run(stage_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    run.font.color.rgb = color
    set_para_spacing(p, before=0, after=30)

# -------------------------------------------------------
# 4.1 Excel preprocessing
# -------------------------------------------------------
add_chapter_heading(doc, '4.1  \u9636\u6bb5\u2460\uff1aExcel\u6570\u636e\u9884\u5904\u7406', level=2)
add_chapter_heading(doc, '4.1.1  \u539f\u59cb\u6570\u636e\u5bfc\u5165', level=3)
add_body_para(doc,
    '\u9176\u6807\u4eea\u68c0\u6d4b\u5b8c\u6bd5\u540e\uff0c\u5c06 OD450 \u8bfb\u6570\u5bfc\u51fa\u4e3a .xls/.xlsx \u6587\u4ef6\u3002\u4ee5\u4e0b\u662f\u6807\u51c6\u6570\u636e\u6574\u7406\u683c\u5f0f\uff1a',
    indent=True)

add_three_line_table(doc,
    ['\u5206\u7ec4', '\u6d53\u5ea6\uff08\u03bcg/mL\uff09', '\u6d53\u5ea6\u5bf9\u6570\uff08log\u2081\u2080\uff09', '\u5b661 OD', '\u5b662 OD', '\u5b663 OD', '\u5e73\u5747OD', 'CV\uff08%\uff09'],
    [
        ['\u7a7a\u767d\u5bf9\u7167\uff08Blank\uff09', '\u2014', '\u2014', '0.082', '0.079', '0.085', '=AVERAGE(D2:F2)', '=STDEV(D2:F2)/AVERAGE(D2:F2)*100'],
        ['\u9634\u6027\u5bf9\u7167\uff080\u03bcg/mL DMSO\uff09', '0', '\u2014', '0.756', '0.743', '0.761', '=AVERAGE(D3:F3)', '=STDEV(D3:F3)/AVERAGE(D3:F3)*100'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C1', '400', '2.602', '0.521', '0.508', '0.515', '=AVERAGE(D4:F4)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C2', '200', '2.301', '0.588', '0.601', '0.592', '=AVERAGE(D5:F5)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C3', '100', '2.000', '0.645', '0.632', '0.650', '=AVERAGE(D6:F6)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C4', '50', '1.699', '0.693', '0.701', '0.688', '=AVERAGE(D7:F7)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C5', '25', '1.398', '0.718', '0.725', '0.712', '=AVERAGE(D8:F8)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C6', '12.5', '1.097', '0.734', '0.728', '0.741', '=AVERAGE(D9:F9)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C7', '6.25', '0.796', '0.745', '0.739', '0.751', '=AVERAGE(D10:F10)', '...'],
        ['\u82cf\u6728\u63d0\u53d6\u7269 C8', '3.125', '0.495', '0.750', '0.748', '0.755', '=AVERAGE(D11:F11)', '...'],
    ],
    '\u88684-1  \u539f\u59cb\u6570\u636e\u6574\u7406\u6a21\u677f\uff08Excel\u683c\u5f0f\uff09'
)

add_chapter_heading(doc, '4.1.2  \u7a7a\u767d\u6821\u6b63\u4e0e\u6291\u5236\u7387\u8ba1\u7b97', level=3)

add_formula_box(doc, '\u6291\u5236\u7387\u8ba1\u7b97\u516c\u5f0f\uff08\u53c2\u7167 GB/T 7714 \u6807\u51c6\u65b9\u6cd5\uff09', [
    '',
    '\u6291\u5236\u7387\uff08%\uff09= [1 - (OD_\u7ed9\u836f\u7ec4 - OD_\u7a7a\u767d) / (OD_\u9634\u6027\u5bf9\u7167 - OD_\u7a7a\u767d)] x 100%',
    '',
    '\u5176\u4e2d\uff1a',
    'OD_\u7ed9\u836f\u7ec4  = \u836f\u7269\u5904\u7406\u5b54\u7684\u5e73\u5747OD450\u503c',
    'OD_\u7a7a\u767d  = \u7a7a\u767d\u5bf9\u7167\u5b54\uff08\u65e0\u7ec6\u80de\uff0c\u4ec5\u57f9\u517b\u57fa+CCK-8\uff09\u7684\u5e73\u5747OD\u503c',
    'OD_\u9634\u6027\u5bf9\u7167 = 0\u03bcg/mL DMSO\u6eb6\u5242\u5bf9\u7167\u5b54\u7684\u5e73\u5747OD\u503c',
    '',
])

add_body_para(doc, '\u6ce8\u610f\u4e8b\u9879\uff1a', bold_prefix='')
for note in [
    '\u7a7a\u767d\u5bf9\u7167\u5b54\u5e94\u5355\u72ec\u653e\u7f6e\u4e8e96\u5b54\u677f\u975e\u8fb9\u7f18\u533a\u57df\uff08Row B-G, Column 2-11\uff09\uff0c\u5efa\u8bae\u8bbe\u7f566\u4e2a\u91cd\u590d\u5b54\u4ee5\u964d\u4f4e\u8bef\u5dee\uff1b',
    '\u8ba1\u7b97\u7ed3\u679c\u82e5\u51fa\u73b0\u8d1f\u503c\uff08\u5373 OD_\u7ed9\u836f < OD_\u7a7a\u767d\uff09\uff0c\u8bf4\u660e\u53d1\u751f\u4e86\u4e25\u91cd\u7ec6\u80de\u6b7b\u4ea1\uff0c\u8be5\u7ec4\u6291\u5236\u7387\u8bb0\u4e3a100%\uff1b',
    '\u82e5 OD_\u7ed9\u836f\u7ec4 > OD_\u9634\u6027\u5bf9\u7167\uff0c\u6291\u5236\u7387\u4e3a\u8d1f\u503c\uff08\u5373\u4fc3\u8fdb\u589e\u6b96\u6548\u5e94\uff09\uff0c\u9700\u8bb0\u5f55\u5e76\u5728\u8ba8\u8bba\u4e2d\u89e3\u91ca\uff0c\u4f46\u4ecd\u4fdd\u7559\u7528\u4e8e4PL\u62df\u5408\u3002',
]:
    add_body_para(doc, note, indent=False)

add_chapter_heading(doc, '4.1.3  Excel\u64cd\u4f5c\u6b65\u9aa4\uff1a\u6291\u5236\u7387\u6279\u91cf\u8ba1\u7b97', level=3)

add_numbered_step(doc, '1', '\u5efa\u7acb\u8ba1\u7b97\u5217', [
    '\u5728\u539f\u59cb\u6570\u636e\u8868\u4e2d\u65b0\u589e\u4e00\u5217\u3010\u6291\u5236\u7387\uff08%\uff09\u3011',
    '\u5728\u7b2c\u4e00\u4e2a\u836f\u7269\u6d53\u5ea6\u884c\uff08\u5982C2\u884c\uff09\u8f93\u5165\uff1a',
    '=\uff081-\uff08G4-$G$2\uff09/\uff08$G$3-$G$2\uff09\uff09*100',
    '\u5176\u4e2dG\u5217\u4e3a\u5e73\u5747OD\uff0c$G$2\u4e3a\u7a7a\u767d\u5bf9\u7167\u7edd\u5bf9\u5f15\u7528\uff0c$G$3\u4e3a\u9634\u6027\u5bf9\u7167\u7edd\u5bf9\u5f15\u7528'])

add_numbered_step(doc, '2', '\u6279\u91cf\u586b\u5145', [
    '\u9009\u4e2d\u8be5\u5355\u5143\u683c\uff0c\u5411\u4e0b\u62d6\u52a8\u586b\u5145\u81f3\u6240\u6709\u6d53\u5ea6\u7ec4',
    '\u68c0\u67e5\u5404\u6d53\u5ea6\u5bf9\u5e94\u6291\u5236\u7387\u662f\u5426\u5448\u73b0\u6d53\u5ea6\u4f9d\u8d56\u6027\uff08\u6d53\u5ea6\u8d8a\u9ad8\u2192\u6291\u5236\u7387\u8d8a\u9ad8\uff09',
    '\u82e5\u67d0\u6d53\u5ea6\u7684\u6291\u5236\u7387\u4f4e\u4e8e\u524d\u4e00\u6d53\u5ea6\uff08\u975e\u5355\u8c03\uff09\uff0c\u9700\u6807\u8bb0\u5ba1\u67e5'])

add_numbered_step(doc, '3', 'CV\u503c\u8d28\u91cf\u63a7\u5236', [
    '\u8ba1\u7b97\u6bcf\u7ec43\u4e2a\u91cd\u590d\u5b54\u7684CV\uff1a=STDEV(\u51161:\u51163)/AVERAGE(\u51161:\u51163)*100',
    'CV \u2264 10%\uff1a\u63a5\u53d7\uff0c\u6570\u636e\u53ef\u7528',
    'CV 10-15%\uff1a\u6807\u6ce8\u4e3a\u300c\u8fb9\u754c\u300d\uff0c\u5206\u6790\u540e\u8c28\u614e\u4f7f\u7528',
    'CV > 15%\uff1a\u6807\u6ce8\u4e3a\u300c\u5f02\u5e38\u300d\uff0c\u6267\u884cGrubbs\u68c0\u9a8c\uff08\u89c15.2.1\u8282\uff09'])

add_warning_box(doc, '\u6291\u5236\u7387\u8ba1\u7b97\u5e38\u89c1\u9519\u8bef', [
    '\u9519\u8bef1\uff1a\u5fd8\u8bb0\u7a7a\u767d\u6821\u6b63\uff08Blank Subtraction\uff09\u2014\u2014\u76f4\u63a5\u7528OD_\u7ed9\u836f/OD_\u9634\u6027\u5bf9\u7167\uff0c\u5bfc\u81f4IC\u2085\u2080\u504f\u4f4e',
    '\u9519\u8bef2\uff1a\u5c06\u7a7a\u767d\u5bf9\u7167\u884c\u8bef\u8bbe\u4e3a\u9634\u6027\u5bf9\u7167\u2014\u2014\u9020\u6210\u6291\u5236\u7387\u7cfb\u7edf\u6027\u504f\u5dee',
    '\u9519\u8bef3\uff1a\u8fb9\u7f18\u6548\u5e94\u672a\u7ea0\u6b63\u2014\u2014\u8fb9\u7f18\u5b54OD\u504f\u9ad8\uff0c\u9ad8\u4f30\u9634\u6027\u5bf9\u7167\u503c\uff0c\u5bfc\u81f4IC\u2085\u2080\u504f\u9ad8',
    '\u9519\u8bef4\uff1aCV\u8ba1\u7b97\u65f6\u7528\u603b\u6807\u51c6\u5dee\u800c\u975e\u6837\u672c\u6807\u51c6\u5dee\u2014\u2014Excel\u5e94\u4f7f\u7528STDEV\uff08\u800c\u975eSTDEVP\uff09',
], color='FF0000', bg='FFF2CC')

# -------------------------------------------------------
# 4.2 GraphPad Prism
# -------------------------------------------------------
add_chapter_heading(doc, '4.2  \u9636\u6bb5\u2461\uff1aGraphPad Prism 9.0\u64cd\u4f5c\u6307\u5357', level=2)

add_body_para(doc,
    'GraphPad Prism 9.0\u662f\u76ee\u524d\u751f\u547d\u79d1\u5b66\u9886\u57df\u6700\u5e7f\u6cdb\u4f7f\u7528\u7684\u7edf\u8ba1\u7ed8\u56fe\u8f6f\u4ef6\uff0c\u672c\u8282\u63d0\u4f9b\u4ece\u96f6\u5f00\u59cb\u7684\u70b9\u51fb\u5f0f\u64cd\u4f5c\u6307\u5357\u3002'
    '\u672c\u5b9e\u9a8c\u4f7f\u7528\u7684\u6a21\u578b\u4e3a\u3010log(inhibitor) vs. normalized response - Variable slope\u3011'
    '\uff08\u56db\u53c2\u6570 Logistic \u6a21\u578b\uff0c4PL\uff09\u3002',
    indent=True)

add_chapter_heading(doc, '4.2.1  \u8f6f\u4ef6\u754c\u9762\u7b80\u4ecb', level=3)
add_body_para(doc, 'GraphPad Prism 9.0\u4e3b\u754c\u9762\u7531\u4ee5\u4e0b\u533a\u57df\u7ec4\u6210\uff1a', indent=True)

add_three_line_table(doc,
    ['\u754c\u9762\u533a\u57df', '\u529f\u80fd\u8bf4\u660e'],
    [
        ['\u5de6\u4fa7\u5bfc\u822a\u680f\uff08Navigator\uff09', '\u6587\u4ef6\u6811\u7ed3\u6784\uff0c\u663e\u793a\u6240\u6709\u6570\u636e\u8868\u3001\u5206\u6790\u7ed3\u679c\u3001\u56fe\u5f62'],
        ['\u6570\u636e\u8868\u533a\uff08Spreadsheet\uff09', '\u8f93\u5165\u539f\u59cb\u6570\u636e\u7684\u8868\u683c\u533a\u57df\uff0c\u7c7b\u4f3cExcel'],
        ['\u5de5\u5177\u680f\uff08Toolbar\uff09', '\u5305\u62ec[Analyze]\u3001[Change]\u7b49\u5feb\u6377\u6309\u9215'],
        ['\u56fe\u5f62\u9884\u89c8\u533a\uff08Graph\uff09', '\u5b9e\u65f6\u663e\u793a\u5f53\u524d\u6570\u636e\u7684\u56fe\u5f62'],
    ],
    '\u88684-2  GraphPad Prism 9.0\u4e3b\u754c\u9762\u7ec4\u6210\u8bf4\u660e'
)

add_chapter_heading(doc, '4.2.2  \u65b0\u5efa\u9879\u76ee\u6587\u4ef6', level=3)

add_numbered_step(doc, '1', '\u542f\u52a8\u8f6f\u4ef6', [
    '\u53cc\u51fb\u684c\u9762GraphPad Prism 9.0\u56fe\u6807\u542f\u52a8\u8f6f\u4ef6',
    '\u5f39\u51fa[Welcome to Prism]\u5bf9\u8bdd\u6846'])

add_numbered_step(doc, '2', '\u9009\u62e9\u6570\u636e\u7c7b\u578b', [
    '\u5728[Welcome]\u5bf9\u8bdd\u6846\u4e2d\uff0c\u9009\u62e9[XY]\u7c7b\u578b\uff08\u5de6\u4fa7\u7b2c\u4e00\u4e2a\u56fe\u6807\uff09',
    '\u6b64\u7c7b\u578b\u9002\u7528\u4e8eX\u8f74\u4e3a\u8fde\u7eed\u53d8\u91cf\uff08\u6d53\u5ea6\u5bf9\u6570\uff09\u3001Y\u8f74\u4e3a\u6291\u5236\u7387\u7684\u5242\u91cf-\u6548\u5e94\u5206\u6790'])

add_numbered_step(doc, '3', '\u8bbe\u7f6eX\u8f74\u683c\u5f0f', [
    '\u5728[What are X data?]\u90e8\u5206\uff0c\u9009\u62e9[Numbers]',
    'X\u8f74\u6570\u636e\u5c06\u8f93\u5165 log10(\u6d53\u5ea6/\u03bcg/mL)\uff0c\u5982 log10(400)=2.602'])

add_numbered_step(doc, '4', '\u8bbe\u7f6eY\u8f74\u683c\u5f0f', [
    '\u5728[Enter Y values]\u90e8\u5206\uff0c\u9009\u62e9[Enter and plot a mean with SD and N]',
    '\u6b64\u683c\u5f0f\u5141\u8bb8\u76f4\u63a5\u8f93\u51653\u6b21\u91cd\u590d\u7684\u539f\u59cb\u6570\u636e\uff0cPrism\u81ea\u52a8\u8ba1\u7b97\u5747\u503c\u00b1SD',
    '\u6216\u9009\u62e9[Enter replicate values in side-by-side subcolumns]\u76f4\u63a5\u8f93\u51653\u5217\u91cd\u590d\u503c'])

add_numbered_step(doc, '5', '\u547d\u540d\u6587\u4ef6', [
    '\u5c06\u9879\u76ee\u540d\u79f0\u8bbe\u4e3a\uff1a\u82cf\u6728-CNE2-CCK8-IC50\u5206\u6790-24h\uff08\u635648h/72h\uff09',
    '\u70b9\u51fb[Create]\u8fdb\u5165\u6570\u636e\u8868'])

add_chapter_heading(doc, '4.2.3  \u6570\u636e\u8f93\u5165', level=3)
add_body_para(doc, '\u8fdb\u5165\u6570\u636e\u8868\u540e\uff0c\u6309\u4ee5\u4e0b\u89c4\u8303\u8f93\u5165\u6570\u636e\uff1a', indent=True)

add_three_line_table(doc,
    ['X\uff08log\u6d53\u5ea6\uff09', 'Y1\uff1aRep1\u6291\u5236\u7387\uff08%\uff09', 'Y2\uff1aRep2\u6291\u5236\u7387\uff08%\uff09', 'Y3\uff1aRep3\u6291\u5236\u7387\uff08%\uff09', '\u8bf4\u660e'],
    [
        ['0.495', '2.5', '3.1', '2.8', '3.125 \u03bcg/mL'],
        ['0.796', '5.3', '6.2', '5.7', '6.25 \u03bcg/mL'],
        ['1.097', '12.1', '11.5', '12.8', '12.5 \u03bcg/mL'],
        ['1.398', '23.4', '22.8', '24.1', '25 \u03bcg/mL'],
        ['1.699', '41.2', '40.5', '42.3', '50 \u03bcg/mL'],
        ['2.000', '58.9', '57.6', '60.1', '100 \u03bcg/mL'],
        ['2.301', '72.5', '71.8', '73.4', '200 \u03bcg/mL'],
        ['2.602', '84.3', '83.9', '85.1', '400 \u03bcg/mL'],
    ],
    '\u88684-3  GraphPad Prism 9.0\u6570\u636e\u8f93\u5165\u793a\u4f8b\uff08\u6a21\u62df\u6570\u636e\uff0c\u4ec5\u4f9b\u6f14\u793a\uff09'
)

add_warning_box(doc, '\u6570\u636e\u8f93\u5165\u6ce8\u610f\u4e8b\u9879', [
    'X\u8f74\u5fc5\u987b\u8f93\u5165 log10(\u6d53\u5ea6) \u800c\u975e\u539f\u59cb\u6d53\u5ea6\u2014\u2014\u82e5\u8f93\u5165\u539f\u59cb\u6d53\u5ea6\u9700\u989d\u5916\u8bbe\u7f6eX\u8f74\u5bf9\u6570\u523b\u5ea6\uff0c\u6613\u51fa\u9519',
    '0\u03bcg/mL\u5bf9\u7167\u7ec4\uff08DMSO\uff09\u4e0d\u5e94\u51fa\u73b0\u5728\u62df\u5408\u6570\u636e\u4e2d\u2014\u2014\u5176log\u503c\u65e0\u6cd5\u5b9a\u4e49',
    '\u6291\u5236\u7387\u5fc5\u987b\u662f\u767e\u5206\u6bd5\u7b26\u53f7\uff080-100\uff09\uff0c\u800c\u975e\u5c0f\u6570\uff080-1\uff09\u2014\u2014\u907f\u514dIC\u2085\u2080\u5355\u4f4d\u9519\u8bef',
    '\u786e\u8ba43\u5217\u91cd\u590d\u503c\u5747\u5df2\u8f93\u5165\u2014\u2014\u82e5\u6709\u7f3a\u5931\uff08CV>15%\u88ab\u6392\u9664\u7684\u5b54\uff09\uff0c\u9700\u5728\u76f8\u5e94\u5217\u8f93\u5165N/A\u6216\u7559\u7a7a',
], color='CC4400', bg='FFF2CC')

add_chapter_heading(doc, '4.2.4  \u8fd0\u884c\u975e\u7ebf\u6027\u56de\u5f52\uff084PL\u6a21\u578b\uff09', level=3)
add_body_para(doc, '\u8fd9\u662f\u83b7\u53d6 IC\u2085\u2080 \u7684\u5173\u952e\u6b65\u9aa4\uff0c\u8bf7\u4e25\u683c\u6309\u7167\u4ee5\u4e0b\u64cd\u4f5c\u6267\u884c\uff1a', indent=True)

add_numbered_step(doc, '1', '\u70b9\u51fbAnalyze\u6309\u9215', [
    '\u5728\u9876\u90e8\u5de5\u5177\u680f\u70b9\u51fb[Analyze]\u6309\u9215\uff08\u5feb\u6377\u952eCtrl+Y\uff09',
    '\u5f39\u51fa[Analyze Data]\u5bf9\u8bdd\u6846'])

add_numbered_step(doc, '2', '\u9009\u62e9\u5206\u6790\u7c7b\u578b', [
    '\u5728\u5de6\u4fa7\u5217\u8868\u4e2d\u627e\u5230[XY analyses]',
    '\u9009\u62e9[Nonlinear regression (curve fit)]',
    '\u70b9\u51fb[OK]\u8fdb\u5165\u6a21\u578b\u9009\u62e9\u754c\u9762'])

add_numbered_step(doc, '3', '\u9009\u62e9\u62df\u5408\u6a21\u578b', [
    '\u5728[Choose an equation]\u641c\u7d22\u6846\u4e2d\u8f93\u5165\uff1adose',
    '\u5728\u641c\u7d22\u7ed3\u679c\u4e2d\u627e\u5230\u5e76\u9009\u62e9\uff1a',
    '[Dose-response -- Inhibition > log(inhibitor) vs. normalized response -- Variable slope]',
    '\u3010\u91cd\u8981\u3011\u6b64\u5373\u56db\u53c2\u6570Logistic\uff084PL\uff09\u6a21\u578b\uff0c\u4e5f\u79f0Hill\u65b9\u7a0b'])

add_numbered_step(doc, '4', '\u68c0\u67e5\u6a21\u578b\u53c2\u6570\u8bbe\u7f6e', [
    '\u786e\u8ba4\u4ee5\u4e0b\u53c2\u6570\u8bbe\u7f6e\uff1a',
    '  Bottom\uff08Ymin\uff09\uff1a0\uff08\u56fa\u5b9a\uff0c\u4ee3\u8868100%\u6291\u5236\u540e\u6b8b\u4f59\u4fe1\u53f7\u4e3a0\uff09',
    '  Top\uff08Ymax\uff09\uff1a100\uff08\u56fa\u5b9a\uff0c\u4ee3\u8868\u65e0\u836f\u65f6\u6291\u5236\u7387=0%\uff09',
    '  HillSlope\uff1a\u6d6e\u52a8\uff08Prism\u81ea\u52a8\u62df\u5408\uff0c\u901a\u5e380.5-3\u4e4b\u95f4\uff09',
    '  LogEC50\uff1a\u6d6e\u52a8\uff08\u5373log10(IC\u2085\u2080)\uff0cPrism\u81ea\u52a8\u62df\u5408\uff09',
    '\u3010\u8bf4\u660e\u3011Bottom=0\u4e14Top=100\u662f\u300c\u5f52\u4e00\u5316\u54cd\u5e94\u300d\u6a21\u578b\uff0c\u9002\u7528\u4e8e\u5df2\u5c06\u5bf9\u7167\u7ec4\u5f52\u4e00\u5316\u7684\u6291\u5236\u7387\u6570\u636e'])

add_numbered_step(doc, '5', '\u9009\u9879\u8bbe\u7f6e', [
    '\u5c55\u5f00[Options]\u9762\u677f\uff0c\u8fdb\u884c\u4ee5\u4e0b\u8bbe\u7f6e\uff1a',
    'Confidence Interval\uff1a\u9009\u62e9[95% CI]',
    'Method\uff1a\u4fdd\u6301[Least squares regression]',
    'Initial values\uff1a\u9009\u62e9[Automatic]\uff08Prism\u81ea\u52a8\u4f30\u7b97\u521d\u59cb\u53c2\u6570\uff09'])

add_numbered_step(doc, '6', '\u8fd0\u884c\u62df\u5408', [
    '\u70b9\u51fb[OK]\u6267\u884c\u975e\u7ebf\u6027\u56de\u5f52',
    'Prism\u81ea\u52a8\u751f\u6210Results\u8868\u683c\u548c\u62df\u5408\u66f2\u7ebf\u56fe'])

add_tip_box(doc, '\u4ec0\u4e48\u662f4PL\u56db\u53c2\u6570Logistic\u6a21\u578b\uff1f', [
    '4PL\u6a21\u578b\u65b9\u7a0b\u4e3a\uff1a',
    'Y = Bottom + (Top - Bottom) / [1 + 10^((LogEC50 - X) x HillSlope)]',
    '\u56db\u4e2a\u53c2\u6570\u7684\u542b\u4e49\uff1a',
    '  Bottom\uff08Ymin\uff09\uff1aS\u5f62\u66f2\u7ebf\u7684\u4e0b\u6e10\u8fd1\u7ebf\uff08\u6700\u5927\u6291\u5236\u6548\u5e94\uff09',
    '  Top\uff08Ymax\uff09\uff1aS\u5f62\u66f2\u7ebf\u7684\u4e0a\u6e10\u8fd1\u7ebf\uff08\u65e0\u836f\u7269\u65f6\u7684\u57fa\u7ebf\uff09',
    '  LogEC50\uff1a\u66f2\u7ebf\u4e2d\u70b9\u5904\u7684\u6d53\u5ea6\u5bf9\u6570\u503c\uff08\u5373 log10(IC\u2085\u2080)\uff09',
    '  HillSlope\uff08Hill\u7cfb\u6570\uff09\uff1a\u63cf\u8ff0S\u5f62\u66f2\u7ebf\u7684\u9661\u5ce8\u7a0b\u5ea6',
    'IC\u2085\u2080 = 10^(LogEC50)\uff0c\u5355\u4f4d\u4e0e\u8f93\u5165X\u7684\u6d53\u5ea6\u5355\u4f4d\u4e00\u81f4\uff08\u672c\u5b9e\u9a8c\u4e3a\u03bcg/mL\uff09',
])

add_chapter_heading(doc, '4.2.5  \u8bfb\u53d6IC\u2085\u2080\u62df\u5408\u7ed3\u679c', level=3)
add_body_para(doc, '\u62df\u5408\u5b8c\u6210\u540e\uff0cPrism\u7684Results\u6807\u7b7e\u9875\u5c06\u663e\u793a\u8be6\u7ec6\u7684\u7edf\u8ba1\u7ed3\u679c\u3002\u5173\u952e\u53c2\u6570\u8bfb\u53d6\u65b9\u6cd5\u5982\u4e0b\uff1a', indent=True)

add_three_line_table(doc,
    ['Results\u53c2\u6570', '\u53c2\u6570\u8bf4\u660e', '\u5224\u8bfb\u6807\u51c6', '\u793a\u4f8b\u503c'],
    [
        ['LogEC50', 'IC\u2085\u2080\u7684\u5bf9\u6570\u503c', '\u2014', '1.871'],
        ['IC\u2085\u2080\uff08antilog\uff09', 'IC\u2085\u2080\u503c\uff08Prism\u81ea\u52a8\u6362\u7b97\uff09', '\u6570\u503c\u5408\u7406\u6027\u9a8c\u8bc1', '74.3 \u03bcg/mL'],
        ['HillSlope', 'Hill\u7cfb\u6570', '\u901a\u5e380.5~3\uff0c\u82e5>5\u63d0\u793a\u6a21\u578b\u4e0d\u5408\u9002', '1.24'],
        ['Bottom', '\u66f2\u7ebf\u4e0b\u6e10\u8fd1\u7ebf', '\u82e5\u56fa\u5b9a\u4e3a0\uff0c\u5e94\u663e\u793a0', '0 (fixed)'],
        ['Top', '\u66f2\u7ebf\u4e0a\u6e10\u8fd1\u7ebf', '\u82e5\u56fa\u5b9a\u4e3a100\uff0c\u5e94\u663e\u793a100', '100 (fixed)'],
        ['Goodness of fit: R2', '\u62df\u5408\u4f18\u5ea6\uff08\u8d8a\u63a51\u8d8a\u597d\uff09', 'R2 >= 0.95 \u4e3a\u5408\u683c', '0.9873'],
        ['Goodness of fit: RMSE', '\u5747\u65b9\u6839\u8bef\u5dee', 'RMSE\u8d8a\u5c0f\u8d8a\u597d\uff0c<5%\u53ef\u63a5\u53d7', '2.14'],
        ['95% CI of LogEC50', 'LogEC50\u795e95%\u7f6e\u4fe1\u533a\u95f4', 'CI\u8f83\u7a84\u8868\u660e\u62df\u5408\u7a33\u5065', '[1.843, 1.899]'],
        ['95% CI of IC50', 'IC\u2085\u2080\u795e95%\u7f6e\u4fe1\u533a\u95f4\uff08\u03bcg/mL\uff09', '\u2014', '[69.6, 79.3]'],
    ],
    '\u88684-4  GraphPad Prism\u975e\u7ebf\u6027\u56de\u5f52\u5173\u952e\u7ed3\u679c\u53c2\u6570\u8bf4\u660e'
)

add_warning_box(doc, 'R2<0.95\u65f6\u7684\u5904\u7406\u6d41\u7a0b', [
    '1. \u68c0\u67e5\u662f\u5426\u6709\u660e\u663e\u5f02\u5e38\u70b9\uff1a\u67e5\u770b\u6b8b\u5dee\u56fe\uff08Residuals plot\uff09\uff0c\u5bfb\u627e\u504f\u79bb>2SD\u7684\u6570\u636e\u70b9',
    '2. \u68c0\u67e5\u6570\u636e\u5355\u8c03\u6027\uff1a\u5982\u679c\u5728\u67d0\u6d53\u5ea6\u51fa\u73b0\u6291\u5236\u7387\u53cd\u5f39\uff0c\u53ef\u80fd\u662f\u836f\u7269\u6c89\u6de0\u6216\u5b9e\u9a8c\u64cd\u4f5c\u95ee\u9898',
    '3. \u68c0\u67e5Hill\u7cfb\u6570\uff1a\u82e5HillSlope < 0\uff08\u659c\u7387\u65b9\u5411\u9519\u8bef\uff09\uff0c\u8bf4\u660e\u6570\u636e\u8f93\u5165\u65b9\u5411\u6709\u8bef',
    '4. \u8003\u8651\u91ca\u653eBottom\u7ea6\u675f\uff1a\u5c1d\u8bd5\u8ba9Bottom\u81ea\u7531\u6d6e\u52a8\uff0c\u89c2\u5bdfR2\u662f\u5426\u663e\u8457\u63d0\u5347',
    '5. \u6700\u7ec8\u65e0\u6cd5\u8fbe\u5230R2>=0.95\u65f6\uff1a\u8865\u5145\u66f4\u9ad8\u6d53\u5ea6\u7ec4\uff08C0: 800\u03bcg/mL\u62161600\u03bcg/mL\uff09\u6269\u5c55\u66f2\u7ebf\u8303\u56f4',
], color='FF0000', bg='FFF2CC')

add_chapter_heading(doc, '4.2.6  \u5242\u91cf-\u6548\u5e94\u66f2\u7ebf\u7684\u5bfc\u51fa\u4e0e\u7f8e\u5316', level=3)

add_numbered_step(doc, '1', '\u4fee\u6539X\u8f74\u6807\u7b7e', [
    '\u53cc\u51fb\u56fe\u5f62\u7684X\u8f74\uff0c\u5f39\u51fa[Format Axes]\u5bf9\u8bdd\u6846',
    '\u5c06\u8f74\u6807\u9898\u6539\u4e3a\uff1alog10(\u82cf\u6728\u63d0\u53d6\u7269\u6d53\u5ea6 / \u03bcg/mL)',
    '\u786e\u8ba4\u6570\u5024\u8303\u56f4\u8986\u76d6\u6240\u6709\u5b9e\u9a8c\u6d53\u5ea6\u70b9'])

add_numbered_step(doc, '2', '\u4fee\u6539Y\u8f74\u6807\u7b7e', [
    '\u5c06Y\u8f74\u6807\u9898\u6539\u4e3a\uff1a\u7ec6\u80de\u589e\u6b96\u6291\u5236\u7387\uff08%\uff09',
    'Y\u8f74\u8303\u56f4\u8bbe\u4e3a -5 \u81f3 110\uff08\u7559\u51fa\u4f59\u91cf\uff09',
    '\u5728Y\u8f7450%\u5904\u6dfb\u52a0\u6c34\u5e73\u865a\u7ebf\uff08\u6807\u6ce8IC50\u4f4d\u7f6e\uff09'])

add_numbered_step(doc, '3', '\u6dfb\u52a0IC50\u6807\u6ce8', [
    '\u70b9\u51fb[Insert > Annotation]\uff08\u6ce8\u91ca\u5de5\u5177\uff09',
    '\u5728IC50\u5bf9\u5e94\u7684\u66f2\u7ebf\u4f4d\u7f6e\u6dfb\u52a0\u6587\u5b57\uff1aIC50 = XX.X ug/mL',
    '\u7528\u7bad\u5934\u6307\u5411\u66f2\u7ebf\u4e0eY=50%\u7684\u4ea4\u53c9\u70b9'])

add_numbered_step(doc, '4', '\u56fe\u5f62\u7f8e\u5316', [
    '\u8bbe\u7f6e\u6570\u636e\u70b9\u6837\u5f0f\uff1a\u5b9e\u5fc3\u5706\u5f62\uff08\u25cf\uff09\uff0c\u5927\u5c0f12pt',
    '\u8bef\u5dee\u7ebf\uff1a\u9009\u62e9SD\uff08\u6807\u51c6\u5dee\uff09',
    '\u62df\u5408\u66f2\u7ebf\uff1a\u5b9e\u7ebf\uff0c\u5bbd\u5ea61.5pt\uff0c\u4e0e\u6570\u636e\u70b9\u540c\u8272',
    '\u56fe\u4f8b\uff1a\u53f3\u4e0a\u89d2\uff0c\u5305\u62ec\u300c\u82cf\u6728\u63d0\u53d6\u7269\uff0824h\uff09\u300d\u7b49\u6807\u6ce8'])

add_numbered_step(doc, '5', '\u5bfc\u51fa\u56fe\u5f62', [
    '\u70b9\u51fb[File > Export]',
    '\u683c\u5f0f\u9009\u62e9\uff1aTIFF\uff08300 DPI\uff09\u7528\u4e8e\u8bba\u6587\u6295\u7a3f\uff0c\u6216PNG\uff08150 DPI\uff09\u7528\u4e8e\u62a5\u544a',
    '\u6587\u4ef6\u547d\u540d\uff1a\u82cf\u6728-CNE2-CCK8-\u5242\u91cf\u6548\u5e94\u66f2\u7ebf-24h-\u65e5\u671f.tiff'])

# -------------------------------------------------------
# 4.3 Multi-experiment statistics
# -------------------------------------------------------
add_chapter_heading(doc, '4.3  \u9636\u6bb5\u2462\uff1a\u591a\u6b21\u72ec\u7acb\u91cd\u590d\u7684\u7edf\u8ba1\u6c47\u603b', level=2)

add_body_para(doc,
    '\u6839\u636e\u5b9e\u9a8c\u65b9\u6848\u8981\u6c42\uff0c\u672c\u5b9e\u9a8c\u9700\u8fdb\u884c\u22653\u6b21\u5b8c\u5168\u72ec\u7acb\u91cd\u590d\uff08Independent Biological Replicates\uff09\u3002'
    '\u6bcf\u6b21\u72ec\u7acb\u91cd\u590d\u5747\u9700\u8fdb\u884c\u4e00\u6b214PL\u62df\u5408\uff0c\u83b7\u5f97\u5404\u81ea\u7684 IC\u2085\u2080 \u503c\uff0c'
    '\u6700\u7ec8\u62a5\u544a3\u6b21\u91cd\u590d\u7684\u5747\u503c\u00b1\u6807\u51c6\u5dee\uff08Mean \u00b1 SD\uff09\u3002',
    indent=True)

add_chapter_heading(doc, '4.3.1  \u72ec\u7acb\u91cd\u590d\u7684\u5b9a\u4e49', level=3)
add_body_para(doc,
    '\u5728\u672c\u5b9e\u9a8c\u4e2d\uff0c\u4e00\u6b21\u300c\u72ec\u7acb\u751f\u7269\u5b66\u91cd\u590d\u300d\u5b9a\u4e49\u4e3a\uff1a\u91cd\u65b0\u590d\u82cf\u7ec6\u80de\uff08\u6216\u4ece\u4e0d\u540c\u4f20\u4ee3\u6279\u6b21\u53d6\u6837\uff09'
    '\u2192\u5b8c\u6574\u57f9\u517b\u2192\u72ec\u7acb\u7ed9\u836f\u5904\u7406\u2192\u72ec\u7acb CCK-8\u68c0\u6d4b\u2192\u72ec\u7acb\u6570\u636e\u5206\u6790\u7684\u5b8c\u6574\u6d41\u7a0b\u3002'
    '\u540c\u4e00\u6b21\u7ec6\u80de\u57f9\u517b\u4e2d\u91cd\u590d\u94fa\u677f\uff08Technical Replicates\uff09\u4e0d\u7b49\u4e8e\u72ec\u7acb\u91cd\u590d\u3002',
    indent=True)
add_body_para(doc,
    '\u6bcf\u6b21\u72ec\u7acb\u91cd\u590d\u4e2d\uff0c\u6bcf\u4e2a\u6d53\u5ea6\u7ec4\u8bbe\u7f6e3\u4e2a\u6280\u672f\u91cd\u590d\u5b54\uff08n=3 wells\uff09\uff0c'
    '\u7528\u4e8e\u8ba1\u7b97\u7ec4\u5185CV\uff1b3\u6b21\u72ec\u7acb\u91cd\u590d\u7528\u4e8e\u8ba1\u7b97\u7ec4\u95f4\u53d8\u5f02\u548c\u6700\u7ec8\u62a5\u544a\u7684 IC\u2085\u2080 \u00b1 SD\u3002',
    indent=True)

add_chapter_heading(doc, '4.3.2  IC50\u7edf\u8ba1\u6c47\u603b\u8868', level=3)

add_three_line_table(doc,
    ['\u72ec\u7acb\u91cd\u590d\u6b21\u6570', '24h IC50\uff08\u03bcg/mL\uff09', '48h IC50\uff08\u03bcg/mL\uff09', '72h IC50\uff08\u03bcg/mL\uff09', 'R2\uff08\u5404\u65f6\u95f4\u70b9\uff09'],
    [
        ['\u91cd\u590d1', '\u2014', '\u2014', '\u2014', '\uff08\u5e94>=0.95\uff09'],
        ['\u91cd\u590d2', '\u2014', '\u2014', '\u2014', '\uff08\u5e94>=0.95\uff09'],
        ['\u91cd\u590d3', '\u2014', '\u2014', '\u2014', '\uff08\u5e94>=0.95\uff09'],
        ['Mean \u00b1 SD', '\u2014', '\u2014', '\u2014', '\u2014'],
        ['95% CI', '\u2014', '\u2014', '\u2014', '\u2014'],
    ],
    '\u88684-5  IC50\u591a\u6b21\u72ec\u7acb\u91cd\u590d\u7edf\u8ba1\u6c47\u603b\uff08\u5f85\u586b\u5199\uff09'
)

add_formula_box(doc, 'IC50\u7edf\u8ba1\u62a5\u544a\u683c\u5f0f\u8981\u6c42', [
    '',
    '\u62a5\u544a\u683c\u5f0f\u793a\u4f8b\uff1a',
    'IC50\uff0824h\uff09= XX.X +/- X.X ug/mL (n = 3, Mean +/- SD)',
    'IC50\uff0848h\uff09= XX.X +/- X.X ug/mL (n = 3, Mean +/- SD)',
    'IC50\uff0872h\uff09= XX.X +/- X.X ug/mL (n = 3, Mean +/- SD)',
    '',
    '95% CI\u4f7f\u7528\u516c\u5f0f\uff1aIC50 +/- t(0.025, df=2) x SD / sqrt(n)',
    '\u5176\u4e2d t(df=2, \u53cc\u5c3e) = 4.303 (\u67e5t\u5206\u5e03\u8868)',
    '',
])

add_chapter_heading(doc, '4.3.3  \u65f6\u95f4\u4f9d\u8d56\u6027\u5206\u6790', level=3)
add_body_para(doc, '\u83b7\u5f97\u4e09\u4e2a\u65f6\u95f4\u70b9IC50\u540e\uff0c\u5206\u6790\u82cf\u6728\u63d0\u53d6\u7269\u5bf9CNE-2\u7ec6\u80de\u7684\u65f6\u95f4\u4f9d\u8d56\u6027\u589e\u6b96\u6291\u5236\u89c4\u5f8b\uff1a', indent=True)
for bullet_text in [
    'IC50(24h) > IC50(48h) > IC50(72h)\uff1a\u8868\u660e\u82cf\u6728\u63d0\u53d6\u7269\u5177\u6709\u65f6\u95f4\u4f9d\u8d56\u6027\u6291\u5236\u4f5c\u7528\uff0c\u968f\u5904\u7406\u65f6\u95f4\u5ef6\u957f\uff0c\u6240\u9700\u6291\u523650%\u7ec6\u80de\u6240\u9700\u6d53\u5ea6\u9010\u6e10\u964d\u4f4e',
    '\u82e5IC50(72h) \u2248 IC50(48h) >> IC50(24h)\uff1a\u63d0\u793a\u5728\u4e8e48-72h\u8303\u56f4\u5185\u7ec6\u80de\u53ef\u80fd\u8fdb\u5165\u7ec6\u80de\u5468\u671f\u963b\u6ede\u6216\u8fbe\u5230\u6700\u5927\u6291\u5236',
    '\u82e5\u4e09\u4e2a\u65f6\u95f4\u70b9IC50\u5dee\u5f02<10%\uff1a\u63d0\u793a\u82cf\u6728\u63d0\u53d6\u7269\u7684\u7ec6\u80de\u6bd2\u6027\u6548\u5e94\u5728\u4e8e24h\u5185\u5df2\u57fa\u672c\u5b8c\u6210\uff0c\u5c5e\u5feb\u901f\u7ec6\u80de\u6bd2\u673a\u5236',
]:
    add_body_para(doc, bullet_text, indent=False)

# -------------------------------------------------------
# 4.4 Data verification checklist
# -------------------------------------------------------
add_chapter_heading(doc, '4.4  \u6570\u636e\u5206\u6790\u8d28\u91cf\u9a8c\u8bc1\u6e05\u5355', level=2)
add_body_para(doc, '\u5b8c\u6210\u6570\u636e\u5206\u6790\u540e\uff0c\u5728\u63d0\u4ea4\u5b9e\u9a8c\u62a5\u544a\u524d\uff0c\u8bf7\u9010\u9879\u6838\u5bf9\u4ee5\u4e0b\u8d28\u91cf\u9a8c\u8bc1\u6e05\u5355\uff1a', indent=True)

checklist_items = [
    ('OD\u7a7a\u767d\u6821\u6b63\u5df2\u5b8c\u6210', '\u6240\u6709\u6570\u636e\u5df2\u51cf\u53bb\u7a7a\u767dOD\u5747\u503c', '\u2610 \u5df2\u5b8c\u6210'),
    ('\u9634\u6027\u5bf9\u7167CV \u2264 10%', 'CV = ____%\uff0c\u5e94\u22641 0%', '\u2610 \u901a\u8fc7'),
    ('\u6bcf\u7ec4\u836f\u7269\u6d53\u5ea6CV \u2264 10%', '\u68c0\u67e58\u4e2a\u6d53\u5ea6\u7ec4CV\u503c', '\u2610 \u901a\u8fc7'),
    ('\u6291\u5236\u7387\u5448\u5355\u8c03\u9012\u589e\u8d8b\u52bf', '\u9ad8\u6d53\u5ea6\u2192\u9ad8\u6291\u5236\u7387', '\u2610 \u901a\u8fc7'),
    ('4PL\u62df\u5408R2 >= 0.95', 'R2 = ____\uff08\u4e09\u4e2a\u65f6\u95f4\u70b9\u5747\u5e94>=0.95\uff09', '\u2610 \u901a\u8fc7'),
    ('HillSlope\u8303\u56f4\u5408\u7406', 'HillSlope = ____ (0.5~3\u4e4b\u95f4)', '\u2610 \u901a\u8fc7'),
    ('IC50\u5728\u6d53\u5ea6\u8303\u56f4\u5185', 'IC50\u4e0d\u8d85\u51fa\u5b9e\u9a8c\u6d53\u5ea6\u8303\u56f4', '\u2610 \u901a\u8fc7'),
    ('\u5df2\u5b8c\u62103\u6b21\u72ec\u7acb\u91cd\u590d', '3\u6b21\u91cd\u590dIC50\u5747\u5df2\u83b7\u5f97', '\u2610 \u901a\u8fc7'),
    ('IC50 +/- SD\u5df2\u8ba1\u7b97', 'Mean +/- SD\u683c\u5f0f\u786e\u8ba4', '\u2610 \u5b8c\u6210'),
    ('\u56fe\u5f62\u683c\u5f0f\u89c4\u8303', '300 DPI\uff0c\u8f74\u6807\u7b7e\u5b8c\u6574\uff0cIC50\u6807\u6ce8\u6e05\u6670', '\u2610 \u901a\u8fc7'),
    ('\u539f\u59cb\u6570\u636e\u5df2\u5b58\u6863', 'Excel+Prism\u6587\u4ef6\u4fdd\u5b58\u81f3\u6307\u5b9a\u6587\u4ef6\u5939', '\u2610 \u5b8c\u6210'),
    ('\u5b9e\u9a8c\u8bb0\u5f55\u672c\u5df2\u66f4\u65b0', '\u6240\u6709IC50\u503c\u7b7e\u5b57\u8bb0\u5f55', '\u2610 \u5b8c\u6210'),
]

checklist_table = doc.add_table(rows=1 + len(checklist_items), cols=3)
checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(['\u9a8c\u8bc1\u9879\u76ee', '\u6807\u51c6\u8981\u6c42', '\u5b8c\u6210\u72b6\u6001']):
    cell = checklist_table.rows[0].cells[i]
    set_cell_bg(cell, 'D6E4F7')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = '\u5b8b\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

for r_i, (item, std, status) in enumerate(checklist_items):
    row = checklist_table.rows[r_i + 1]
    bg = 'FFFFFF' if r_i % 2 == 0 else 'F5F9FF'
    for c_i, val in enumerate([item, std, status]):
        cell = row.cells[c_i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10.5)
        r.font.name = '\u5b8b\u4f53'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

doc.add_paragraph()
add_body_para(doc, '\u6ce8\uff1a\u8be5\u6e05\u5355\u9700\u5728\u6bcf\u6b21\u72ec\u7acb\u91cd\u590d\u5b8c\u6210\u6570\u636e\u5206\u6790\u540e\u6253\u5370\u5e76\u586b\u5199\uff0c\u7b7e\u5b57\u540e\u9644\u4e8e\u5b9e\u9a8c\u8bb0\u5f55\u672c\u4e2d\u3002', indent=True)

# -------------------------------------------------------
# 4.5 Prism file management
# -------------------------------------------------------
add_chapter_heading(doc, '4.5  Prism\u9879\u76ee\u6587\u4ef6\u7ba1\u7406\u89c4\u8303', level=2)
add_body_para(doc,
    'GraphPad Prism\u9879\u76ee\u6587\u4ef6\uff08.pzfx\u683c\u5f0f\uff09\u5305\u542b\u6240\u6709\u539f\u59cb\u6570\u636e\u3001\u5206\u6790\u7ed3\u679c\u548c\u56fe\u5f62\uff0c'
    '\u662f\u5b9e\u9a8c\u6570\u636e\u6eaf\u6e90\u7684\u91cd\u8981\u4f9d\u636e\uff0c\u9700\u6309\u4ee5\u4e0b\u89c4\u8303\u7ba1\u7406\uff1a',
    indent=True)

add_three_line_table(doc,
    ['\u6587\u4ef6\u7c7b\u578b', '\u6587\u4ef6\u547d\u540d\u89c4\u8303', '\u5b58\u653e\u4f4d\u7f6e', '\u5907\u6ce8'],
    [
        ['Prism\u9879\u76ee\u6587\u4ef6', '\u82cf\u6728-CNE2-IC50-\u72ec\u7acb\u91cd\u590dN-\u65e5\u671f.pzfx', '\u6570\u636e\u6587\u4ef6\u5939/Prism/', '\u6bcf\u6b21\u91cd\u590d\u5355\u72ec\u4fdd\u5b58'],
        ['\u5bfc\u51faIC50\u7edf\u8ba1\u56fe', '\u82cf\u6728-CNE2-\u5242\u91cf\u6548\u5e94\u66f2\u7ebf-NNh-\u65e5\u671f.tiff', '\u6570\u636e\u6587\u4ef6\u5939/Figures/', '300 DPI TIFF\u683c\u5f0f'],
        ['Excel\u9884\u5904\u7406\u8868', '\u82cf\u6728-CNE2-CCK8-\u539f\u59cb\u6570\u636e-\u65e5\u671f.xlsx', '\u6570\u636e\u6587\u4ef6\u5939/RawData/', '\u5305\u542b\u516c\u5f0f\uff0c\u52ff\u5220\u9664'],
        ['\u6570\u636e\u6c47\u603b\u8868', '\u82cf\u6728-CNE2-IC50\u6c47\u603b-\u4e09\u6b21\u91cd\u590d.xlsx', '\u6570\u636e\u6587\u4ef6\u5939/Summary/', '\u586b\u5165\u6240\u6709\u91cd\u590d\u7ed3\u679c'],
    ],
    '\u88684-6  \u6570\u636e\u6587\u4ef6\u7ba1\u7406\u547d\u540d\u89c4\u8303'
)

add_tip_box(doc, 'GraphPad Prism\u6587\u4ef6\u5907\u4efd\u5efa\u8bae', [
    '\u6bcf\u6b21\u5b8c\u6210\u5206\u6790\u540e\u7acb\u5373\u53e6\u5b58\u4e3a\uff08File > Save As\uff09\uff0c\u4fdd\u7559\u7248\u672c\u53f7\u540e\u7f00',
    '\u5c06Prism\u6587\u4ef6\u540c\u6b65\u81f3\u4e91\u76d8\uff08\u5982\u575a\u679c\u4e91/iCloud\uff09\u6216\u4f7f\u7528U\u76d8\u53cc\u5907\u4efd',
    '\u4e0d\u8981\u4fee\u6539\u5df2\u786e\u8ba4\u7684\u539f\u59cb\u6570\u636e\u8868\uff0c\u82e5\u9700\u91cd\u65b0\u5206\u6790\uff0c\u590d\u5236\u6570\u636e\u8868\u65b0\u5efa\u5206\u6790',
    '\u5bfc\u5e08\u5ba1\u9605\u65f6\u901a\u5e38\u9700\u8981\u67e5\u770bPrism\u539f\u59cb\u6587\u4ef6\uff0c\u8bf7\u786e\u4fdd\u6587\u4ef6\u5b8c\u6574\u53ef\u6253\u5f00',
])

doc.save(DOC_PATH)
print("Chapter 4 complete and saved.")
