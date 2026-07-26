#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Chapter 5: Quality Control, Red Lines, and Troubleshooting Guide
Appends to existing v2.0 document
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "final/\u82cf\u6728\u5bf9CNE-2\u7ec6\u80de\u589e\u6b96\u6291\u5236\u7387\u53caIC50\u6d4b\u5b9a\u5168\u6d41\u7a0b\u64cd\u4f5c\u624b\u518c_v2.0.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_spacing(para, before=0, after=0, line_rule=None, line_val=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line_rule and line_val:
        spacing.set(qn('w:lineRule'), line_rule)
        spacing.set(qn('w:line'), str(line_val))
    pPr.append(spacing)

def add_chapter_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        set_para_spacing(p, before=240, after=120)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        set_para_spacing(p, before=180, after=60)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        set_para_spacing(p, before=120, after=60)
    run.font.name = '\u9ed1\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    return p

def add_body_para(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '480')
        pPr.append(ind)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=60, line_rule='auto', line_val=360)
    return p

def add_red_line_box(doc, title, items):
    """Red box for absolute red lines/requirements"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'FFE0E0')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '16')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'C00000')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    p_title = cell.paragraphs[0]
    r = p_title.add_run('\u274c ' + title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    set_para_spacing(p_title, before=60, after=30)
    for item in items:
        p2 = cell.add_paragraph()
        r2 = p2.add_run('\u2718 ' + item)
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.name = '\u5b8b\u4f53'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        r2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        set_para_spacing(p2, before=0, after=30)
    doc.add_paragraph()

def add_warning_box(doc, title, lines, color='FF8C00', bg='FFF2CC'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)
    p_title = cell.paragraphs[0]
    r = p_title.add_run('\u26a0 ' + title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
    set_para_spacing(p_title, before=60, after=30)
    for line in lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run('\u25b6 ' + line)
        r2.font.size = Pt(10.5)
        r2.font.name = '\u5b8b\u4f53'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p2, before=0, after=30, line_rule='auto', line_val=340)
    doc.add_paragraph()

def add_green_box(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'E2EFDA')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '548235')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    p_title = cell.paragraphs[0]
    r = p_title.add_run('\u2713 ' + title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(0x54, 0x82, 0x35)
    set_para_spacing(p_title, before=60, after=30)
    for line in lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run('  ' + line)
        r2.font.size = Pt(10.5)
        r2.font.name = '\u5b8b\u4f53'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p2, before=0, after=30, line_rule='auto', line_val=340)
    doc.add_paragraph()

def add_three_line_table(doc, headers, rows, caption=None):
    if caption:
        p_cap = doc.add_paragraph()
        r_cap = p_cap.add_run(caption)
        r_cap.bold = True
        r_cap.font.size = Pt(10.5)
        r_cap.font.name = '\u5b8b\u4f53'
        r_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_cap, before=60, after=30)
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name, sz in [('top', '12'), ('bottom', '12')]:
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    for border_name in ['insideH', 'insideV']:
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, 'D6E4F7')
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bot_b = OxmlElement('w:bottom')
        bot_b.set(qn('w:val'), 'single')
        bot_b.set(qn('w:sz'), '6')
        bot_b.set(qn('w:space'), '0')
        bot_b.set(qn('w:color'), '000000')
        tcBorders.append(bot_b)
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '\u5b8b\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F5F9FF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '\u5b8b\u4f53'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    doc.add_paragraph()

def add_trouble_section(doc, problem_num, problem, causes, solutions, prevention=None):
    """Add a structured troubleshooting entry"""
    # Problem header
    p_prob = doc.add_paragraph()
    set_cell_bg_para = False
    pPr = p_prob._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '24')
    left_b.set(qn('w:space'), '4')
    left_b.set(qn('w:color'), 'C00000')
    pBdr.append(left_b)
    pPr.append(pBdr)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '240')
    pPr.append(ind)
    r = p_prob.add_run('\u95ee\u9898' + str(problem_num) + '\uff1a' + problem)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = '\u9ed1\u4f53'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    set_para_spacing(p_prob, before=120, after=30)

    # Causes
    p_cause_hdr = doc.add_paragraph()
    pPr2 = p_cause_hdr._p.get_or_add_pPr()
    ind2 = OxmlElement('w:ind')
    ind2.set(qn('w:left'), '240')
    pPr2.append(ind2)
    r2 = p_cause_hdr.add_run('\u3010\u53ef\u80fd\u539f\u56e0\u3011')
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.name = '\u9ed1\u4f53'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r2.font.color.rgb = RGBColor(0x84, 0x36, 0x0C)
    set_para_spacing(p_cause_hdr, before=30, after=0)

    for cause in causes:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '480')
        ind.set(qn('w:firstLine'), '0')
        pPr.append(ind)
        run = p.add_run('\u25c6 ' + cause)
        run.font.size = Pt(10.5)
        run.font.name = '\u5b8b\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p, before=0, after=20, line_rule='auto', line_val=340)

    # Solutions
    p_sol_hdr = doc.add_paragraph()
    pPr3 = p_sol_hdr._p.get_or_add_pPr()
    ind3 = OxmlElement('w:ind')
    ind3.set(qn('w:left'), '240')
    pPr3.append(ind3)
    r3 = p_sol_hdr.add_run('\u3010\u89e3\u51b3\u65b9\u6848\u3011')
    r3.bold = True
    r3.font.size = Pt(11)
    r3.font.name = '\u9ed1\u4f53'
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    r3.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    set_para_spacing(p_sol_hdr, before=30, after=0)

    for i, sol in enumerate(solutions, 1):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '480')
        ind.set(qn('w:firstLine'), '0')
        pPr.append(ind)
        run = p.add_run(str(i) + '. ' + sol)
        run.font.size = Pt(10.5)
        run.font.name = '\u5b8b\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        set_para_spacing(p, before=0, after=20, line_rule='auto', line_val=340)

    if prevention:
        p_prev_hdr = doc.add_paragraph()
        pPr4 = p_prev_hdr._p.get_or_add_pPr()
        ind4 = OxmlElement('w:ind')
        ind4.set(qn('w:left'), '240')
        pPr4.append(ind4)
        r4 = p_prev_hdr.add_run('\u3010\u9884\u9632\u63aa\u65bd\u3011')
        r4.bold = True
        r4.font.size = Pt(11)
        r4.font.name = '\u9ed1\u4f53'
        r4._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
        r4.font.color.rgb = RGBColor(0x54, 0x82, 0x35)
        set_para_spacing(p_prev_hdr, before=30, after=0)
        for prev in prevention:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '480')
            ind.set(qn('w:firstLine'), '0')
            pPr.append(ind)
            run = p.add_run('\u2713 ' + prev)
            run.font.size = Pt(10.5)
            run.font.name = '\u5b8b\u4f53'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
            run.font.color.rgb = RGBColor(0x54, 0x82, 0x35)
            set_para_spacing(p, before=0, after=20, line_rule='auto', line_val=340)

    doc.add_paragraph()


# ============================================================
# MAIN
# ============================================================
print("Loading document for Chapter 5 addition...")
doc = Document(DOC_PATH)

doc.add_page_break()
add_chapter_heading(doc, '\u7b2c5\u7ae0  \u8d28\u91cf\u63a7\u5236\u3001\u5b9e\u9a8c\u7ea2\u7ebf\u4e0e\u6545\u969c\u6392\u67e5\u6307\u5357', level=1)

add_body_para(doc,
    '\u672c\u7ae0\u5305\u62ec\u4e24\u4e2a\u90e8\u5206\uff1a\u7b2c\u4e00\u90e8\u5206\u4e3a\u672c\u5b9e\u9a8c\u7684\u7ed3\u679c\u5224\u5b9a\u6807\u51c6\u4e0e\u7edd\u5bf9\u7ea2\u7ebf'
    '\uff08\u51fa\u73b0\u7ea2\u7ebf\u60c5\u51b5\u5219\u5b9e\u9a8c\u5fc5\u987b\u91cd\u505a\uff09\uff1b\u7b2c\u4e8c\u90e8\u5206\u4e3a15\u4e2a\u5e38\u89c1\u5b9e\u9a8c\u95ee\u9898\u7684\u6545\u969c\u6392\u67e5\u6307\u5357\u3002',
    indent=True)

# -------------------------------------------------------
# 5.1 Result acceptance criteria / Red lines
# -------------------------------------------------------
add_chapter_heading(doc, '5.1  \u5b9e\u9a8c\u7ed3\u679c\u5224\u5b9a\u6807\u51c6\u4e0e\u7ea2\u7ebf', level=2)

add_chapter_heading(doc, '5.1.1  \u7ed3\u679c\u63a5\u53d7\u6807\u51c6\u603b\u8868', level=3)
add_body_para(doc, '\u4ee5\u4e0b\u662f\u672c\u5b9e\u9a8c\u5168\u6d41\u7a0b\u7684\u7ed3\u679c\u5224\u5b9a\u6807\u51c6\u3002\u6240\u6709\u3010\u5fc5\u987b\u6ee1\u8db3\u3011\u6761\u4ef6\u5747\u4e3a\u7ea2\u7ebf\uff0c\u4efb\u4e00\u4e0d\u6ee1\u8db3\u5219\u8be5\u6b21\u5b9e\u9a8c\u6570\u636e\u4e0d\u53ef\u63a5\u53d7\u3002', indent=True)

add_three_line_table(doc,
    ['\u6307\u6807\u7c7b\u578b', '\u6307\u6807\u540d\u79f0', '\u63a5\u53d7\u6807\u51c6', '\u91cd\u8981\u7ea7\u522b'],
    [
        ['\u7ec6\u80de\u8d28\u91cf', '\u94fa\u677f\u524d\u6d3b\u529b', '>= 95%', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u7ec6\u80de\u8d28\u91cf', '\u6d3b\u529b\u68c0\u9a8c(\u53f0\u76fc\u84dd\u6392\u9664\u6cd5)', '< 5%\u6b7b\u4ea1\u7387', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u7ec6\u80de\u8d28\u91cf', '\u652f\u539f\u4f53\u5f62\u6001', '\u8d34\u58c1\u751f\u957f\uff0c\u5f62\u6001\u6b63\u5e38', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u7ec6\u80de\u8d28\u91cf', '\u6d53\u5ea6\uff08\u94fa\u677f\u65f6\uff09', '2x10^4 - 3x10^4\u4e2a/mL', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u5bf9\u7167\u7ec4', '\u9634\u6027\u5bf9\u7167OD\u5747\u503c', '>= 0.4', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u5bf9\u7167\u7ec4', '\u9634\u6027\u5bf9\u7167CV', '<= 10%', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u5bf9\u7167\u7ec4', 'DMSO\u5bf9\u7167 vs \u9634\u6027\u5bf9\u7167', '\u5dee\u5f02 < 5%', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u836f\u7269\u5904\u7406', 'DMSO\u6700\u7ec8\u6d53\u5ea6', '<= 0.1%\uff08v/v\uff09', '\u7ea2\u7ebf'],
        ['\u836f\u7269\u5904\u7406', '\u5404\u7ec4\u6291\u5236\u7387CV', '<= 10%', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u836f\u7269\u5904\u7406', '\u6291\u5236\u7387\u5355\u8c03\u6027', '\u6d53\u5ea6\u4f9d\u8d56\u6027\u6b63\u786e', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u62df\u5408\u8d28\u91cf', 'R2\uff08\u4e09\u4e2a\u65f6\u95f4\u70b9\uff09', '>= 0.95', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u62df\u5408\u8d28\u91cf', 'HillSlope\u8303\u56f4', '0.5 ~ 3.0', '\u9700\u7b26\u5408'],
        ['\u91cd\u590d\u6027', '\u751f\u7269\u5b66\u91cd\u590d\u6b21\u6570', '>= 3\u6b21\u72ec\u7acb\u91cd\u590d', '\u5fc5\u987b\u6ee1\u8db3'],
        ['\u91cd\u590d\u6027', '\u4e09\u6b21IC50\u7ec4\u95f4CV', '<= 20%', '\u5efa\u8bae\u6ee1\u8db3'],
    ],
    '\u88685-1  \u5b9e\u9a8c\u7ed3\u679c\u5224\u5b9a\u6807\u51c6\u603b\u8868'
)

add_chapter_heading(doc, '5.1.2  \u7edd\u5bf9\u7ea2\u7ebf\uff08\u5fc5\u987b\u91cd\u505a\u6761\u4ef6\uff09', level=3)

add_red_line_box(doc, '\u7edd\u5bf9\u7ea2\u7ebf\uff1a\u51fa\u73b0\u4ee5\u4e0b\u4efb\u4e00\u60c5\u51b5\u5fc5\u987b\u5e9f\u5f03\u8be5\u6b21\u5b9e\u9a8c\u6570\u636e\u5e76\u91cd\u505a', [
    '\u7ea2\u7ebf\u2460\uff1aDMSO\u6700\u7ec8\u6d53\u5ea6 > 0.1%\uff08\u5bfc\u81f4\u8f66\u8f7d\u8fa9\u5269\u4e0d\u53ef\u6392\u9664\uff09',
    '\u7ea2\u7ebf\u2461\uff1a\u94fa\u677f\u524d\u7ec6\u80de\u6d3b\u529b < 90%\uff08CCK-8\u57fa\u7ebf\u4fe1\u53f7\u4e0d\u8db3\uff09',
    '\u7ea2\u7ebf\u2462\uff1a\u9634\u6027\u5bf9\u7167OD450 < 0.3\uff08\u7ec6\u80de\u589e\u6b96\u4e25\u91cd\u4e0d\u8db3\uff09',
    '\u7ea2\u7ebf\u2463\uff1a\u9634\u6027\u5bf9\u7167CV > 20%\uff08\u6a21\u677f\u5185\u5747\u4e00\u6027\u4e25\u91cd\u5dee\u5f02\uff09',
    '\u7ea2\u7ebf\u2464\uff1a\u5b9e\u9a8c\u4e2d\u53d1\u73b0\u8096\u6d46\u6c61\u67d3\uff08\u54cd\u6682\u505c\u5b9e\u9a8c\uff0c\u5904\u7406\u6c61\u67d3\u540e\u91cd\u65b0\u57f9\u517b\uff09',
    '\u7ea2\u7ebf\u2465\uff1a\u4efb\u4e00\u65f6\u95f4\u70b9R2 < 0.90\uff08\u62df\u5408\u66f2\u7ebf\u6570\u636e\u8d28\u91cf\u4e0d\u8fbe\u6807\uff09',
    '\u7ea2\u7ebf\u2466\uff1a\u5b9e\u9a8c\u7ec4\u548c\u5bf9\u7167\u7ec4\u94fa\u677f\u8bef\u64cd\u4f5c\uff08\u5982\u8fb9\u7f18\u5b54\u672a\u52a0PBS\u800c\u76f4\u63a5\u7ec6\u80de\uff09',
    '\u7ea2\u7ebf\u2467\uff1a\u7ec6\u80de\u672a\u6e05\u6d17\u76f4\u63a5\u7ed9\u836f\uff08DMSO\u8840\u6e05\uff0c\u836f\u7269\u548c\u6b8b\u4f59\u57f9\u517b\u57fa\u76f8\u4e92\u5e72\u6270\uff09',
])

add_chapter_heading(doc, '5.1.3  \u9ed1\u7ebf\u6307\u6807\uff08\u63a5\u53d7\u4f46\u9700\u5907\u6ce8\u8bf4\u660e\uff09', level=3)
add_body_para(doc, '\u4ee5\u4e0b\u60c5\u51b5\u4e0d\u5c5e\u4e8e\u7ea2\u7ebf\uff0c\u4f46\u9700\u5728\u5b9e\u9a8c\u62a5\u544a\u4e2d\u8be6\u7ec6\u8bf4\u660e\u5e76\u5206\u6790\u53ef\u80fd\u5f71\u54cd\uff1a', indent=True)
for item in [
    '10% < \u9634\u6027\u5bf9\u7167CV <= 20%\uff1a\u9700\u5206\u6790\u539f\u56e0\uff08\u5982\u6c29\u5316\u4e0d\u5747\uff09\u5e76\u5907\u6ce8',
    '0.90 <= R2 < 0.95\uff1a\u63a5\u53d7\u4f46\u9700\u5907\u6ce8\u6570\u636e\u62df\u5408\u8d28\u91cf\u5076\u5c14\u504f\u5dee\u7684\u53ef\u80fd\u6027',
    'IC50\u4f4d\u4e8e\u5b9e\u9a8c\u6d53\u5ea6\u8303\u56f4\u7684\u8fb9\u754c\uff08\u5982\u63a5\u8fd13.125\u6216400\u03bcg/mL\uff09\uff1a\u63a5\u53d7\u4f46\u5efa\u8bae\u8865\u5145\u8fb9\u754c\u6d53\u5ea6',
    '\u4e09\u6b21\u72ec\u7acb\u91cd\u590d\u95f4IC50\u7ec4\u95f4CV > 20%\uff1a\u9700\u5206\u6790\u539f\u56e0\uff0c\u5efa\u8bae\u589e\u52a0\u91cd\u590d\u6b21\u6570',
]:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    run = p.add_run('\u25e6 ' + item)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=30)

# -------------------------------------------------------
# 5.2 Outlier detection
# -------------------------------------------------------
add_chapter_heading(doc, '5.2  \u5f02\u5e38\u5024\u68c0\u9a8c\u4e0e\u6570\u636e\u6e05\u6d17', level=2)

add_chapter_heading(doc, '5.2.1  Grubbs\u68c0\u9a8c\u6cd5\uff08G\u68c0\u9a8c\uff09', level=3)
add_body_para(doc,
    'Grubbs\u68c0\u9a8c\u662f\u5c0f\u6837\u672c\uff08n=3-10\uff09\u5f02\u5e38\u5024\u68c0\u9a8c\u7684\u9996\u9009\u65b9\u6cd5\u3002\u5f53\u67d0\u5b54CV > 15%\u65f6\uff0c\u5c1d\u8bd5\u4f7f\u7528Grubbs\u68c0\u9a8c\u786e\u5b9a\u6700\u53ef\u75591\u4e2a\u5f02\u5e38\u5024\u5e76\u6392\u9664\u3002',
    indent=True)

add_three_line_table(doc,
    ['n\uff08\u91cd\u590d\u5b54\u6570\uff09', 'G\u4e34\u754c\u5024\uff08\u03b1=0.05\uff09', 'G\u4e34\u754c\u5024\uff08\u03b1=0.01\uff09'],
    [
        ['3', '1.155', '1.155'],
        ['4', '1.481', '1.496'],
        ['5', '1.715', '1.764'],
        ['6', '1.887', '1.973'],
        ['8', '2.126', '2.274'],
        ['10', '2.290', '2.482'],
    ],
    '\u88685-2  Grubbs\u68c0\u9a8c G\u7edf\u8ba1\u91cf\u4e34\u754c\u5024\u8868'
)

add_body_para(doc, 'Grubbs\u68c0\u9a8c\u64cd\u4f5c\u6b65\u9aa4\uff1a', indent=True)
steps_grubbs = [
    '\u7b2c\u4e00\u6b65\uff1a\u8ba1\u7b97n\u4e2a\u91cd\u590d\u5b54\u7684\u5747\u503c\uff08mean\uff09\u548c\u6807\u51c6\u5dee\uff08SD\uff09',
    '\u7b2c\u4e8c\u6b65\uff1a\u627e\u51fa\u5047\u60f3\u5f02\u5e38\u5024\uff08\u5373\u5411\u5747\u5026\u504f\u5dee\u6700\u5927\u7684\u5b54\uff09',
    '\u7b2c\u4e09\u6b65\uff1a\u8ba1\u7b97G\u7edf\u8ba1\u91cf = |\u5047\u60f3\u5f02\u5e38\u5024 - mean| / SD',
    '\u7b2c\u56db\u6b65\uff1a\u67e5\u4e0a\u8868\uff0c\u82e5G > G\u4e34\u754c\uff08\u03b1=0.05\uff09\uff0c\u5219\u8be5\u5b54\u4e3a\u5f02\u5e38\u5024\uff0c\u53ef\u6392\u9664',
    '\u7b2c\u4e94\u6b65\uff1a\u6392\u9664\u5f02\u5e38\u5024\u540e\uff0c\u7528\u5269\u4f59n-1\u4e2a\u6570\u636e\u91cd\u65b0\u8ba1\u7b97\u5747\u503c\uff0cCV\u5e94\u964d\u81f3<=10%',
]
for step in steps_grubbs:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    run = p.add_run('\u2022 ' + step)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=30)

add_warning_box(doc, '\u5f02\u5e38\u5024\u5904\u7406\u89c4\u8303', [
    '\u6bcf\u4e2a\u6d53\u5ea6\u7ec4\u6700\u591a\u53ea\u80fd\u6392\u9664\u4e00\u4e2a\u5f02\u5e38\u5024\uff08\u5373Grubbs\u68c0\u9a8c\u6bcf\u6b21\u53ea\u8bc6\u522b\u6700\u5f02\u5e38\u7684\u5b54\uff09',
    '\u6392\u9664\u5f02\u5e38\u5024\u540en=2\u65f6\uff0cCV\u65e0\u6cd5\u8ba1\u7b97\uff1b\u82e5\u5269\u4f59\u4e24\u5b542\u5b54CV\u4ecd>15%\uff0c\u5e94\u8003\u8651\u91cd\u505a\u8be5\u6d53\u5ea6\u7ec4',
    '\u5f02\u5e38\u5024\u6392\u9664\u60c5\u51b5\u5fc5\u987b\u5728\u5b9e\u9a8c\u8bb0\u5f55\u672c\u4e2d\u8be6\u7ec6\u8bb0\u5f55\uff08\u5305\u62ec\u88ab\u6392\u9664\u7684OD\u5747\u503c\u3001G\u7edf\u8ba1\u91cf\u548c\u4e34\u754c\u5024\uff09',
    '\u4e0d\u5f97\u4ee5\u63d0\u9ad8IC50\u6216\u6539\u5584\u62df\u5408\u4e3a\u76ee\u7684\u4e3b\u89c2\u5220\u9664\u5b54\u6570\u636e',
], color='FF8C00', bg='FFF2CC')

# -------------------------------------------------------
# 5.3 Troubleshooting guide
# -------------------------------------------------------
add_chapter_heading(doc, '5.3  \u5e38\u89c1\u5b9e\u9a8c\u95ee\u9898\u6545\u969c\u6392\u67e5\u6307\u5357', level=2)

add_body_para(doc,
    '\u672c\u8282\u6536\u5f5515\u4e2a\u6700\u5e38\u89c1\u7684\u5b9e\u9a8c\u95ee\u9898\uff0c\u6bcf\u4e2a\u95ee\u9898\u5747\u63d0\u4f9b\u53ef\u80fd\u539f\u56e0\u3001\u89e3\u51b3\u65b9\u6848\u548c\u9884\u9632\u63aa\u65bd\uff0c\u9002\u5408\u5b9e\u9a8c\u51fa\u73b0\u610f\u5916\u65f6\u5bf9\u7167\u4f7f\u7528\u3002',
    indent=True)

# Trouble 1
add_trouble_section(doc, 1, '\u9634\u6027\u5bf9\u7167OD\u5048\u5c0f\uff08<0.3\uff09',
    ['\u94fa\u677f\u7ec6\u80de\u5bc6\u5ea6\u8fc7\u4f4e\uff08\u9519\u8bef\u8ba1\u7b97\u6216\u62df\u677f\u64cd\u4f5c\u4e0d\u5f53\uff09',
     '\u7ec6\u80de\u6d3b\u529b\u5dee\uff08\u590d\u82cf\u540e\u51fa\u73b0\u5e94\u6fc0\u635f\u4f24\uff09',
     'CCK-8\u8bd5\u5242\u4e59\u7ef4\u65f6\u95f4\u4e0d\u8db3',
     '\u9176\u6807\u4eea\u68c0\u6d4b\u6ce2\u957f\u9519\u8bef\uff08\u5e94\u4e3a450nm\uff09'],
    ['\u68c0\u67e5\u94fa\u677f\u65f6\u7ec6\u80de\u8ba1\u6570\u8bb0\u5f55\uff0c\u786e\u8ba4\u5bc6\u5ea6\u8fbe2x10^4/mL',
     '\u91cd\u65b0\u590d\u82cf\u7ec6\u80de\uff0c\u590d\u82cf24h\u540e\u68c0\u6d4b\u6d3b\u529b\u518d\u94fa\u677f',
     '\u5ef6\u957fCCK-8\u5b54\u5185\u5b5035-45min\u540e\u91cd\u65b0\u68c0\u6d4b',
     '\u786e\u8ba4\u9176\u6807\u4eea\u8bbe\u7f6e\u4e3aOD450nm\uff0c\u5e76\u786e\u8ba4\u53c2\u8003\u6ce2\u957f620nm'],
    ['\u94fa\u677f\u524d\u6ce8\u610f\u7ec6\u8083\u548c\u8ba1\u6570\u7cbe\u5ea6\uff0c\u4e0d\u5f97\u63a8\u7b97',
     '\u590d\u82cf\u540e\u4e0d\u5c11\u4e8e30min\u624d\u5f00\u59cb\u62cd\u677f\uff0c\u786e\u8ba4\u7ec6\u80de\u8d34\u58c1'])

# Trouble 2
add_trouble_section(doc, 2, '\u9634\u6027\u5bf9\u7167CV > 15%',
    ['\u94fa\u677f\u4e0d\u5747\u5300\uff08\u7ec6\u80de\u60ac\u6d6e\u6db2\u6df7\u5747\u4e0d\u5145\u5206\uff09',
     '\u8fb9\u7f18\u6548\u5e94\u672a\u5145\u5206\u9884\u9632\uff08\u8fb9\u7f18\u5b54\u6c34\u5206\u8b70\u53d1\uff09',
     '\u591a\u9053\u6db2\u5c55\u52a0\u65f6\u5404\u9053\u6d41\u901f\u4e0d\u4e00\u81f4',
     '\u7ec6\u80de\u5728\u96e2\u5fc3\u540e\u8fc7\u957f\u65f6\u95f4\u672a\u91cd\u60ac\u6d6e'],
    ['\u91cd\u65b0\u5236\u5907\u7ec6\u80de\u60ac\u6d6e\u6db2\uff0c\u51b3\u5c71\u5f0f\u6df7\u5300(\u77ac\u65e0\uff0810s+\u5012\u7f6e10x)\u540e\u518d\u62bd\u53d6',
     '\u786e\u8aa49x12\u8fb9\u7f18\u5b54\u52a0\u5165200uL PBS\u540e\u518d\u94fa\u677f',
     '\u68c0\u67e5\u591a\u9053\u5c55\u52a0\u76841-8\u9053\u6db2\u5c55\u91cf\u5dee\u5f02\uff08\u4f7f\u7528\u9009\u9898\u5b9a\u91cf\u5c55\u52a0\u786e\u8ba4\uff09',
     '\u62bd\u53d6\u7ec6\u80de\u540e10min\u5185\u5b8c\u6210\u94fa\u677f\u64cd\u4f5c'],
    ['\u5404\u6db2\u5c55\u65b0\u5e72\u9759\u540e\u518d\u7528\uff0c\u907f\u514d\u5934\u90e8\u6c14\u6ce1',
     '\u94fa\u677f\u65f6\u4fdd\u6301\u5bb9\u5668\u5782\u76f4'])

# Trouble 3
add_trouble_section(doc, 3, 'CCK-8\u52a0\u5165\u540e\u989c\u8272\u5de2\u5f02\uff08\u5c40\u90e8\u6df1\u8272\u6216\u6d45\u8272\uff09',
    ['\u591a\u9053\u5c55\u52a0CCK-8\u65f6\u5c40\u90e8\u6df1\u5c42\u6df7\u5300\u4e0d\u5145\u5206',
     'CCK-8\u8bd5\u5242\u5728\u51b0\u4e0a\u56de\u6e29\u4e0d\u5145\u5206\uff0c\u5c40\u90e8\u6e38\u79bb\u9178',
     '\u9a77\u5c55\u5934\u5982\u5c16\u7aef\u5185\u6709\u6c14\u6ce1\uff0c\u62c9\u5165\u91cf\u4e0d\u51c6\u786e'],
    ['\u52a0\u5165CCK-8\u540e\u8f7b\u8f7b\u62cd\u7ef4\u660410x\u6df7\u5300\uff0c\u518d\u653e\u56de\u57f9\u517b\u7b9550min\u540e\u68c0\u6d4b',
     '\u63d0\u524d30min\u53d6\u51faCCK-8\u5c31\u5b563\uff0c\u5c55\u5f00\u76d6\u5b50\u56de\u6e29\u5230\u5ba4\u6e29',
     '\u4e0d\u8981\u4f7f\u7528\u5c16\u5934\u8fd4\u5c16\u5c55\u52a0\u5934'],
    ['\u6bcf\u6b21\u4f7f\u7528CCK-8\u524d\u5148\u9493\u52a8\u6df7\u5300\uff0c\u786e\u8ba4\u65e0\u6c89\u6de0'])

# Trouble 4
add_trouble_section(doc, 4, 'DMSO\u8f66\u8f7d\u5bf9\u7167\u7ec4OD\u660e\u663e\u4f4e\u4e8e\u9634\u6027\u5bf9\u7167\uff08DMSO\u6bd2\u6027\uff09',
    ['\u836f\u7269\u6bcd\u6db2\u4e2dDMSO\u6d53\u5ea6\u8fc7\u9ad8\uff08>10%\uff09',
     '\u4e0d\u540c\u6d53\u5ea6\u7ec4DMSO\u6700\u7ec8\u6d53\u5ea6\u4e0d\u4e00\u81f4',
     '\u5404\u6d53\u5ea6\u7ec4\u4e4b\u95f4DMSO\u878d\u5242\u91cf\u8ba1\u7b97\u9519\u8bef'],
    ['\u91cd\u65b0\u68c0\u67e5\u5404\u6d53\u5ea6\u7ec4\u836f\u7269\u5de5\u4f5c\u6db2\u4e2dDMSO\u5b9e\u9645\u6d53\u5ea6',
     '\u786e\u4fdd\u8f66\u8f7d\u5bf9\u7167\u7ec4DMSO\u6d53\u5ea6\u7b49\u4e8e\u836f\u7269\u7ec4\u6700\u9ad8\u6d53\u5ea6\u7ec4\u7684DMSO\u6d53\u5ea6',
     '\u4ee5C1\uff08400ug/mL\uff09\u5de5\u4f5c\u6db2\u4e3a\u4f8b\uff1a\u5468DMSO\u6d53\u5ea6=0.04%\uff0c\u628a\u8f66\u8f7d\u5bf9\u7167\u7ec4DMSO\u8c03\u81f3\u540c\u4e00\u6d53\u5ea6'],
    ['\u5236\u5907\u836f\u7269\u6bcd\u6db2\u65f6\u4e25\u683c\u6309\u7535\u5b50\u79d8\u4e2d\u516c\u5f0f\u9a8c\u8bc1DMSO\u6d53\u5ea6',
     '\u5efa\u7acb\u5bfc\u6c34\u5c42\u6d53\u5ea6\u8868\uff0c\u6bcf\u4e2a\u6d53\u5ea6\u7ec4DMSO\u6d53\u5ea6\u5e94\u76f8\u540c'])

# Trouble 5
add_trouble_section(doc, 5, '\u5242\u91cf-\u6548\u5e94\u66f2\u7ebf\u4e0d\u6210S\u5f62\uff08R2<0.90\uff09',
    ['\u6d53\u5ea6\u8303\u56f4\u4e0d\u591f\u5bbd\uff08\u672a\u56ca\u62ec IC50\uff09',
     '\u9ad8\u6d53\u5ea6\u7ec4\u836f\u7269\u6c89\u6de0\u6216\u5b57\u6e90',
     '\u5b57\u6e90 DMSO \u81ea\u8eab\u5177\u6709\u7ec6\u80de\u6bd2\u6027',
     '\u5bf9\u7167\u7ec4\u9489\u7267\u5316\u6216\u6c61\u67d3'],
    ['\u6839\u636e\u9884\u5b9e\u9a8c\u7ed3\u679c\u8c03\u6574\u6d53\u5ea6\u8303\u56f4',
     '\u68c0\u67e5\u836f\u7269\u5728\u57f9\u517b\u57fa\u4e2d\u7684\u6eb6\u89e3\u6027\uff0c\u5fc5\u8981\u65f6\u8fc7\u6ee4\u6d88\u6bd2',
     '\u5c1d\u8bd5\u91ca\u653eBottom/Top\u7ea6\u675f\uff0c\u5206\u6790\u662f\u5426\u53ef\u80fd\u8fbe\u52060%\u6216100%\u6291\u5236',
     '\u589e\u52a0\u8fb9\u754c\u6d53\u5ea6\u70b9\uff08C0: 800ug/mL\uff09'],
    ['\u5b9e\u9a8c\u524d\u5e94\u8fdb\u884c\u9884\u5b9e\u9a8c\u786e\u8ba4\u6d53\u5ea6\u8303\u56f4'])

# Trouble 6
add_trouble_section(doc, 6, '\u84b8\u53d1\u6f0f\u6c14\u6d4e\u5bfc\u81f4\u5b54\u5185\u6db2\u4f53\u5927\u91cf\u6d88\u5931',
    ['\u5b54\u677f\u5c01\u53e3\u81a8\u8fc7\u677f\u76d803-01\u5c01\u53e3\u6000\u5c01\u53e3\u6761\u6f0f\u6c14',
     '\u5b54\u677f\u4e0d\u5e73\u7f6e\uff0c\u80dc\u3001\u5c01\u53e3\u6761\u7136\u5c40\u90e8\u671d\u4e0a',
     'CO2\u7f2a\u5c42\u638c\u63a1\u5b9c\u6e20\u5c04\uff0c\u5bfc\u81f4\u5c40\u90e8\u6e29\u5ea6\u504f\u9ad8'],
    ['\u66f4\u6362\u5c01\u53e3\u6761\uff0c\u68c0\u67e5\u5c01\u6c14\u52b4\u6307\u538b\u529b',
     '\u786e\u8ba4\u5b54\u677f\u6c34\u5e73\u653e\u7f6e\uff0c\u5185\u5c42\u88c5\u6c34\u84c9\u6c34',
     '\u5982\u5b54\u5185\u6db2\u4f53\u4f53\u79ef\u5c0f\u4e8e100uL\uff0c\u9700\u5e9f\u5f03\u8be5\u5b54\u6570\u636e'],
    ['\u6bcf\u6b21\u5b9e\u9a8c\u524d\u68c0\u67e5\u5c01\u53e3\u6761\u72b6\u6001\uff0c\u8001\u5c04\u624b\u5e94\u6bcf3-4\u5929\u66f4\u6362\u4e00\u6b21\u5c01\u53e3\u6761'])

# Trouble 7
add_trouble_section(doc, 7, '\u67d3\u8272\u5e72\u6270\uff08OD\u5168\u5b54\u5040\u9ad8\uff09',
    ['\u57f9\u517b\u57fa\u672c\u8eab\u5c31\u6709\u8272\u5b57\uff08\u5982\u542b\u9996\u2044\u6c26\u54c1\u7ea2\u3001\u5439\u6c27\u6307\u793a\u5242\uff09',
     '\u82cf\u6728\u63d0\u53d6\u7269\u5728\u67d0\u6d53\u5ea6\u4e0b\u5728450nm\u6709\u5403\u5145\u5438\u6536',
     'CCK-8\u5e95\u7269\u548cCCK-8\u6bcd\u6db2\u6211\u524d\u6df7\u5300\u4e0d\u5783'],
    ['\u7531\u4e8e\u82cf\u6728\u63d0\u53d6\u7269\u5728450nm\u7684\u5149\u5438\u6536\u60c5\u51b5\uff0c\u9700\u8bbe\u7f6e\u7a7a\u767d\u5bf9\u7167\u9a8c\u8bc1\uff1a\u836f\u7269\u5de5\u4f5c\u6db2\u52a0CCK-8\uff0c\u4f46\u4e0d\u542b\u7ec6\u80de',
     '\u82e5\u7a7a\u767d\u5bf9\u7167OD\u663e\u8457\u5347\u9ad8\uff0c\u8bf4\u660e\u836f\u7269\u5e72\u6270\uff0c\u9700\u5c06\u7a7a\u767d\u5bf9\u7167\u503c\u4ece\u6240\u6709\u5b54OD\u4e2d\u51cf\u53bb\uff08\u5b58\u5728\u5904\u7406\u4e4b\u540e\u518d\u5206\u6790\uff09'],
    ['\u5148\u6d4b\u8bd5\u82cf\u6728\u63d0\u53d6\u7269\u81ea\u8eab\u5728450nm\u7684\u5438\u5149\u5ea6\uff0c\u82e5OD>0.05\u5219\u9700\u5bf9\u6570\u636e\u8fdb\u884c\u8272\u5dee\u6821\u6b63'])

# Trouble 8
add_trouble_section(doc, 8, '\u5168\u677f\u6df1\u8272\uff08\u5305\u62ec\u7a7a\u767d\u5b54\uff09',
    ['CCK-8\u8bd5\u5242\u53d7\u6c61\u67d3\u6216\u53d1\u751f\u81ea\u5347\u53cd\u5e94',
     '\u5949\u7eb3\u7b49\u5e72\u6270\uff08\u975e\u7b54\u6c14\u5206\u6e14\u6db2\u53cd\u5e94\uff09',
     '\u9176\u6807\u4eea\u673a\u68b0\u6545\u969c\uff08\u8d3c\u8d1f\u53c2\u8003\uff09'],
    ['\u66f4\u6362\u65b0\u7684CCK-8\u8bd5\u5242\u81ea\u5c01\u3001\u540c\u5f39\u89e3\u5c01',
     '\u7a7a\u767d\u5b54\u7ec4OD>0.3\u60c5\u51b5\u4e0b\uff0c\u68c0\u67e5\u9176\u6807\u4eea\u53c2\u8003\u5e73\u5766\uff08\u8be6\u8bf4\u660e\u4e66\uff09',
     '\u5982\u9176\u6807\u4eea\u65e0\u6545\u969c\uff0c\u8bc1\u5b9eRCCK-8\u5df2\u53d7\u6c61\u67d3\uff0c\u4e0e\u5382\u5546\u8054\u7cfb'],
    ['\u6bcf\u6b21\u4f7f\u7528CCK-8\u65f6\u9700\u76ee\u89c2\u5224\u65ad\u989c\u8272\uff08\u5c90\u6c60\u8272\uff0c\u65e0\u6d51\u6d4a\uff09'])

# Trouble 9
add_trouble_section(doc, 9, '\u9634\u6027\u5bf9\u7167\u548cDMSO\u5bf9\u7167\u5dee\u5f02 > 15%',
    ['\u5404\u6d53\u5ea6\u7ec4DMSO\u8d44\u91cf\u6dfb\u52a0\u4e0d\u4e00\u81f4',
     'DMSO\u8f66\u8f7d\u6d53\u5ea6\u672a\u5339\u914d\u6700\u9ad8\u836f\u7269\u6d53\u5ea6\u7ec4',
     '\u9634\u6027\u5bf9\u7167\u5b54\u5185\u8ffd\u52a0\u4e86\u989d\u5916\u7684\u6eb6\u5242\u6216PBS'],
    ['\u9654\u5f52\u6240\u6709\u6d53\u5ea6\u7ec4\u7684\u5b9e\u9645DMSO\u6d53\u5ea6\u8ba1\u7b97\u8868',
     '\u786e\u8ba4\u8f66\u8f7d\u5bf9\u7167\u7ec4DMSO\u6d53\u5ea6=\u836f\u7269\u6700\u9ad8\u6d53\u5ea6\u7ec4DMSO\u6d53\u5ea6'],
    ['\u4e0b\u6b21\u8bbe\u8ba1\u5b9e\u9a8c\u65f6\u5e94\u7ecf\u7531\u4e2d\u95f4\u5c42\u6d53\u5ea6\u8868\u786e\u8ba4DMSO\u6d53\u5ea6'])

# Trouble 10
add_trouble_section(doc, 10, '\u7ec6\u80de\u90e4\u5f62\u3001\u6f02\u6d6e\u548c\u5806\u53e0\uff08\u5f02\u5e38\u5e8f\u80bd\u3016\uff09',
    ['\u7ec6\u80de\u4f20\u4ee3\u6b21\u6570\u8fc7\u591a\uff08> P30\uff09',
     '\u8096\u6d46\u6c61\u67d3',
     '\u57f9\u517b\u57fa\u548cCO2\u6d53\u5ea6\u9519\u8bef\uff08pH\u504f\u79fb\uff09'],
    ['\u69c8\u6e05\u7ec6\u80de\u4f20\u4ee3\u6b21\u6570\u5e76\u91cd\u65b0\u590d\u82cf\u65b0\u7279\u4ee3',
     '\u68c0\u6d4b\u8096\u6d46\uff08PCR\u6216\u8096\u6d46\u68c0\u6d4b\u8bd5\u5242\u76d2\uff09\uff0c\u4e00\u7ecf\u8bc1\u5b9e\u5201\u67d3\u5237\u7b39\u8bed\u5e9e',
     '\u6821\u9a8cCO2\u6d53\u5ea6\u5c06\u4e3a5%\uff0c\u78c1\u67e5\u5c01\u53e3\u6761'],
    ['\u6bcf2\u5468\u68c0\u6d4b\u8096\u6d46\uff0c\u7ec4\u56de\u57f9\u517b\u65f6\u69d8\u54c1\u5c31\u68c0',
     '\u5efa\u7acb\u7ec6\u80de\u4f20\u4ee3\u8bb0\u5f55\uff0c\u66ff\u6362\u59cf\u6bcd\u5c901\u4e2a\u6708\u66f4\u65b0\u4e00\u6b21'])

# Trouble 11
add_trouble_section(doc, 11, 'GraphPad HillSlope\u5c0f\u4e8e0\uff08\u5426\u6027Hill\u7cfb\u6570\uff09',
    ['\u4f7f\u7528\u4e86\u632a\u52a8\u578b\u6a21\u578b\uff08Agonist\uff09\u800c\u975e\u6291\u5236\u578b\uff08Inhibitor\uff09',
     'Y\u8f74\u6570\u636e\u8f93\u5165\u65b9\u5411\u9519\u8bef\uff08\u5e94\u8f93\u5165\u6291\u5236\u7387\uff0c\u8f93\u5165\u4e86\u5b58\u6d3b\u7387\uff09'],
    ['\u786e\u8ba4\u9009\u62e9\u6a21\u578b\u4e3a[log(inhibitor) vs. normalized response - Variable slope]',
     '\u68c0\u67e5Y\u8f74\u6570\u636e\u662f\u5426\u4e3a\u6291\u5236\u7387\uff08\u9ad8\u6d53\u5ea6\u5e94\u63a5\u8fd190-100%\uff09'],
    [])

# Trouble 12
add_trouble_section(doc, 12, '\u4e09\u6b21IC50\u91cd\u590d\u6027\u5dee\uff08\u7ec4\u95f4CV>30%\uff09',
    ['\u90e8\u5206\u91cd\u590d\u4e3a\u6280\u672f\u91cd\u590d\uff08\u540c\u4e00\u6279\u7ec6\u80de\uff09\u800c\u975e\u751f\u7269\u5b66\u72ec\u7acb\u91cd\u590d',
     '\u6bcf\u6b21\u5b9e\u9a8c\u7ec6\u80de\u4f20\u4ee3\u6570\u76f8\u5dee\u8fc7\u5927\uff08>5\u4ee3\uff09',
     '\u6bcf\u6b21\u5b9e\u9a8c\u836f\u7269\u5de5\u4f5c\u6db2\u6765\u6e90\u4e0d\u540c\uff08\u4e0d\u540c\u6279\u53f7\u6bcd\u6db2\uff09'],
    ['\u786e\u8ba4\u4e09\u6b21\u5747\u4e3a\u5b8c\u6574\u751f\u7269\u5b66\u91cd\u590d\uff08\u91cd\u65b0\u590d\u82cf\u7ec6\u80de\uff09',
     '\u7edf\u4e00\u4e09\u6b21\u7ec6\u80de\u4f20\u4ee3\u6570\u8303\u56f4\uff08P3-P10\uff09',
     '\u7edf\u4e00\u4e09\u6b21\u836f\u7269\u5de5\u4f5c\u6db2\u6765\u6e90\uff08\u540c\u4e00\u6279\u8d45\u6bcd\u6db2\uff09'],
    ['\u5b9e\u9a8c\u8bb0\u5f55\u672c\u5e94\u8bb0\u5f55\u6bcf\u6b21\u7ec6\u80de\u4f20\u4ee3\u6570\u548c\u5ef6\u9b45\u5206\u6cd5\u6848\u53f7'])

# Trouble 13
add_trouble_section(doc, 13, 'CCK-8\u5b54\u8272\u6df1\u4f46\u5bf9\u5e94\u5b54\u9a8c\u8bc1\u5716\u5bcc\u5b54\u6d43\u957fOD\u5df2\u8d85\u8fc7\u6a21\u677f\u8303\u56f4',
    ['CCK-8\u53cd\u5e94\u65f6\u95f4\u8fc7\u957f\u5bfc\u81f4\u589e\u5468\u7ed3\u678b\u5c40\u90e8\u5230\u8d85\u79e7\u73b0\u8c61',
     '\u4e0d\u540c\u6d53\u5ea6\u7ec4OD\u5dee\u5f02\u518d\u4e0d\u5206\u660e\u663e'],
    ['\u7f29\u77edCCK-8\u5b54\u5185\u5b5014\u5c66\u5c31\u9762\u8bf7\u5c55\u4fdd\u8bc1\u5b54\u5c405\uff0c\u5bfc\u6d59\u7ea7\u5f00\u59cb\u68c0\u6d4b',
     '\u68c0\u6d4b\u65f6\u95f4\u53ef\u9009\u62e92h\u67e5\u770b\u5feb\u589e\u6b96\u7ec4\u548c\u6291\u5236\u7ec4\u7684OD\u533a\u5206\u5ea6'],
    ['\u9732\u51faOD\u8d84\u51fa2.0\u7684\u4e1c\u897f\u76845\u5b54\u8303\u56f4\uff0c\u9700\u63d0\u524d\u68c0\u6d4b'])

# Trouble 14
add_trouble_section(doc, 14, '\u5b9e\u9a8c\u53ef\u91cd\u590d\u6027\u597d\u4f46\u5f97\u5230\u7684IC50\u548c\u6587\u732e\u62a5\u9053\u5dee\u5f02\u5927',
    ['\u82cf\u6728\u63d0\u53d6\u7269\u7684\u7269\u7406\u5316\u5b66\u5c5e\u6027\uff08\u6df7\u5408\u7269\u8fd8\u662f\u5355\u8d28\uff09\u4e0e\u6587\u732e\u4e0d\u540c',
     'CNE-2\u7ec6\u80de\u7ecf\u8fc7\u591a\u4ee3\u6b21\u57f9\u517b\uff0c\u906d\u4f20\u8868\u578b\u6f02\u79fb',
     '\u6587\u732e\u4e2d\u4f7f\u7528\u7684\u5b9e\u9a8c\u6761\u4ef6\u4e0d\u540c\uff08CCK-8 vs MTT\u6cd5\uff0c24h vs 48h\uff09'],
    ['\u6838\u67e5\u6587\u732e\u5c0f\u8282\u6cd5\u7684\u5b9e\u9a8c\u65b9\u6848\u5e76\u5bf9\u8bd5\u4f5c\u6bd4\u8f83',
     '\u6307\u5b9a\u4f7f\u7528\u7684\u82cf\u6728\u63d0\u53d6\u7269\u7b49\u5c42\uff0c\u5b9a\u8d28\u5b9a\u91cf\u6a19\u51c6',
     '\u5145\u5206\u9605\u8bfb\u591a\u7bc7\u6587\u732e\u5e76\u4e0e\u5bfc\u5e08\u8ba8\u8bba\u5dee\u5f02\u539f\u56e0'],
    [])

# Trouble 15
add_trouble_section(doc, 15, '\u6570\u636e\u5206\u6790\u5b8c\u6210\u540e\u65e0\u6cd5\u627e\u5230Prism\u6587\u4ef6',
    ['\u6587\u4ef6\u672a\u53ca\u65f6\u4fdd\u5b58\uff0c\u8ba1\u7b97\u673a\u7f16\u5d29\u6e83\u5bfc\u81f4\u4e22\u5931',
     '\u6587\u4ef6\u547d\u540d\u4e0d\u89c4\u8303\uff0c\u65e0\u6cd5\u5339\u914d'],
    ['\u5c1d\u8bd5\u8fdb\u5165\u8ba1\u7b97\u673a\u56de\u6536\u7f16\u6216\u4e34\u65f6\u6587\u4ef6\u5939',
     '\u8fdb\u884c\u540e\u7eed\u5b9e\u9a8c\u65f6\u5c5e\u4e8e\u8fd9\u4e0b\uff0c\u9c7c\u5212\u5c0f\u7528\u5904\u4e3b\u52a8\u8d4b\u4e88\u7b26\u5408\u89c4\u8303\u547d\u540d'],
    ['\u6bcf\u6b21\u5206\u6790\u5b8c\u6210\u540e\u7acb\u5373\u4fdd\u5b58\u5e76\u590d\u5236\u5230\u5907\u4efd\u76d8\uff0c\u8be6\u89c1\u7b2c4.5\u8282'])

# -------------------------------------------------------
# 5.4 Biosafety and waste disposal
# -------------------------------------------------------
add_chapter_heading(doc, '5.4  \u751f\u7269\u5b89\u5168\u548c\u5e9f\u5f03\u7269\u5904\u7f6e', level=2)

add_chapter_heading(doc, '5.4.1  \u751f\u7269\u5b89\u5168\u6805\u7ebf', level=3)
add_body_para(doc, '\u672c\u5b9e\u9a8c\u6d89\u53ca CNE-2 \u80ff\u4e2a\u8086\u7ec6\u80de\u7cfb\uff08\u8fb9\u754c \u6717 Epstein-Barr \u75c5\u6bd2\uff0cEBV\uff09\u53ca\u8096\u6d46\u575e\u67d3\u98ce\u9669\uff0c\u5b9e\u9a8c\u64cd\u4f5c\u5fc5\u987b\u5728\u4e8c\u7ea7\u751f\u7269\u5b89\u5168\u67dc\uff08BSC-II A2\u578b\uff09\u5185\u6267\u884c\u3002', indent=True)

add_red_line_box(doc, '\u751f\u7269\u5b89\u5168\u64cd\u4f5c\u7edd\u5bf9\u7981\u6b62', [
    '\u981a\u77fb\u5168\u7a0b\u5b9e\u9a8c\u64cd\u4f5c\u5fc5\u987b\u5e26\u624b\u5957\uff0c\u4e0d\u5f97\u7528\u88f8\u624b\u63a5\u89e6\u7ec6\u80de\u548c\u57f9\u517b\u57fa',
    '\u7ec6\u80de\u60ac\u6d6e\u6db2\u6284\u78b1\u4e0d\u5f97\u5728\u5b89\u5168\u67d3\u5916\u6f33\u6b63\u6c14\u6d41\u64cd\u4f5c',
    '\u53d1\u751f\u610f\u5916\u6cc4\u6f0f\uff0c\u5fc5\u987b\u7acb\u5373\u4f7f\u79e3\u7b26\u5408\u7cfb\u5217\uff0c\u5c11\u7528\u5ea6\u6c92\u9152\u6216\u6b63\u786e\u6d88\u6bd2\u5242\u5904\u7406',
    '\u5b9e\u9a8c\u4e2d\u4ea7\u751f\u7684\u6240\u6709\u548c\u5176\u624b\u53ca\u59fd\u6750\u6599\u5fc5\u987b\u6309\u5e9f\u5f03\u7269\u5904\u7f6e\u75c5\u539f\u4f53\u7ea7\u5904\u7f6e\uff0c\u4e0d\u5f97\u76f4\u63a5\u5012\u5165\u666e\u901a\u751f\u6d3b\u5783\u573e',
])

add_chapter_heading(doc, '5.4.2  \u5e9f\u5f03\u7269\u5206\u7c7b\u5904\u7f6e\u89c4\u8303', level=3)

add_three_line_table(doc,
    ['\u5e9f\u5f03\u7269\u7c7b\u578b', '\u5177\u4f53\u5185\u5bb9', '\u5904\u7f6e\u65b9\u6cd5', '\u5bb9\u5668\u6807\u8bc6'],
    [
        ['\u611f\u67d3\u6027\u5e9f\u5f03\u7269', '\u7ec6\u80de\u5830\u662f\u5c64\u3001\u6b8b\u4f59\u57f9\u517b\u57fa\u3001\u5c85\u5e03\u3001\u624b\u5957', '121\u00b0C\u9ad8\u538b\u84b8\u6c7120min\u9664\u83cc\u540e\u52a0\u76d6\u624b\u5957\u6536\u96c6\u8fd0\u9001\u5e9f\u5f03', '\u9ec4\u8272\u611f\u67d3\u6027\u5e9f\u5f03\u7269\u888b'],
        ['\u68d2\u9488\u3001\u5c16\u952e\u5e9f\u5f03\u7269', '9mL/P1000\u5c16\u5634\u5934\u3001\u9488\u5c16\uff08\u5982\u6709\u4f7f\u7528\uff09', '\u8033\u5c16\u5934\u88c5\u5165\u575a\u786c\u8010\u523a\u5371\u9669\u5e9f\u5f03\u7269\u7b52', '\u7ea2\u8272\u5371\u5bb3\u5e9f\u5f03\u7269\u7b52'],
        ['DMSO\u5316\u5b66\u5e9f\u5f03\u7269', '\u8fc7\u5c71\u7684\u82cf\u6728\u5de5\u4f5c\u6db2\u3001\u7a0b\u5f0f\u66f4\u6362\u5c71\u5c71\u7684DMSO', '\u6536\u96c6\u5230\u5c01\u53e3\u5bb9\u5668\uff0c\u6807\u6ce8\u201c\u6709\u673a\u5e9f\u6db2\u201d\uff0c\u4ea4\u5b66\u6821\u5e9f\u5f03\u7269\u5904\u7f6e\u6d88\u5a01', '\u6807\u6ce8\u6709\u673a\u5e9f\u6db2\u7684\u5c01\u53e3\u6d17\u6da4\u5b50'],
        ['\u666e\u901a\u751f\u6d3b\u5783\u573e', '\u5bf9\u7167\u5c04\u5f39\u7c97\u7684\u7b80\u4ecb\u5305\u88c5\u3001\u64cd\u4f5c\u624b\u5957\u7b49', '\u6536\u96c6\u5230\u9ed1\u8272\u5783\u573e\u888b\uff0c\u4e0d\u9700\u7279\u6b8a\u5904\u7f6e', '\u666e\u901a\u5783\u573e'],
    ],
    '\u88685-3  \u5b9e\u9a8c\u5e9f\u5f03\u7269\u5206\u7c7b\u5904\u7f6e\u89c4\u8303'
)

add_chapter_heading(doc, '5.4.3  \u7d27\u6025\u4e8b\u4ef6\u5904\u7f6e\u6d41\u7a0b', level=3)
for item in [
    '\u76ae\u80a4\u6216\u7c98\u819c\u6ee1\u6db2\uff1a\u7acb\u5373\u7528\u6e05\u6d01\u6d41\u5169\u5145\u5206\u51b2\u6d17>=15min\uff0c\u5e72\u8471\u6d88\u9999\u3001\u62a5\u544a\u5b9e\u9a8c\u5ba4\u8d1f\u8d23\u4eba',
    '\u4e00\u6b21\u6027\u6bee\u9489\u523a\u4f24\uff1a\u6c34\u51b2\u6d17\u4f24\u53e3\uff0c\u62a5\u544a\uff0c\u5fc5\u8981\u65f6\u4e0e\u533b\u7597\u673a\u6784\u8054\u7cfb',
    '\u5927\u91cf\u5bee\u6d69\u6cc4\u6f0f\uff1a\u5c11\u7528\u5ac9\u5f0f\u5e03\u5438\u5185\u5bf9\u9519\u5c55\u5e9f\u5f03\u7269\uff0c\u7528\u8131\u5b89\u6b87\u6d88\u6bd2\u89e3\u5c71\u5e73\uff0c\u6dd8\u6d17\u5c71\u5c40\u90e8\u4e1a\u4e8f\u8bbe\u5907',
]:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    run = p.add_run('\u25b6 ' + item)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=30)

# -------------------------------------------------------
# 5.5 Experiment record template requirements
# -------------------------------------------------------
add_chapter_heading(doc, '5.5  \u5b9e\u9a8c\u8bb0\u5f55\u4e0e\u5b58\u6863\u8981\u6c42', level=2)

add_body_para(doc,
    '\u89c4\u8303\u7684\u5b9e\u9a8c\u8bb0\u5f55\u662f\u7891\u9a8c\u8bba\u6587\u60c5\u62a5\u548c\u5b66\u672f\u8bda\u4fe1\u7684\u57fa\u7840\u3002\u672c\u8282\u89c4\u5b9a\u5b9e\u9a8c\u8bb0\u5f55\u7684\u6700\u4f4e\u6807\u51c6\u548c\u5b58\u6863\u8981\u6c42\u3002',
    indent=True)

add_three_line_table(doc,
    ['\u8bb0\u5f55\u5185\u5bb9', '\u5fc5\u987b\u8bb0\u5f55\u9879\u76ee', '\u5907\u6ce8'],
    [
        ['\u5b9e\u9a8c\u65e5\u671f\u4e0e\u64cd\u4f5c\u8005', '\u65e5\u671f\u3001\u8d77\u59cb/\u7ed3\u675f\u65f6\u95f4\u3001\u64cd\u4f5c\u8005\u59d3\u540d', '\u5b9e\u9a8c\u540e\u7b7e\u5b57\u786e\u8ba4'],
        ['\u7ec6\u80de\u4fe1\u606f', '\u4f20\u4ee3\u6b21\u6570\u3001\u590d\u82cf\u65e5\u671f\u3001\u6d3b\u529b\u68c0\u6d4b\u7ed3\u679c', '\u6bcf\u6b21\u94fa\u677f\u524d\u5fc5\u8bb0'],
        ['\u836f\u7269\u4fe1\u606f', '\u6279\u53f7\u3001\u5236\u5907\u65e5\u671f\u3001\u6bcd\u6db2\u6d53\u5ea6\u3001DMSO\u6d53\u5ea6\u9a8c\u8bc1', '\u5c0f\u6570\u70b9\u540e4\u4f4d\u5c0f\u6570'],
        ['\u5b9e\u9a8c\u8bbe\u8ba1\u56fe', '96\u5b54\u677f\u6392\u5217\u56fe\uff08\u624b\u7ed8\u6216\u6253\u5370\uff09', '\u6240\u6709\u5b54\u5747\u5e94\u6807\u6ce8'],
        ['\u539f\u59cb\u6570\u636e', '\u5168\u90e8OD\u8bfb\u6570\uff08\u624b\u5199\u6216\u6253\u5370\u7c98\u8d34\uff09', '\u4e0d\u5f97\u4e8b\u540e\u4fee\u6539'],
        ['\u5f02\u5e38\u5904\u7406', '\u4efb\u4f55\u88ab\u6392\u9664\u7684\u5b54\u6570\u636e\u5fc5\u987b\u8bb0\u5f55\u539f\u56e0', '\u5305\u62ec Grubbs \u68c0\u9a8c\u8fc7\u7a0b'],
        ['\u5f02\u5e38\u4e8b\u4ef6', '\u5b9e\u9a8c\u8fc7\u7a0b\u4e2d\u53d1\u751f\u7684\u4efb\u4f55\u610f\u5916\uff0c\u5982\u5c0f\u7c73\u6d88\u6bd2\u4e86\u3001\u5c81\u548c\u4e86\u3001\u6c22\u6c14\u6de8', '\u5929\u8fb9\u62a5\u5bfc\u4e0d\u5f97\u9690\u7792'],
        ['IC50\u6c47\u603b', '\u4e09\u6b21\u91cd\u590d\u7684 IC50 \u5747\u503c\u00b1SD\uff0c95%CI', '\u5bfc\u5e08\u7b7e\u5b57\u786e\u8ba4'],
    ],
    '\u88685-4  \u5b9e\u9a8c\u8bb0\u5f55\u5fc5\u5f55\u9879\u76ee\u6e05\u5355'
)

add_green_box(doc, '\u5b9e\u9a8c\u8bb0\u5f55\u7684\u6838\u5fc3\u539f\u5219', [
    '\u53ef\u6e90\u6027\uff1a\u4ed6\u4eba\u6309\u7167\u8bb0\u5f55\u5e94\u80fd\u5c06\u5b9e\u9a8c\u91cd\u73b0',
    '\u771f\u5b9e\u6027\uff1a\u5b9e\u9a8c\u65f6\u5373\u65f6\u8bb0\u5f55\uff0c\u4e0d\u5f97\u4e8b\u540e\u627e\u6b65\u8865\u8bb0\u6216\u4fee\u6539',
    '\u5b8c\u6574\u6027\uff1a\u8bb0\u5f55\u5e94\u5305\u62ec\u6240\u6709\u9519\u8bef\u548c\u5f02\u5e38\uff0c\u4e0d\u64a4\u9664\u5931\u8d25\u6570\u636e',
    '\u4e09\u5e74\u5b58\u6863\uff1a\u6240\u6709\u5b9e\u9a8c\u8bb0\u5f55\u5e94\u81f3\u5c11\u4fdd\u5b58\u4e09\u5e74\uff08\u552e\u820c\u5b78\u671f\u95f4\uff09',
])

doc.save(DOC_PATH)
print("Chapter 5 complete and saved.")
