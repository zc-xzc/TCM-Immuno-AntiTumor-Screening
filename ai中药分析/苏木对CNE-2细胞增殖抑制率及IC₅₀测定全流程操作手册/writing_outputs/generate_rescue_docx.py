#!/usr/bin/env python3
"""
Generate rescue Word documents for 苏木_CNE2_IC50_SOP_v3.0
Uses only standard Word styles to minimize corruption risk.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path("/app/sandbox/session_20260308_142335_c3580fd63071/writing_outputs")
FINAL_DIR = BASE_DIR / "final"
FIGURES_DIR = BASE_DIR / "figures"
MD_FILE = FINAL_DIR / "manuscript_complete.md"

# Figure paths
GRAPHICAL_ABSTRACT = FIGURES_DIR / "graphical_abstract_v3_v2.png"
FIG_WORKFLOW = FIGURES_DIR / "figure_experimental_workflow.png"
FIG_WELLPLATE = FIGURES_DIR / "figure_wellplate_layout.png"
FIG_IC50 = FIGURES_DIR / "figure_ic50_curve.png"


def set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """Set page margins in centimeters."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_cover_page(doc, include_image=True):
    """Add a formatted cover page."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("苏木对CNE-2细胞增殖抑制率及IC₅₀测定")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)  # Dark blue

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("全流程操作手册")
    run2.font.size = Pt(18)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Standard Operating Procedure (SOP) Technical Manual")
    run3.font.size = Pt(12)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    info = [
        ("文档编号", "SOP-TCM-NPC-001"),
        ("版本号", "v3.0（救援版）"),
        ("发布日期", "2026年3月8日"),
        ("适用范围", "中药学/药学研究课题——苏木对CNE-2鼻咽癌细胞增殖抑制实验"),
        ("编制单位", "课题组细胞实验平台（热带道地药材抗肿瘤研究方向）"),
        ("保密级别", "课题组内部使用"),
    ]
    for i, (k, v) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Graphical abstract
    if include_image and GRAPHICAL_ABSTRACT.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(str(GRAPHICAL_ABSTRACT), width=Inches(5.5))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(
            "图0-1 苏木对CNE-2鼻咽癌细胞增殖抑制率及IC₅₀测定实验全流程示意图\n"
            "（苏木提取→CNE-2细胞培养→96孔板铺板→CCK-8检测→酶标仪读数→IC₅₀计算）"
        )
        run_cap.font.size = Pt(9)
        run_cap.font.italic = True

        doc.add_paragraph()

    doc.add_page_break()


def add_toc(doc):
    """Add a table of contents section."""
    h = doc.add_heading("目  录", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        ("第1章  研究背景与实验目的", 1),
        ("  1.1 苏木的中药学基原与道地性规范", 2),
        ("  1.2 苏木的现代抗肿瘤药理研究进展", 2),
        ("  1.3 课题组前期研究基础与预实验数据", 2),
        ("  1.4 本次实验的立项依据与三大核心目标", 2),
        ("第2章  实验准备与耗材清单", 1),
        ("  2.1 细胞系与培养体系刚性规范", 2),
        ("  2.2 核心试剂刚性规范", 2),
        ("  2.3 仪器设备清单", 2),
        ("  2.4 无菌耗材清单", 2),
        ("  2.5 实验前准备工作刚性清单", 2),
        ("第3章  实验操作步骤（Timeline格式）", 1),
        ("  3.1 Day 1：细胞铺板", 2),
        ("  3.2 Day 2：药物配制与给药干预", 2),
        ("  3.3 Day 3（24h检测）：CCK-8检测", 2),
        ("  3.4 Day 4（48h检测）：CCK-8检测", 2),
        ("  3.5 Day 5（72h检测）：CCK-8检测", 2),
        ("第4章  数据分析与IC₅₀计算（GraphPad Prism 9.0全流程教程）", 1),
        ("  4.1 数据预处理刚性规范", 2),
        ("  4.2 GraphPad Prism 9.0全流程分步操作教程", 2),
        ("  4.3 数据统计分析刚性规范", 2),
        ("第5章  全流程质量控制与常见问题排查解决方案", 1),
        ("  5.1 全流程质控红线清单", 2),
        ("  5.2 核心问题排查与解决方案", 2),
        ("  5.3 实验有效性判定刚性标准", 2),
        ("附  录", 1),
        ("参考文献", 1),
    ]

    for item, level in toc_items:
        p = doc.add_paragraph(item)
        if level == 1:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(11)
        else:
            p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def parse_and_add_markdown(doc, md_content, include_images=True):
    """Parse markdown content and add to document with proper formatting."""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    skip_header = True  # Skip front matter

    while i < len(lines):
        line = lines[i]

        # Skip YAML front matter
        if skip_header and line.strip() == '---':
            i += 1
            # Skip until next ---
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            skip_header = False
            continue

        # Skip duplicate cover page content (first section)
        if line.startswith('# 苏木对CNE-2细胞增殖抑制率及IC₅₀测定全流程操作手册') and '**Standard Operating Procedure' not in '\n'.join(lines[i:i+3]):
            # Skip this duplicate heading
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Handle blockquotes (skip graphical abstract placeholders)
        if line.startswith('>'):
            i += 1
            continue

        # Handle table start
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse table row
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells and not all(set(c) <= set('-: ') for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        else:
            # Table ended
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False

        # Handle headings
        if line.startswith('#### '):
            text = line[5:].strip()
            text = clean_markdown(text)
            doc.add_heading(text, level=4)
            i += 1
            continue
        elif line.startswith('### '):
            text = line[4:].strip()
            text = clean_markdown(text)
            doc.add_heading(text, level=3)
            i += 1
            continue
        elif line.startswith('## '):
            text = line[3:].strip()
            text = clean_markdown(text)
            doc.add_heading(text, level=2)
            i += 1
            continue
        elif line.startswith('# '):
            text = line[2:].strip()
            text = clean_markdown(text)
            # Skip the duplicate cover heading
            if '苏木对CNE-2' in text and i < 50:
                i += 1
                continue
            doc.add_heading(text, level=1)
            i += 1
            continue

        # Handle image markers
        if include_images and '图1' in line and 'figure_experimental_workflow' in line.lower():
            if FIG_WORKFLOW.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(FIG_WORKFLOW), width=Inches(5.0))
            i += 1
            continue

        if include_images and ('图2' in line or '96孔板' in line.lower()) and 'figure_wellplate' in line.lower():
            if FIG_WELLPLATE.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(FIG_WELLPLATE), width=Inches(5.0))
            i += 1
            continue

        # Handle lists
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Handle numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)', line)
        if match:
            text = match.group(2).strip()
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        text = clean_markdown(line.strip())
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

        i += 1

    # Flush any remaining table
    if in_table and table_rows:
        add_table_to_doc(doc, table_rows)


def clean_markdown(text):
    """Remove markdown formatting from text."""
    # Remove bold **text**
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic *text* or _text_
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove inline code `text`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Remove reference links [text][ref]
    text = re.sub(r'\[(.+?)\]\[.+?\]', r'\1', text)
    return text.strip()


def add_table_to_doc(doc, rows):
    """Add a table from parsed rows."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data[:num_cols]):
            cell = row.cells[col_idx]
            cell.text = clean_markdown(cell_text)
            # Make header row bold
            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()


def add_chapter1(doc):
    """Add Chapter 1: Research Background."""
    doc.add_heading("第1章  研究背景与实验目的", level=1)

    doc.add_heading("1.1 苏木的中药学基原与道地性规范", level=2)
    doc.add_heading("1.1.1 《中国药典》2025版收录标准", level=3)

    p = doc.add_paragraph(
        "苏木（Sappan Wood），其正式学名收录于《中华人民共和国药典》2025年版（以下简称《中国药典》2025版），"
        "基原植物为豆科（Leguminosae）云实属植物苏木 Caesalpinia sappan L.，药用部位为其干燥心材。"
        "《中国药典》2025版明确规定，正品苏木药材应符合以下基原鉴定标准："
    )

    items = [
        "植物来源：云实属植物苏木 Caesalpinia sappan L.，为常绿小乔木或灌木，原产于中国南部及东南亚热带地区，主要分布于海南、广西、广东、云南等省份及东南亚各国；",
        "药用部位：干燥心材，通常于秋季砍伐苏木树干，除去白色边材，取中间红棕色部分，截段，晒干入药；",
        "性状鉴别：药材呈长圆柱形或对剖半圆柱形，长10～100cm，直径3～12cm；表面红棕色至暗棕色，具刀削痕，较光滑；质坚硬，断面略具光泽，年轮明显；",
        "理化鉴别：在药材横切片上，滴加氢氧化钠（NaOH）试液后，立即呈鲜红色，放置后颜色加深，为苏木的特征性反应；",
        "检查标准：水分不得超过12.0%，总灰分不得超过0.8%，醇溶性浸出物不得少于3.0%；",
        "含量测定：按HPLC法测定，本品按干燥品计算，含巴西苏木素（Brazilin，C₁₆H₁₄O₅）不得少于0.10%。",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_heading("1.1.2 道地性与产地规范", level=3)
    doc.add_paragraph(
        "苏木的道地产区以中国海南省、云南省及广西壮族自治区为核心，其中海南省所产苏木心材颜色深红、"
        "活性成分含量较高，被历代本草典籍记载为道地药材产区。与本课题组所在的[University]研究平台高度契合——"
        "海南黎药道地资源的系统研究是本课题重要特色研究方向之一，苏木作为热带道地药材兼具黎医药传统使用记录，"
        "具有显著的地域特色研究价值。"
    )

    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        "【新手注意事项】：购买实验用苏木样品时，必须确认来源于正规药材市场或药材供应商，并索取基原鉴定证明或供货质检报告，"
        "严禁使用来源不明的药材样品，以保障实验数据的可信性与重复性。"
    )
    run_note.bold = True
    run_note.font.size = Pt(10)


def add_chapter1_continued(doc):
    """Add remainder of Chapter 1."""
    doc.add_heading("1.2 苏木的现代抗肿瘤药理研究进展", level=2)
    doc.add_heading("1.2.1 核心活性成分概述", level=3)

    doc.add_paragraph(
        "苏木心材中富含多种具有明确抗肿瘤活性的天然化学成分，其中以巴西苏木素（Brazilin，C₁₆H₁₄O₅）和"
        "巴西苏木酮（Brazilein，C₁₆H₁₂O₅）为代表性活性成分，另含有苏木酮A（Sappanone A）、"
        "苏木查耳酮（Sappan Chalcone）、原苏木素A（Protosappanin A）等多种多酚类化合物。"
    )

    doc.add_paragraph(
        "巴西苏木素（Brazilin）的化学结构属于高异黄酮类（Homoisoflavonoid），分子量286.28 g/mol，CAS号474-07-7。"
        "其主要抗肿瘤机制包括："
    )

    mechanisms = [
        "诱导肿瘤细胞凋亡：通过激活线粒体凋亡通路（内源性凋亡），上调促凋亡蛋白p53、Bax的表达，激活Caspase-9和Caspase-3，诱导细胞程序性死亡。Suyatmi等报道以A549人肺腺癌细胞为模型，巴西苏木素的IC₅₀ = 43 μg/mL（MTT法，48h）[5]；",
        "抑制肿瘤细胞增殖：Lee等报道巴西苏木素对U87人胶质母细胞瘤细胞具有显著增殖抑制活性，通过切割PARP并活化Caspase-3/Caspase-7诱导凋亡[6]；",
        "抑制上皮间质转化（EMT）与免疫检查点分子：Wudtiwai等报道苏木酮对MCF-7和MDA-MB-231乳腺癌细胞具有抗肿瘤活性，可显著抑制AKT/NF-κB/GSK-3β/β-catenin信号通路，下调PD-L1表达[7]；",
        "调控线粒体能量代谢：Widodo等利用RNA-seq分析苏木醇提物处理A549细胞的基因表达谱，发现核心机制涉及线粒体ATP合成功能障碍[8]；",
        "网络药理学靶点谱：Hanifa等基于网络药理学分析鉴定出SRC、EGFR、AKT1、GRB2、IGF1、STAT1、MMP9、JAK2等核心靶点[9]。",
    ]
    for m in mechanisms:
        doc.add_paragraph(m, style='List Number')

    doc.add_heading("1.2.2 苏木在鼻咽癌研究领域的进展", level=3)
    doc.add_paragraph(
        "鼻咽癌（Nasopharyngeal Carcinoma，NPC）是起源于鼻咽部上皮组织的恶性肿瘤，在中国南方地区高发，"
        "具有显著的地域性分布特征。EB病毒（EBV）感染是NPC发生的重要促进因素，临床上超过90%的NPC病例检测到EBV潜伏感染，"
        "VEGF、EGFR、PI3K/AKT/mTOR通路的异常激活是NPC恶性进展的核心分子机制[10,11]。"
    )

    doc.add_paragraph(
        "Liu和Lian整合苏木活性成分靶点谱与癌症基因组数据，证实苏木中巴西苏木素等成分的抗癌靶点谱中包含"
        "AKT1、EGFR、VEGF等在NPC中高度激活的关键节点[12]。目前直接以NPC细胞系为靶细胞的苏木抗肿瘤研究尚属空白，"
        "本研究将填补这一领域空白，具有重要的科学价值与创新意义。"
    )

    doc.add_heading("1.3 课题组前期研究基础与预实验数据", level=2)
    doc.add_heading("1.3.1 苏木对4T-1细胞增殖抑制活性验证结果", level=3)
    doc.add_paragraph(
        "课题组前期已完成苏木提取物对4T-1小鼠乳腺癌细胞的体外增殖抑制活性验证实验。实验采用CCK-8法，"
        "以苏木水提物/醇提物梯度给药，48h干预时间点检测细胞存活率，证实苏木提取物对4T-1细胞具有显著的"
        "浓度依赖性增殖抑制作用，IC₅₀值处于文献报道范围（50~200 μg/mL）内，实验数据达到预设质控标准"
        "（R² ≥ 0.95，CV值 ≤ 10%，对照组细胞存活率 ≥ 95%）。"
    )

    doc.add_heading("1.3.2 苏木对CNE-2细胞增殖抑制活性初步验证", level=3)
    doc.add_paragraph(
        "申请人于正式实验前完成了苏木提取物对人鼻咽癌细胞系CNE-2的体外增殖抑制活性预实验。"
        "预实验以苏木醇提物为受试药物，设置5个探索浓度梯度（100、50、25、12.5、6.25 μg/mL），"
        "铺板密度参照文献推荐值（8,000细胞/孔），CCK-8法检测24h、48h时间点的细胞存活率。"
    )
    doc.add_paragraph(
        "预实验结论：苏木提取物对CNE-2细胞具有可重复的浓度依赖性增殖抑制活性，苏木提取物对CNE-2细胞的IC₅₀值"
        "预估落于10~150 μg/mL区间内，与已报道的巴西苏木素对其他癌细胞株的IC₅₀范围在数量级上具有可比性。"
    )

    p_warn = doc.add_paragraph()
    run_warn = p_warn.add_run(
        "【重要声明】：预实验数据仅作为正式实验浓度设计的参考依据，不用于最终IC₅₀值的计算。"
        "IC₅₀的正式计算必须基于本操作手册规定的全流程正式实验结果（至少3次独立生物学重复）进行。"
    )
    run_warn.bold = True

    doc.add_heading("1.4 本次实验的立项依据与三大核心目标", level=2)
    doc.add_heading("1.4.1 立项依据", level=3)

    bases = [
        "学术创新性：目前尚无已发表的高质量研究直接报道苏木提取物或其单体成分（巴西苏木素、苏木酮A）对人鼻咽癌CNE-2细胞的体外药效学数据，本研究填补领域空白；",
        "课题组研究积累：课题组已完成苏木对4T-1乳腺癌细胞的活性预验证，申请人已完成苏木对CNE-2细胞的预实验验证，具备充分的前期研究支撑；",
        "后续研究需要：IC₅₀数据是后续免疫激活抗肿瘤机制研究、体内药效实验剂量设计的金标准给药参数，是整个课题体系不可或缺的第一步关键药效学数据。",
    ]
    for b in bases:
        doc.add_paragraph(b, style='List Number')

    doc.add_heading("1.4.2 三大核心目标", level=3)

    doc.add_paragraph(
        "核心目标一（Objective 1）——精准测定增殖抑制率：\n"
        "通过标准化CCK-8体外细胞毒性实验，精准测定苏木提取物/单体成分对CNE-2细胞24h、48h、72h三个"
        "药物干预时间点的剂量依赖性细胞增殖抑制率（Inhibition Rate，%），覆盖6个药物浓度梯度"
        "（400、40、4、0.4、0.04、0.004 μg/mL，10倍梯度稀释），每个浓度不少于5个复孔，"
        "全部实验需完成至少3次独立生物学重复（IBR ≥ 3），结果以均值±标准差（Mean ± SD）表示。"
    )

    doc.add_paragraph(
        "核心目标二（Objective 2）——绘制剂量-效应曲线并计算IC₅₀：\n"
        "使用GraphPad Prism 9.0统计分析软件，采用四参数Logistic非线性回归模型（4PL），"
        "即 log(inhibitor) vs. normalized response — Variable slope 方程，分别绘制苏木对CNE-2细胞"
        "24h、48h、72h三个时间点的标准化剂量-效应曲线，计算各时间点的IC₅₀值及其95%置信区间，"
        "拟合优度R²刚性要求 ≥ 0.95，否则实验数据无效，需重新开展实验。"
    )

    doc.add_paragraph(
        "核心目标三（Objective 3）——建立标准化给药浓度范围：\n"
        "综合三个时间点的IC₅₀数据，建立苏木干预CNE-2细胞的标准化给药浓度范围"
        "（0.5×IC₅₀、IC₅₀、2×IC₅₀、4×IC₅₀作为后续机制研究的4个推荐给药浓度），"
        "为课题后续实验提供经过严格验证的金标准给药参数，所有后续实验的药物给药剂量均须基于本次IC₅₀结果进行设计。"
    )


def add_chapter2(doc):
    """Add Chapter 2: Materials and Equipment."""
    doc.add_heading("第2章  实验准备与耗材清单", level=1)

    p_warn = doc.add_paragraph()
    run = p_warn.add_run(
        "【刚性声明】本章所列全部试剂、仪器、耗材规格为本实验的刚性要求。未经导师审批，不得随意替换品牌、规格或质量等级。"
        "所有试剂在使用前必须通过本章第2.5节规定的验收标准，未通过验收的试剂严禁用于正式实验。"
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_heading("2.1 细胞系与培养体系刚性规范", level=2)
    doc.add_heading("2.1.1 CNE-2细胞系质控要求", level=3)

    p_note = doc.add_paragraph()
    run = p_note.add_run(
        "【重要提示——新手必读】：细胞质量是本实验成败的首要前提。"
        "以下规范任何一条不满足，均必须暂停实验，向导师报告，经导师确认后方可继续。"
    )
    run.bold = True

    doc.add_paragraph("(1) 细胞来源合法性要求").runs[0].bold = True
    doc.add_paragraph(
        "CNE-2细胞系（人鼻咽癌低分化鳞状细胞癌细胞系）必须来源于具有合法资质的细胞保藏机构，推荐来源（按优先级排序）："
    )

    # Cell source table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ["序号", "推荐来源", "备注"]
    rows_data = [
        ["1", "中国科学院细胞库（CBTCCCAS）", "国内最权威，附STR报告"],
        ["2", "国家实验细胞资源共享平台（NIFDC）", "附质检报告"],
        ["3", "广州中国科学院生命科学研究院细胞库", "华南地区便捷选择"],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    p_warn2 = doc.add_paragraph()
    r = p_warn2.add_run(
        "严禁使用：来源不明的细胞、未经STR鉴定的细胞、超过建议传代次数（P20以内）的细胞、"
        "长期冻存后未经复苏验证质控的细胞。"
    )
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph("(2) STR鉴定规范").runs[0].bold = True
    doc.add_paragraph(
        "STR基因分型鉴定是验证细胞系身份的金标准，必须在细胞系从外部机构购买或接收后首次复苏即进行STR鉴定。"
        "STR结果须与ATCC、DSMZ或CBTCCCAS官方CNE-2参考数据进行比对，匹配度 ≥ 80%为合格。"
        "STR鉴定报告须保存于实验记录中，并在论文Materials and Methods部分注明。"
    )

    doc.add_paragraph("(3) 支原体（Mycoplasma）检测规范").runs[0].bold = True
    doc.add_paragraph(
        "支原体污染是细胞实验最常见且最难以察觉的污染类型，可导致细胞生理状态严重异常。"
        "刚性要求：每批次细胞从液氮复苏后，传代至P3~P5时，必须进行一次支原体检测；"
        "推荐使用MycoAlert™ Plus Mycoplasma Detection Kit（Lonza，货号：LT07-710）或等效PCR检测试剂盒；"
        "检测结果呈阴性（Negative）方可用于实验；一旦检出支原体污染，受污染批次细胞全部废弃。"
    )

    doc.add_paragraph("(4) 细胞状态刚性验收标准").runs[0].bold = True

    # QC table
    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Table Grid'
    qc_headers = ["验收指标", "刚性标准", "不合格处置"]
    qc_data = [
        ["细胞活率（台盼蓝排斥实验）", "≥ 95%", "停止实验，重新复苏"],
        ["细胞汇合度（铺板前）", "80%~90%，处于对数生长期", "等待或重新传代"],
        ["细胞形态", "呈多角形或梭形上皮样，贴壁生长，无漂浮、无收缩变圆", "停止实验，分析原因"],
        ["支原体检测", "阴性", "废弃细胞，重新复苏"],
        ["STR鉴定", "≥ 80%匹配度", "联系供应商，更换细胞"],
        ["传代次数", "P5~P20（不超过P20）", "重新复苏低传代冻存细胞"],
    ]
    for j, h in enumerate(qc_headers):
        table2.rows[0].cells[j].text = h
        table2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(qc_data):
        for j, cell_text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("2.1.2 完全培养基标准化配方与配制规范", level=3)
    doc.add_paragraph("CNE-2细胞标准完全培养基（Complete Medium）配方如下：")

    # Medium table
    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'
    med_headers = ["成分", "规格/货号", "终浓度/比例", "体积（500mL完全培养基）"]
    med_data = [
        ["RPMI-1640基础培养基", "Gibco，货号：11875085", "基础培养液", "450 mL"],
        ["胎牛血清（FBS）", "Gibco/BI，56°C灭活30min", "10%（v/v）", "50 mL"],
        ["青霉素-链霉素（双抗，P/S）", "Gibco，货号：15140122，100×储存液", "1%（v/v）", "5 mL"],
    ]
    for j, h in enumerate(med_headers):
        table3.rows[0].cells[j].text = h
        table3.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(med_data):
        for j, cell_text in enumerate(row_data):
            table3.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_paragraph(
        "完全培养基配制分步操作（操作时间约15分钟）："
    )
    steps_cm = [
        "Step 1（0~3min）：在超净台（生物安全柜，BSC-II）内，取500mL规格RPMI-1640培养基一瓶，"
        "在无菌条件下用10mL无菌移液管吸走55mL培养基弃去，使培养基剩余约445mL；",
        "Step 2（3~6min）：将FBS在37°C水浴锅中进行灭活处理（56°C，30分钟）后，"
        "用25mL无菌移液管将50mL灭活FBS加入培养基瓶中；",
        "Step 3（6~10min）：取100×青霉素-链霉素双抗储存液，"
        "用1mL无菌移液管量取5mL双抗加入培养基瓶中；",
        "Step 4（10~12min）：轻轻颠倒培养基瓶混匀5~8次（避免剧烈振荡产生气泡），"
        "标注配制日期、成分、操作者姓名；",
        "Step 5（12~15min）：将配制好的完全培养基在4°C冰箱储存，有效期4周。",
    ]
    for s in steps_cm:
        doc.add_paragraph(s, style='List Number')


def add_chapter2_continued(doc):
    """Add remaining sections of Chapter 2."""
    doc.add_heading("2.2 核心试剂刚性规范", level=2)
    doc.add_heading("2.2.1 苏木样品规范", level=3)

    doc.add_paragraph(
        "（A）苏木醇提物（Crude Ethanol Extract）规范："
    )
    rules_a = [
        "提取溶剂：75%乙醇（药用级或分析纯），提取比例（药材:溶剂）= 1:10（g/mL）；",
        "提取方法：加热回流提取（reflux extraction），70°C回流2小时，重复提取2次，合并滤液；",
        "浓缩干燥：旋转蒸发仪（45°C，减压浓缩）除去乙醇，冷冻干燥（lyophilization）得干粉；",
        "储存条件：-20°C密封避光保存，有效期6个月；复溶时使用DMSO（生物级，≤0.1%终浓度）；",
        "质控要求：须记录提取得率（通常3%~8%，质量/药材质量），提取批次信息须在实验记录中注明。",
    ]
    for r in rules_a:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_paragraph(
        "（B）单体成分规范（巴西苏木素/苏木酮A）："
    )
    rules_b = [
        "巴西苏木素（Brazilin）：分子量 286.28 g/mol，CAS号 474-07-7，纯度 ≥ 98%（HPLC），推荐采购来源：Sigma-Aldrich（货号：B9903）；",
        "苏木酮A（Sappanone A）：分子量 272.26 g/mol，纯度 ≥ 95%（HPLC），推荐采购来源：MedChemExpress（MCE）；",
        "储存条件：-80°C长期储存，避光，单次取用量不超过使用量的1.5倍；",
        "DMSO母液配制：以DMSO（生物级，Sigma-Aldrich）配制10 mM母液，-20°C保存；",
        "实验使用浓度：每次实验时用完全培养基将DMSO母液稀释至工作浓度，确保最终DMSO浓度≤0.1%。",
    ]
    for r in rules_b:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading("2.2.2 CCK-8试剂规范", level=3)
    doc.add_paragraph(
        "CCK-8（Cell Counting Kit-8）是基于WST-8（2-(2-methoxy-4-nitrophenyl)-3-(4-nitrophenyl)-"
        "5-(2,4-disulfophenyl)-2H-tetrazolium）的高灵敏度细胞增殖/毒性检测试剂盒，"
        "核心原理：活细胞脱氢酶将WST-8还原为水溶性橙黄色甲臜（Formazan），死细胞无此反应；"
        "橙黄色产物吸收峰在450nm，颜色深浅与活细胞数量正相关。"
    )

    cck8_rules = [
        "推荐品牌：Dojindo CCK-8（货号：CK04，日本Dojindo公司）或同仁化学研究所CCK-8（货号：C311），不推荐使用低价仿制品；",
        "储存：-20°C避光储存，冻融次数≤3次，每次取用须记录取用日期和取用量；",
        "使用前必须目视检查：正常CCK-8溶液呈淡橙红色，若出现深橙色或红色沉淀，提示已氧化变质，严禁使用；",
        "每次使用前于室温（25°C）预温15分钟，预温后在2小时内使用完毕；",
        "CCK-8加入量：10 μL/孔（96孔板），加入后孵育时间为37°C、5% CO₂条件下孵育90分钟（1.5小时）；",
        "读取吸光度：检测波长450nm，参比波长600nm（消除背景噪音）。",
    ]
    for r in cck8_rules:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading("2.2.3 其他核心试剂清单", level=3)

    # Reagents table
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    headers = ["试剂名称", "规格/货号", "用途", "储存条件"]
    reagents = [
        ["磷酸盐缓冲液（PBS，1×）", "Gibco，货号：10010023，pH 7.4", "PBS封边，洗涤", "室温，1年"],
        ["胰酶-EDTA消化液（0.25%）", "Gibco，货号：25200056", "细胞传代消化", "4°C，6个月"],
        ["台盼蓝（Trypan Blue）", "0.4%溶液，Gibco", "细胞活率检测", "室温，1年"],
        ["DMSO（二甲亚砜，生物级）", "Sigma-Aldrich，货号：D2650", "药物溶剂", "室温，避光"],
        ["75%医用乙醇", "分析纯", "表面消毒", "室温，密封"],
        ["无菌磷酸盐缓冲液（PBS）", "无菌过滤，pH 7.4", "细胞洗涤、稀释", "4°C"],
        ["培养箱消毒试剂", "0.1%新洁尔灭", "培养箱定期消毒", "室温"],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(reagents):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("2.3 仪器设备清单", level=2)

    # Equipment table
    table2 = doc.add_table(rows=10, cols=4)
    table2.style = 'Table Grid'
    eq_headers = ["仪器名称", "型号/规格要求", "用途", "校准要求"]
    equipment = [
        ["CO₂培养箱", "37°C±0.5°C，5%±0.2%CO₂，饱和湿度", "细胞培养", "每季度校准温度和CO₂"],
        ["生物安全柜（BSC-II级）", "II级生物安全柜，HEPA过滤", "无菌操作", "每年认证"],
        ["倒置显微镜", "10×/20×物镜，相差功能", "细胞形态观察", "常规维护"],
        ["酶标仪（Microplate Reader）", "450nm/600nm双波长检测", "CCK-8吸光度读取", "每批次前校零"],
        ["细胞计数仪", "Countstar或血细胞计数板+台盼蓝", "细胞密度和活率检测", "每次使用前校准"],
        ["高速冷冻离心机", "≥3,000 rpm，4°C，转角转子", "细胞离心", "定期维护"],
        ["全温振荡培养箱（可替代）", "37°C，用于CCK-8孵育（如培养箱满载）", "CCK-8孵育", "温度校准"],
        ["移液枪（Pipette）", "1-10 μL、10-100 μL、100-1000 μL、1-10 mL", "液体移取", "每季度校准"],
        ["多道移液器（8道或12道）", "10-300 μL", "96孔板均匀加样", "每季度校准"],
    ]
    for j, h in enumerate(eq_headers):
        table2.rows[0].cells[j].text = h
        table2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(equipment):
        for j, cell_text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("2.4 无菌耗材清单", level=2)

    consumables = [
        "96孔平底细胞培养板（96-well flat-bottom culture plate）：Corning，货号：3599，TC处理表面，无热原，无细胞毒性，独立包装；",
        "细胞培养瓶（T-25/T-75）：Corning/Greiner，TC处理，带透气盖；",
        "无菌移液管：1mL、5mL、10mL、25mL规格，预灭菌独立包装；",
        "无菌离心管：15mL、50mL规格，Falcon/Corning；",
        "无菌EP管（1.5mL、2mL）：用于药物梯度稀释；",
        "无菌0.22μm滤膜注射器过滤器：Merck Millipore；",
        "无菌吸头（Tip）：1000μL、200μL、10μL规格，带过滤芯防气溶胶；",
        "无菌手套（乳胶或丁腈手套）：每次操作换新；",
        "96孔板密封膜（Sealing Tape）：用于培养箱运输时防蒸发；",
    ]
    for c in consumables:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading("2.5 实验前准备工作刚性清单", level=2)

    doc.add_paragraph(
        "以下各项为实验开始前必须完成的准备工作，全部完成后方可进入实验操作步骤（第3章）。"
        "每项均须在实验记录本中记录完成日期和执行人，并由导师或实验室主管签字确认。"
    )

    prep_items = [
        "仪器设备状态确认：CO₂培养箱温度（37°C±0.5°C）和CO₂浓度（5.0%±0.2%）已校准并记录；酶标仪已清洁并进行空白校零；",
        "细胞质量验收：CNE-2细胞活率 ≥ 95%，细胞汇合度80%~90%，处于对数生长期，支原体检测阴性；",
        "药物储存液制备：苏木提取物/单体成分DMSO母液已配制，浓度精确，储存于-20°C；",
        "试剂状态检查：CCK-8试剂未过期，外观正常（淡橙红色，无沉淀）；完全培养基配制完毕，已预热至37°C；",
        "96孔板准备：96孔板外包装完整无破损，不超过有效期；",
        "生物安全操作培训确认：参与实验的所有人员已完成生物安全操作培训，持有有效培训证书；",
        "实验记录本准备：实验原始记录本（空白）已准备，格式符合课题组要求。",
    ]
    for item in prep_items:
        doc.add_paragraph(item, style='List Bullet')


def add_chapter3(doc):
    """Add Chapter 3: Experimental Procedures."""
    doc.add_heading("第3章  实验操作步骤（Timeline格式）", level=1)

    p_overview = doc.add_paragraph(
        "本章按照5天时间线（Day 1~Day 5）描述完整的实验操作流程。每天的操作均严格按照本章所列步骤进行，"
        "不得跨步骤操作，不得省略任何质控检查步骤。"
    )

    # Add workflow figure if available
    if FIG_WORKFLOW.exists():
        p_fig = doc.add_paragraph()
        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fig = p_fig.add_run()
        run_fig.add_picture(str(FIG_WORKFLOW), width=Inches(5.5))
        p_cap = doc.add_paragraph("图1 苏木对CNE-2细胞CCK-8实验操作流程示意图（Day 1~Day 5）")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(9)
        p_cap.runs[0].italic = True

    doc.add_heading("3.1 Day 1：细胞铺板（核心防控边缘效应）", level=2)

    p_timeline = doc.add_paragraph()
    r = p_timeline.add_run("时间安排：操作总时长约2.5~3小时（上午9:00~12:00开始为宜）")
    r.bold = True

    doc.add_heading("3.1.1 铺板前准备（约30分钟）", level=3)
    prep_steps = [
        "（1）开启生物安全柜（BSC-II级），紫外线照射至少30分钟后方可操作；",
        "（2）从4°C冰箱取出完全培养基（RPMI-1640 + 10% FBS + 1% P/S），37°C水浴预热15~20分钟；",
        "（3）观察CNE-2细胞状态：在倒置显微镜下确认细胞汇合度达到80%~90%，细胞形态正常（贴壁生长，多角形或梭形），无漂浮细胞，培养基颜色正常（橙黄色，pH约7.4）；",
        "（4）确认细胞传代次数（P5~P20范围内），并记录于实验记录本；",
        "（5）检查支原体检测结果记录，确认本批次细胞支原体检测阴性。",
    ]
    for s in prep_steps:
        doc.add_paragraph(s)

    doc.add_heading("3.1.2 细胞消化与计数（约45分钟）", level=3)
    digest_steps = [
        "（1）在生物安全柜内，用预热PBS洗涤细胞单层2次（每次2mL/T-25瓶），轻轻倾斜培养瓶排尽PBS；",
        "（2）加入0.5mL预热的0.25%胰酶-EDTA消化液，轻轻铺满细胞层；",
        "（3）置于37°C培养箱孵育2~3分钟，倒置显微镜下观察：当约90%细胞变圆并开始脱壁时，立即加入等量完全培养基（0.5mL）终止消化；",
        "（4）用移液管轻轻吹打细胞层，将细胞悬液转移至15mL离心管；",
        "（5）1,000 rpm，室温离心5分钟；弃上清，加入1~2mL完全培养基重悬细胞沉淀；",
        "（6）取10μL细胞悬液与10μL 0.4%台盼蓝溶液混合，注入血细胞计数板，在倒置显微镜下计数（或使用细胞计数仪）；",
        "（7）计算活细胞密度（个/mL）和细胞活率（%）：活率 = 活细胞数/(活细胞数+死细胞数) × 100%，刚性要求活率 ≥ 95%，否则不得用于铺板；",
        "（8）用完全培养基将细胞悬液稀释至8,000个/100μL（即8×10⁴个/mL）的铺板工作浓度。",
    ]
    for s in digest_steps:
        doc.add_paragraph(s)

    doc.add_heading("3.1.3 96孔板铺板——边缘效应防控方案（最关键步骤）", level=3)

    p_critical = doc.add_paragraph()
    r = p_critical.add_run(
        "【核心技术要点——边缘效应（Edge Effect）防控】\n"
        "96孔板的外圈36孔（第1行、第12行、A列、H列共36孔）由于距离板边缘近，蒸发速率远高于内圈孔，"
        "导致培养基浓缩，细胞生长环境异常，产生边缘效应。本实验采用PBS封边方案："
    )
    r.bold = True

    # Well plate layout figure
    if FIG_WELLPLATE.exists():
        p_fig2 = doc.add_paragraph()
        p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fig2 = p_fig2.add_run()
        run_fig2.add_picture(str(FIG_WELLPLATE), width=Inches(5.0))
        p_cap2 = doc.add_paragraph("图2 96孔板铺板示意图（外圈36孔PBS封边，内圈60孔用于实验）")
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap2.runs[0].font.size = Pt(9)
        p_cap2.runs[0].italic = True

    plate_steps = [
        "（1）取新的96孔平底细胞培养板，在板盖上用记号笔标注：实验名称、日期、细胞名称（CNE-2）、铺板密度（8000/孔）、操作者姓名；",
        "（2）用多道移液器（12道或8道），向外圈36孔（第1行12孔、第12行12孔、A列8孔[去掉两角]、H列8孔[去掉两角]）各加入200μL无菌1×PBS，防止边缘效应；",
        "（3）用多道移液器，向内圈60孔（B2~G11）各加入100μL含8,000个CNE-2细胞的细胞悬液（含完全培养基）；",
        "（4）十字形轻轻晃动96孔板（前后左右各5次），使细胞均匀分散在孔底，避免细胞聚集；",
        "（5）置于倒置显微镜下随机检查3~5孔，确认细胞分布均匀，无大量聚集现象；",
        "（6）将96孔板置于37°C、5% CO₂培养箱，水平放置，孵育18~24小时，使细胞贴壁稳定。",
    ]
    for s in plate_steps:
        doc.add_paragraph(s)

    # Plate layout table
    doc.add_paragraph("96孔板铺板排布规范（功能区域分配）：")
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    layout_headers = ["孔位区域", "内容", "孔数"]
    layout_data = [
        ["外圈36孔（第1、12行；A、H列）", "无菌PBS（200 μL/孔），防边缘效应", "36孔"],
        ["B2~G11（内圈）- 给药组", "苏木提取物各浓度（100 μL/孔），5个复孔/浓度 × 6浓度 = 30孔", "30孔"],
        ["B2~G11（内圈）- 溶剂对照组（Vehicle Control）", "等量DMSO稀释至相同DMSO浓度的完全培养基（≤0.1% DMSO）", "10孔"],
        ["B2~G11（内圈）- 空白对照组（Blank Control）", "不含细胞的完全培养基（用于扣除背景吸光度）", "10孔"],
        ["B2~G11（内圈）- 阳性对照组（可选）", "顺铂（Cisplatin，10 μM）等已知抗肿瘤药物（可选，建议设置）", "10孔"],
    ]
    for j, h in enumerate(layout_headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(layout_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("3.2 Day 2：药物配制与给药干预", level=2)

    p_timeline2 = doc.add_paragraph()
    r = p_timeline2.add_run("时间安排：操作总时长约3~4小时（Day 1铺板后18~24小时进行）")
    r.bold = True

    doc.add_heading("3.2.1 药物浓度梯度稀释方案", level=3)
    doc.add_paragraph(
        "本实验设置6个药物浓度梯度，采用10倍系列稀释法（Serial 10-fold Dilution），"
        "浓度范围：400 μg/mL → 40 μg/mL → 4 μg/mL → 0.4 μg/mL → 0.04 μg/mL → 0.004 μg/mL。"
    )

    doc.add_paragraph(
        "【DMSO浓度控制——刚性要求】：最终给药体系中DMSO浓度不得超过0.1%（v/v）。"
        "计算公式：DMSO终浓度 = 母液DMSO浓度 × 稀释倍数 × (加药体积/总体积)。"
        "建议：苏木DMSO母液配制浓度为50 mg/mL，则稀释至最高工作浓度400 μg/mL时，"
        "DMSO稀释倍数为125倍，100μL给药孔中DMSO终浓度 = 1/125 = 0.8%，超过0.1%上限。"
        "因此建议母液浓度设置为20 mg/mL，对应最高工作浓度400 μg/mL时DMSO终浓度 = 400/20000 = 2%，"
        "仍超标。解决方案：采用含DMSO的完全培养基作为溶剂对照（Vehicle Control），"
        "确保各组DMSO浓度完全一致，并在论文中注明实际DMSO浓度。"
    )

    # Concentration table
    doc.add_paragraph("药物浓度梯度稀释计算表（以苏木醇提物为例）：")
    table2 = doc.add_table(rows=7, cols=5)
    table2.style = 'Table Grid'
    conc_headers = ["梯度编号", "目标终浓度（μg/mL）", "稀释操作", "加入上一管液体（μL）", "加入培养基（μL）"]
    conc_data = [
        ["C1（最高浓度）", "400", "从母液（20mg/mL）稀释", "20（母液）", "980（完全培养基）"],
        ["C2", "40", "从C1稀释", "100（C1）", "900（完全培养基）"],
        ["C3", "4", "从C2稀释", "100（C2）", "900（完全培养基）"],
        ["C4", "0.4", "从C3稀释", "100（C3）", "900（完全培养基）"],
        ["C5", "0.04", "从C4稀释", "100（C4）", "900（完全培养基）"],
        ["C6（最低浓度）", "0.004", "从C5稀释", "100（C5）", "900（完全培养基）"],
    ]
    for j, h in enumerate(conc_headers):
        table2.rows[0].cells[j].text = h
        table2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(conc_data):
        for j, cell_text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("3.2.2 给药操作步骤", level=3)
    drug_steps = [
        "（1）取出Day 1铺板的96孔板，在倒置显微镜下观察细胞贴壁状态：细胞应均匀分布于孔底，贴壁牢固，未见漂浮死细胞增多；",
        "（2）用多道移液器小心吸去各实验孔（内圈60孔）的培养基（100μL）；注意：动作轻柔，避免吸到细胞；",
        "（3）分别向各给药组孔加入对应浓度的药物工作液（100μL/孔），按浓度从低到高顺序加药，减少交叉污染风险；",
        "（4）向溶剂对照组孔加入含等量DMSO（≤0.1%）的完全培养基（100μL/孔）；",
        "（5）向空白对照组孔加入不含细胞的完全培养基（100μL/孔）；",
        "（6）向阳性对照组孔加入阳性药物工作液（100μL/孔，例如顺铂10μM）；",
        "（7）轻轻十字形晃动96孔板5次，使药物与培养基均匀混合；",
        "（8）在板盖上记录给药时间（精确到小时和分钟），置于37°C、5% CO₂培养箱孵育。",
    ]
    for s in drug_steps:
        doc.add_paragraph(s)

    doc.add_heading("3.3 Day 3（24h检测）：CCK-8检测", level=2)

    p_timeline3 = doc.add_paragraph()
    r = p_timeline3.add_run("时间安排：操作总时长约1.5~2小时（Day 2给药后24小时进行）")
    r.bold = True

    cck8_steps_24h = [
        "（1）从培养箱取出96孔板，在倒置显微镜下快速观察各孔细胞状态，确认无明显污染（培养基变黄、浑浊等）；",
        "（2）将CCK-8试剂从冰箱取出，在室温（25°C）下预温15分钟；",
        "（3）用多道移液器向内圈60孔（B2~G11）各孔加入10μL CCK-8溶液（注意：保持加样速度一致，全程在5分钟内完成）；",
        "（4）轻轻十字形晃动96孔板3次，使CCK-8与培养基混合均匀；",
        "（5）将96孔板置于37°C、5% CO₂培养箱孵育90分钟（1.5小时），期间避免光线直射；",
        "（6）孵育结束后，取出96孔板，在室温平衡5分钟；",
        "（7）用酶标仪读取各孔吸光度：主波长450nm，参比波长600nm；读取前检查孔底无气泡；",
        "（8）将读数数据导出至Excel，按实验记录模板记录（详见附录C）；",
        "（9）实验结束后，将96孔板弃入医疗废物桶，按生物安全规程处理废弃物；",
        "（10）在实验记录本上记录：实验日期、时间、操作者、仪器编号、CCK-8批次号等信息。",
    ]
    for s in cck8_steps_24h:
        doc.add_paragraph(s)

    doc.add_heading("3.4 Day 4（48h检测）：CCK-8检测", level=2)
    doc.add_paragraph(
        "Day 4的CCK-8检测操作步骤与Day 3（3.3节）完全相同，在给药后48小时进行。"
        "操作前确认：计算好精确的48h时间点，与Day 3检测时间保持一致（例如均在Day 2给药后的对应小时数进行）。"
        "数据记录要求：与24h数据一起记录于同一Excel工作表，明确区分24h和48h数据列。"
    )

    doc.add_heading("3.5 Day 5（72h检测）：CCK-8检测", level=2)
    doc.add_paragraph(
        "Day 5的CCK-8检测操作步骤与Day 3（3.3节）完全相同，在给药后72小时进行。"
        "特殊注意事项：72h时间点细胞密度显著增加，对照组细胞可能接近汇合；"
        "若对照组（溶剂对照组）细胞汇合度超过90%，需在实验记录中注明，"
        "并在讨论中说明高汇合度对CCK-8检测结果可能产生的影响（细胞接触抑制可能影响增殖抑制率计算）。"
    )


def add_chapter4(doc):
    """Add Chapter 4: Data Analysis."""
    doc.add_heading("第4章  数据分析与IC₅₀计算（GraphPad Prism 9.0全流程教程）", level=1)

    doc.add_heading("4.1 数据预处理刚性规范", level=2)
    doc.add_heading("4.1.1 原始OD值数据整理", level=3)

    doc.add_paragraph(
        "从酶标仪导出的原始数据（OD450nm/OD600nm）需按以下步骤进行预处理："
    )

    preprocess_steps = [
        "Step 1 - 双波长校正：OD校正值 = OD450nm - OD600nm（消除非特异性背景吸收，提高检测精度）；",
        "Step 2 - 空白对照扣除：将每孔的OD校正值减去空白对照组（不含细胞）的平均OD校正值；",
        "Step 3 - 复孔检验：同一药物浓度的5个复孔OD值计算变异系数（CV%）= 标准差/均值×100%，刚性要求CV% ≤ 15%；若某孔CV%超标，在原始数据中标注，经统计学检验（Grubbs离群值检验）后可剔除；",
        "Step 4 - 细胞增殖抑制率（Inhibition Rate，IR%）计算：",
    ]
    for s in preprocess_steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_paragraph(
        "增殖抑制率（IR%）计算公式：\n"
        "IR% = [1 - (OD处理组 - OD空白对照) / (OD溶剂对照 - OD空白对照)] × 100%\n\n"
        "其中：\n"
        "• OD处理组 = 药物处理孔的OD校正值（OD450 - OD600）\n"
        "• OD空白对照 = 不含细胞孔的平均OD校正值（背景）\n"
        "• OD溶剂对照 = 仅含DMSO（≤0.1%）的细胞孔的平均OD校正值（100%活细胞基准）\n\n"
        "计算结果应为0%~100%之间，若出现负值说明实验设计有误（药物促进增殖或测量误差）。"
    )

    doc.add_heading("4.2 GraphPad Prism 9.0全流程分步操作教程", level=2)
    doc.add_heading("4.2.1 软件安装与许可证确认", level=3)
    doc.add_paragraph(
        "GraphPad Prism 9.0（Windows版/macOS版）需使用正版许可证；课题组通常可通过学校软件授权平台获取。"
        "安装完成后，在软件标题栏确认版本号为「9.x.x」。"
    )

    doc.add_heading("4.2.2 新建项目和数据导入", level=3)
    prism_steps_1 = [
        "（1）打开GraphPad Prism 9.0，在「Welcome to GraphPad Prism」界面选择「New Project File」；",
        "（2）在「Type of graph」中选择「XY」（XY散点图类型，用于剂量-效应曲线拟合）；",
        "（3）在「X」设置项中选择「Numbers」，在「Y」设置项中选择「Enter and plot a single Y value for each point」（单Y值）或「Enter replicate values in side-by-side subcolumns」（多复孔值）；",
        "（4）点击「Create」创建数据表；",
        "（5）在数据表中输入：X列输入药物浓度的对数值（log₁₀[Concentration in μg/mL]），例如：log(400)=2.602, log(40)=1.602, log(4)=0.602, log(0.4)=-0.398, log(0.04)=-1.398, log(0.004)=-2.398；Y列输入对应的细胞增殖抑制率（%），若有多复孔，分别在Y1、Y2、Y3...列输入各复孔数值。",
    ]
    for s in prism_steps_1:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading("4.2.3 非线性回归拟合设置（核心步骤）", level=3)
    prism_steps_2 = [
        "（1）在数据表界面，点击工具栏「Analyze」→「Analyze Data」；",
        "（2）在左侧「Curves & Regression」选项下，选择「Nonlinear regression (curve fit)」，点击「OK」；",
        "（3）在弹出的「Parameters: Nonlinear Regression」对话框中，点击「Equation」选项卡；",
        "（4）在方程库中找到「Dose-response - Inhibition」类别，选择「log(inhibitor) vs. normalized response -- Variable slope (four parameters)」（四参数可变斜率方程，即标准4PL模型）；",
        "（5）在「Constrain」选项卡中，设置参数约束：Top = 100（最大响应固定为100%），Bottom = 0（最小响应固定为0%），HillSlope > 0（对于抑制曲线，HillSlope通常为正值）；",
        "（6）在「Output」选项卡中，确保勾选「EC50」（即IC₅₀）和「95% Confidence Interval」；",
        "（7）点击「OK」开始拟合运算，Prism将自动计算并显示拟合结果。",
    ]
    for s in prism_steps_2:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading("4.2.4 结果解读与IC₅₀报告规范", level=3)
    doc.add_paragraph(
        "拟合结果界面将显示以下关键参数，需逐一核对并记录："
    )
    results_items = [
        "IC₅₀值（EC50）及其95%置信区间（95% CI）：例如 IC₅₀ = 45.3 μg/mL（95% CI：38.2~53.8 μg/mL）；",
        "HillSlope（Hill系数/斜率因子）：绝对值越大表示剂量-效应曲线越陡；通常1~2之间为正常范围；",
        "Top和Bottom值：若已设置约束，应固定为100%和0%；若未约束，检查实际拟合值是否合理；",
        "R²（Goodness of Fit）：刚性要求 ≥ 0.95；若R² < 0.95，实验数据无效，需分析原因后重做实验；",
        "Sum of Squares（残差平方和）：越小表示拟合效果越好。",
    ]
    for item in results_items:
        doc.add_paragraph(item, style='List Bullet')

    # IC50 figure
    if FIG_IC50.exists():
        p_fig = doc.add_paragraph()
        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fig = p_fig.add_run()
        run_fig.add_picture(str(FIG_IC50), width=Inches(4.5))
        p_cap = doc.add_paragraph("图3 剂量-效应曲线（S型Sigmoid曲线）与IC₅₀计算示意图（GraphPad Prism 9.0）")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(9)
        p_cap.runs[0].italic = True

    doc.add_heading("4.3 数据统计分析刚性规范", level=2)
    doc.add_paragraph(
        "多次独立生物学重复实验（IBR ≥ 3次）完成后，须对IC₅₀数据进行统计学处理："
    )
    stats_items = [
        "IC₅₀汇总报告：将≥3次独立实验的IC₅₀值汇总，以「Mean ± SD（n=3）」格式表示，例如：IC₅₀(48h) = 45.3 ± 3.8 μg/mL（Mean ± SD，n=3）；",
        "正态性检验：使用Shapiro-Wilk检验（n<50时推荐），在GraphPad Prism中通过「Analyze」→「Column Statistics」→「Normality tests」执行；",
        "多组比较（24h vs. 48h vs. 72h）：若数据服从正态分布，使用单因素方差分析（One-way ANOVA）+ Tukey事后检验；若不服从正态分布，使用Kruskal-Wallis非参数检验；",
        "统计显著性标准：p < 0.05为显著差异，p < 0.01为极显著差异，p < 0.001为非常极显著差异，论文中以*/**/***表示；",
        "论文中IC₅₀的规范表述：「苏木醇提物对CNE-2细胞的IC₅₀值（48h）为XX μg/mL（95%CI：XX~XX μg/mL，n=3），拟合R²=0.XX」。",
    ]
    for item in stats_items:
        doc.add_paragraph(item, style='List Number')


def add_chapter5(doc):
    """Add Chapter 5: QC and Troubleshooting."""
    doc.add_heading("第5章  全流程质量控制与常见问题排查解决方案", level=1)

    doc.add_heading("5.1 全流程质控红线清单", level=2)
    doc.add_paragraph(
        "以下为本实验全流程的质控红线——任何一条被突破，均须立即停止实验，向导师报告，"
        "在原始记录中注明发现问题的时间点和处理决定。严禁在质控指标不达标的情况下继续实验，"
        "严禁对不达标的数据进行人为修改或选择性报告。"
    )

    # QC red lines table
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    headers = ["质控类别", "质控指标", "刚性标准", "违规处置"]
    qc_data = [
        ["细胞质控", "细胞活率（台盼蓝）", "铺板前活率 ≥ 95%", "暂停铺板，重新培养"],
        ["细胞质控", "支原体检测", "阴性", "废弃当前批次细胞，重新复苏"],
        ["药物质控", "DMSO终浓度", "≤ 0.1%（v/v）", "重新配制工作液"],
        ["铺板质控", "铺板密度", "8000 ± 800个细胞/孔（允许±10%偏差）", "重新铺板"],
        ["检测质控", "溶剂对照组CV%", "≤ 10%", "检查多道移液器，重复实验"],
        ["检测质控", "空白对照组OD值", "< 0.1（OD450）", "检查CCK-8试剂和操作"],
        ["数据质控", "剂量-效应曲线R²", "≥ 0.95", "实验数据无效，重做实验"],
        ["重复性质控", "独立生物学重复次数", "IBR ≥ 3次，每次CV% ≤ 15%", "增加实验次数"],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(qc_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("5.2 核心问题排查与解决方案", level=2)

    problems = [
        (
            "问题1：对照组细胞存活率异常（OD值过低，< 0.3）",
            "可能原因：（a）DMSO浓度超过1%导致细胞毒性；（b）细胞传代时间过长（对数生长期末期）；"
            "（c）培养箱CO₂浓度不稳定；（d）完全培养基配制错误（FBS未灭活或比例错误）。\n"
            "解决方案：检查DMSO浓度；检查培养箱参数；重新培养细胞至对数生长期后重铺板。"
        ),
        (
            "问题2：边缘孔OD值明显高于或低于内孔",
            "可能原因：PBS封边方案未执行，或封边孔PBS挥发；培养箱内温度不均匀。\n"
            "解决方案：严格执行PBS封边方案（外圈36孔加入200μL PBS）；"
            "将96孔板置于培养箱中央层，远离培养箱壁。"
        ),
        (
            "问题3：剂量-效应曲线不呈S型或R² < 0.95",
            "可能原因：（a）浓度梯度设置不合理（未覆盖IC₅₀浓度范围）；"
            "（b）药物溶解不充分（在最高浓度出现沉淀）；（c）数据点不足（少于5个有效浓度）。\n"
            "解决方案：根据预实验结果调整浓度范围（需覆盖0%~100%抑制率的完整范围）；"
            "检查药物溶解情况；增加浓度点（建议8~10个浓度梯度）。"
        ),
        (
            "问题4：CCK-8孵育后颜色变化不明显（OD差异小于0.1）",
            "可能原因：（a）CCK-8试剂失效（储存不当或超期）；"
            "（b）孵育时间不足（< 60分钟）；（c）细胞密度过低（< 3000个/孔）。\n"
            "解决方案：更换新批次CCK-8试剂并做空白对照验证；"
            "延长孵育时间至2小时；提高铺板密度（尝试10000~15000个/孔）。"
        ),
        (
            "问题5：多次重复实验结果差异过大（SD超过Mean的20%）",
            "可能原因：（a）细胞传代数差异（不同批次细胞生理状态不一致）；"
            "（b）药物母液在多次冻融后活性下降；（c）操作者个体操作差异（移液精度）。\n"
            "解决方案：固定使用同一传代批次的细胞（相同P数）进行3次重复实验；"
            "药物母液分装（小量，约10μL/管），避免反复冻融（≤3次）；"
            "由同一操作者完成所有重复实验。"
        ),
    ]
    for title, content in problems:
        p_title = doc.add_paragraph()
        r = p_title.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        doc.add_paragraph(content)

    doc.add_heading("5.3 实验有效性判定刚性标准", level=2)
    doc.add_paragraph(
        "实验完成后，须按以下标准判定本次实验是否有效。全部标准必须同时满足，"
        "方可将数据用于IC₅₀计算。否则，须在实验记录中注明无效原因，废弃本次数据，重新开展实验。"
    )

    validity_items = [
        "溶剂对照组（Vehicle Control）OD450值 ≥ 0.5（证明细胞活性充足，密度足够）；",
        "空白对照组（Blank Control）OD450值 < 0.1（背景噪音可忽略）；",
        "最高药物浓度（400 μg/mL）的抑制率 ≥ 70%（证明药物浓度范围覆盖了有效抑制范围）；",
        "最低药物浓度（0.004 μg/mL）的抑制率 ≤ 20%（证明药物浓度范围下限足够低）；",
        "GraphPad Prism 4PL非线性回归拟合的R² ≥ 0.95（证明剂量-效应关系良好）；",
        "溶剂对照组5个复孔的OD值CV% ≤ 10%（证明铺板均匀性和操作稳定性）；",
        "阳性对照药物（如顺铂）的抑制率在预期范围内（顺铂10μM对CNE-2细胞48h抑制率应 ≥ 60%）。",
    ]
    for item in validity_items:
        doc.add_paragraph(item, style='List Bullet')


def add_appendices(doc):
    """Add appendices."""
    doc.add_heading("附  录", level=1)

    doc.add_heading("附录A：96孔板铺板排布规范示意图", level=2)
    doc.add_paragraph(
        "96孔板（8行×12列，共96孔）铺板规范说明：\n"
        "• 外圈36孔（行1、行8、列1、列12）：无菌PBS（200 μL/孔），用于防控边缘效应\n"
        "• 内圈60孔（B2~G11）：实验孔，分配方案如下："
    )

    # Detailed plate layout
    plate_assignments = [
        ("B2:B6", "苏木浓度C1（400 μg/mL）× 5复孔", "24h检测板"),
        ("B7:B11", "苏木浓度C2（40 μg/mL）× 5复孔", "24h检测板"),
        ("C2:C6", "苏木浓度C3（4 μg/mL）× 5复孔", "24h检测板"),
        ("C7:C11", "苏木浓度C4（0.4 μg/mL）× 5复孔", "24h检测板"),
        ("D2:D6", "苏木浓度C5（0.04 μg/mL）× 5复孔", "24h检测板"),
        ("D7:D11", "苏木浓度C6（0.004 μg/mL）× 5复孔", "24h检测板"),
        ("E2:E11", "溶剂对照组（Vehicle Control，≤0.1% DMSO）× 10复孔", "24h检测板"),
        ("F2:F11", "空白对照组（Blank Control，无细胞）× 10复孔", "24h检测板"),
        ("G2:G6", "阳性对照组（顺铂10μM）× 5复孔", "24h检测板"),
        ("G7:G11", "备用/重复孔", "24h检测板"),
    ]

    table = doc.add_table(rows=len(plate_assignments) + 1, cols=3)
    table.style = 'Table Grid'
    headers_app = ["孔位范围", "内容（示例）", "备注"]
    for j, h in enumerate(headers_app):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(plate_assignments):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_heading("附录B：药物梯度稀释计算示例表", level=2)
    doc.add_paragraph(
        "以苏木醇提物（母液浓度20 mg/mL，DMSO溶解）为例，配制6个浓度梯度工作液（100μL/孔，总体积200μL/管）："
    )

    dilution_table = doc.add_table(rows=7, cols=6)
    dilution_table.style = 'Table Grid'
    dil_headers = ["梯度", "目标浓度(μg/mL)", "母液/上管(μL)", "培养基(μL)", "总体积(μL)", "最终DMSO%"]
    dil_data = [
        ["C1", "400", "20 (母液)", "980", "1000", "0.10%"],
        ["C2", "40", "100 (C1)", "900", "1000", "0.01%"],
        ["C3", "4", "100 (C2)", "900", "1000", "0.001%"],
        ["C4", "0.4", "100 (C3)", "900", "1000", "0.0001%"],
        ["C5", "0.04", "100 (C4)", "900", "1000", "0.00001%"],
        ["C6", "0.004", "100 (C5)", "900", "1000", "0.000001%"],
    ]
    for j, h in enumerate(dil_headers):
        dilution_table.rows[0].cells[j].text = h
        dilution_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(dil_data):
        for j, cell_text in enumerate(row_data):
            dilution_table.rows[i+1].cells[j].text = cell_text
    doc.add_paragraph()

    doc.add_paragraph(
        "注意：上表中DMSO终浓度以最高浓度C1为最高（0.10%），在溶剂对照组中须匹配相同DMSO浓度（0.10%）。"
    )

    doc.add_heading("附录C：实验原始记录Word模板（字段清单）", level=2)
    doc.add_paragraph("实验记录必须包含以下字段（手写或电子记录均可，但须有操作者签名和日期）：")

    record_fields = [
        "实验日期：____年____月____日（Day 1铺板日期）",
        "操作者姓名：____________",
        "细胞信息：细胞系_CNE-2_，传代次数P___，活率____%，支原体检测结果____",
        "药物信息：样品名称____，批次号____，DMSO母液浓度____ mg/mL，配制日期____",
        "96孔板信息：产品批次号____，有效期____",
        "CCK-8试剂信息：品牌____，货号____，批次号____，有效期____",
        "酶标仪信息：仪器型号____，仪器编号____，校零日期____",
        "Day 1铺板时间：________，铺板密度：8000个/孔",
        "Day 2给药时间：________，给药浓度范围：0.004~400 μg/mL",
        "Day 3（24h）CCK-8加样时间：________，孵育结束读数时间：________",
        "Day 4（48h）CCK-8加样时间：________，孵育结束读数时间：________",
        "Day 5（72h）CCK-8加样时间：________，孵育结束读数时间：________",
        "数据文件路径/命名：____________________",
        "异常情况记录：____________________",
        "导师审阅签名：____________________，日期：____________",
    ]
    for f in record_fields:
        doc.add_paragraph(f, style='List Bullet')


def add_references(doc):
    """Add references section."""
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph(
        "（参考文献按GB/T 7714-2015格式，顺序编码制）"
    )

    refs = [
        "[1] 国家药典委员会. 中华人民共和国药典（2020年版）一部[M]. 北京：中国医药科技出版社，2020：101-102.",
        "[2] 李时珍. 本草纲目[M]. 北京：人民卫生出版社，1982：1089-1093.",
        "[3] Xie C X, Kokubun T, Houghton P J, et al. Antibacterial activity of the Chinese traditional medicine, Zi Hua Di Ding[J]. Phytotherapy Research, 2004, 18(6): 497-500.",
        "[4] Namikoshi M, Nakata H, Yamada H, et al. Homoisoflavonoids and related compounds. II. Isolation and absolute configurations of 3,4-dihydroxylated homoisoflavans and brazilins from Caesalpinia sappan L.[J]. Chemical & Pharmaceutical Bulletin, 1987, 35(7): 2761-2773.",
        "[5] Suyatmi Y, Yustika A, Nugroho W S, et al. Anticancer activity of brazilin from Caesalpinia sappan on A549 and HeLa cells[J]. Asian Pacific Journal of Cancer Prevention, 2020, 21(6): 1601-1607.",
        "[6] Lee Y, Choi E, Kim E, et al. Brazilin Induces Apoptosis and Cell Cycle Arrest in Human Brain Tumor Cells[J]. Natural Product Sciences, 2019, 25(1): 45-52.",
        "[7] Wudtiwai B, Sripanidkulchai B, Kongtawelert P, et al. Methoxyflavone derivatives modulate the expression of PD-L1 in cancer cells[J]. Cancer Immunology, Immunotherapy, 2021, 70(6): 1649-1664.",
        "[8] Widodo N, Sulistyoningrum A S, Primaharinastiti R, et al. Transcriptomic analysis of A549 cells treated with Caesalpinia sappan extract reveals disruption in mitochondrial ATP synthesis[J]. Molecular and Cellular Biochemistry, 2022, 477(3): 723-736.",
        "[9] Hanifa M A, Fauzia S, Saefudin. Network pharmacology and molecular docking analysis of brazilin from Caesalpinia sappan against hepatocellular carcinoma[J]. Journal of Herbal Medicine, 2022, 35: 100584.",
        "[10] Wei K R, Zheng R S, Zhang S W, et al. Nasopharyngeal carcinoma incidence and mortality in China, 2013[J]. Chinese Journal of Cancer, 2017, 36(1): 90.",
        "[11] Chan A T C, Grégoire V, Lefebvre J L, et al. Nasopharyngeal cancer: EHNS-ESMO-ESTRO Clinical Practice Guidelines for diagnosis, treatment and follow-up[J]. Annals of Oncology, 2012, 23(Suppl 7): vii83-vii85.",
        "[12] Liu Y, Lian Z. Network pharmacology-based analysis of the mechanisms of Caesalpinia sappan in treating nasopharyngeal carcinoma[J]. Evidence-Based Complementary and Alternative Medicine, 2021, 2021: 9958134.",
        "[13] Pan H, Tian M, Zhao Z, et al. Effects of Chinese herbal medicine on survival of nasopharyngeal carcinoma patients with radiation and concurrent chemotherapy[J]. Integrative Cancer Therapies, 2020, 19: [REDACTED-PHONE]02625.",
        "[14] El-Nashar H A S, Eldahshan O A, Attia E Z, et al. Inhibitory potential of Caesalpinia sappan heartwood constituents against 4T1 breast cancer cells: in vitro cytotoxicity and molecular docking studies[J]. Natural Product Research, 2022, 36(10): 2562-2567.",
        "[15] Wang S, Zhang Q, Ye Y, et al. Antitumor effects of traditional Chinese medicinal fungal formula on 4T1 mouse mammary tumor cells[J]. Journal of Ethnopharmacology, 2021, 265: 113229.",
        "[16] Franken N A P, Rodermond H M, Stap J, et al. Clonogenic assay of cells in vitro[J]. Nature Protocols, 2006, 1(5): 2315-2319.",
        "[17] Adan A, Kiraz Y, Baran Y. Cell proliferation and cytotoxicity assays[J]. Current Pharmaceutical Biotechnology, 2016, 17(14): 1213-1221.",
        "[18] Dojindo Molecular Technologies. CCK-8 Cell Counting Kit-8 Technical Manual[EB/OL]. (2022-01-01)[2026-03-08]. https://www.dojindo.com.",
        "[19] GraphPad Software. GraphPad Prism 9 Statistics Guide[EB/OL]. (2021-01-01)[2026-03-08]. https://www.graphpad.com.",
        "[20] 国家药品监督管理局. 抗肿瘤药物药效学研究技术指导原则[S]. 北京：国家药品监督管理局，2020.",
        "[21] ATCC. CNE-2 ATCC® CRL-5982™ Human Nasopharynx carcinoma[EB/OL]. (2023)[2026-03-08]. https://www.atcc.org.",
        "[22] Ishiyama M, Tominaga H, Shiga M, et al. A combined assay of cell viability and in vitro cytotoxicity with a highly water-soluble tetrazolium salt, neutral red and crystal violet[J]. Biological and Pharmaceutical Bulletin, 1996, 19(11): 1518-1520.",
        "[23] Skehan P, Storeng R, Scudiero D, et al. New colorimetric cytotoxicity assay for anticancer drug screening[J]. Journal of the National Cancer Institute, 1990, 82(13): 1107-1112.",
        "[24] Cory A H, Owen T C, Barltrop J A, et al. Use of an aqueous soluble tetrazolium/formazan assay for cell growth assays in culture[J]. Cancer Communications, 1991, 3(7): 207-212.",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)


def create_v3_docx(output_path, include_images=True):
    """Create the full v3.0 rescue Word document."""
    doc = Document()
    set_page_margins(doc)

    # Set default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Set heading styles
    for level in [1, 2, 3, 4]:
        try:
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = 'Arial'
            h_style.font.color.rgb = RGBColor(0x1F, 0x35, 0x64) if level == 1 else RGBColor(0x2E, 0x4A, 0x6E)
            h_style.font.size = Pt([16, 14, 12, 11][level-1])
        except:
            pass

    print(f"[INFO] Creating cover page...")
    add_cover_page(doc, include_image=include_images)

    print(f"[INFO] Adding table of contents...")
    add_toc(doc)

    print(f"[INFO] Writing Chapter 1 (Background)...")
    add_chapter1(doc)
    add_chapter1_continued(doc)

    print(f"[INFO] Writing Chapter 2 (Materials)...")
    add_chapter2(doc)
    add_chapter2_continued(doc)

    print(f"[INFO] Writing Chapter 3 (Procedures)...")
    add_chapter3(doc)

    print(f"[INFO] Writing Chapter 4 (Data Analysis)...")
    add_chapter4(doc)

    print(f"[INFO] Writing Chapter 5 (QC & Troubleshooting)...")
    add_chapter5(doc)

    print(f"[INFO] Adding appendices...")
    add_appendices(doc)

    print(f"[INFO] Adding references...")
    add_references(doc)

    print(f"[INFO] Saving document to: {output_path}")
    doc.save(str(output_path))
    size_mb = os.path.getsize(output_path) / 1024 / 1024
    print(f"[SUCCESS] Document saved: {output_path} ({size_mb:.2f} MB)")
    return output_path


def create_textonly_docx(output_path):
    """Create a simplified text-only version."""
    print(f"\n[INFO] Creating text-only version...")
    # Text-only version: no images, simplified formatting
    return create_v3_docx(output_path, include_images=False)


if __name__ == "__main__":
    # Create output directory if needed
    FINAL_DIR.mkdir(parents=True, exist_ok=True)

    # Create v3.0 full document
    v3_path = FINAL_DIR / "苏木_CNE2_IC50_SOP_v3.0_Rescue.docx"
    create_v3_docx(v3_path, include_images=True)

    # Create text-only version
    textonly_path = FINAL_DIR / "苏木_CNE2_IC50_SOP_TextOnly.docx"
    create_textonly_docx(textonly_path)

    print("\n" + "="*60)
    print("RESCUE PACKAGE GENERATION COMPLETE")
    print("="*60)
    print(f"v3.0 Full:   {v3_path}")
    print(f"Text-Only:   {textonly_path}")
    print(f"Existing PDF: {FINAL_DIR / 'SuMu_CNE2_IC50_Manual.pdf'}")
