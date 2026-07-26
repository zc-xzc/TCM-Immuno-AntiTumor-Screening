#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Continue: Add Chapter 3 (experimental procedures Day 1-5)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_warning_box(doc, text, color="FFF9E6"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_red_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("【刚性要求】 ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(10.5)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    set_para_spacing(p, before=3, after=3)
    return p

def add_numbered_list(doc, items, indent_cm=0.63):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(indent_cm)
        if isinstance(item, tuple):
            run = p.add_run(item[0])
            run.bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)
        set_para_spacing(p, before=1, after=1)

def add_bullet_list(doc, items, indent_cm=0.63):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent_cm)
        if isinstance(item, tuple):
            run = p.add_run(item[0])
            run.bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)
        set_para_spacing(p, before=1, after=1)

def add_three_line_table(doc, headers, rows, col_widths=None):
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(hdr_cells[i], 'D9E2F3')
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    return table

def add_step_header(doc, time_range, action):
    """Add a formatted step header with time range."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    pPr.append(shd)
    r1 = p.add_run(f'⏱ {time_range}  |  ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x39, 0x7D)
    r2 = p.add_run(action)
    r2.bold = True
    r2.font.size = Pt(10.5)
    set_para_spacing(p, before=6, after=2)
    return p

# Load the document
doc = Document('/app/sandbox/session_20260308_142335_c3580fd63071/writing_outputs/final/苏木对CNE-2细胞增殖抑制率及IC50测定全流程操作手册_v2.0.docx')

# ══════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════
doc.add_heading('第3章  实验操作步骤（Time-line格式）', level=1)

p_intro = doc.add_paragraph(
    '本章节以时间线（Time-line）格式详细描述苏木对CNE-2细胞增殖抑制实验的全流程操作步骤。'
    '实验跨度共5天（Day 1至Day 5），每个操作步骤精确到分钟级时间节点、微升级加液量、'
    '以及具体操作规范。所有操作步骤必须在BSL-2级生物安全柜内完成，严格执行无菌操作规程。'
)
set_para_spacing(p_intro, before=4, after=4)

# Overview timeline table
ot = add_three_line_table(doc,
    ['实验日期', '主要任务', '预计耗时', '关键质控节点'],
    [
        ['Day 1（D1）', '细胞铺板（3块96孔板）', '约3小时', '铺板后12h显微镜观察细胞贴壁情况'],
        ['Day 2（D2）', '苏木工作液配制+给药干预', '约2小时', '给药前镜下确认细胞状态；DMSO终浓度≤0.1%'],
        ['Day 3（D3）', '24h时间点CCK-8检测', '约2.5小时', '检测结束立即导出数据并备份'],
        ['Day 4（D4）', '48h时间点CCK-8检测', '约2.5小时', '同Day 3'],
        ['Day 5（D5）', '72h时间点CCK-8检测+数据处理', '约4小时', '同Day 3；数据完整归档'],
    ],
    col_widths=[2.5, 5, 2.5, 5]
)

# ── DAY 1 ──
doc.add_heading('3.1  Day 1：细胞铺板（核心防控边缘效应）', level=2)

p = doc.add_paragraph(
    'Day 1核心任务为CNE-2细胞的消化收集和96孔板铺板。本步骤的关键质控点为：'
    '（1）铺板前细胞活率验证（≥95%）；（2）96孔板边缘效应防控（PBS封边）；'
    '（3）细胞密度精准控制（每孔8,000个细胞）。'
)
set_para_spacing(p, before=4, after=4)

doc.add_heading('3.1.1  铺板前细胞质控SOP', level=3)

qc_steps = [
    ('汇合度要求：', 'Day 1铺板前，CNE-2细胞培养瓶（T-25或T-75）中的细胞汇合度必须达到80%-90%。'
     '使用倒置显微镜（10×物镜）观察细胞形态，细胞应呈典型的上皮样贴壁形态，正常生长，'
     '无漂浮死细胞聚集，无培养基颜色异常（正常RPMI-1640含酚红为红色/橙红色，若变黄则pH过低，提示污染或需换液）；'),
    ('细胞活率验证（台盼蓝染色法，Trypan Blue Exclusion Assay）：',
     '取50μL细胞悬液（消化收集后，见步骤3.1.2），与50μL 0.4%台盼蓝染色液混合（比例1:1），'
     '轻轻混匀后静置1-2分钟。将混合液滴入血球计数板，在倒置显微镜下计数：'
     '活细胞（折光性强，透明）和死细胞（被台盼蓝染成蓝色）。'
     '计算活率 = 活细胞数 ÷ （活细胞数 + 死细胞数） × 100%。'
     '活率≥95%方可进行铺板；活率＜95%的细胞需更换培养基，继续培养至活率达标，或重新复苏冻存细胞；'),
    ('【新手示例】细胞计数：', '在血球计数板4个大格（每个大格含16个小格）中分别计数，假设计数结果为：'
     '大格1：活细胞42个、死细胞2个；大格2：活细胞38个、死细胞1个；大格3：活细胞40个、死细胞2个；大格4：活细胞39个、死细胞1个。'
     '总活细胞数=42+38+40+39=159；总死细胞数=2+1+2+1=6。'
     '活率=159/(159+6)×100%=96.4%≥95%，合格，可进行铺板。'
     '细胞浓度（个/mL）= (159/4) × 10^4 × 2（稀释倍数）= 7.95 × 10^5/mL。'),
]
add_numbered_list(doc, qc_steps)

doc.add_heading('3.1.2  96孔板铺板刚性排布规范（PBS封边方案）', level=3)

add_red_warning(doc,
    '96孔板最外围36孔（第1列A1-H1、第12列A12-H12、第A行A1-A12、第H行H1-H12，共36孔）'
    '全部加入200μL无菌PBS（磷酸缓冲盐溶液，pH 7.4）进行封边处理，不种细胞、不加药物。'
    '仅使用中间区域（B2至G11，共60孔）用于实验，以彻底消除96孔板边缘效应（Edge Effect）[22]。'
    '违反本规定的实验数据无效！')

# 96-well plate diagram using table
p_diag = doc.add_paragraph()
p_diag.add_run('96孔板铺板排布示意图（文字版）：').bold = True
set_para_spacing(p_diag, before=6, after=2)

# Create plate diagram
plate_rows = ['A','B','C','D','E','F','G','H']
plate_cols = [str(i) for i in range(1,13)]
plate_table = doc.add_table(rows=9, cols=13)
plate_table.style = 'Table Grid'
plate_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row (column numbers)
plate_table.rows[0].cells[0].text = ''
for j, col in enumerate(plate_cols):
    c = plate_table.rows[0].cells[j+1]
    c.text = col
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].runs[0].font.size = Pt(9)
    c.paragraphs[0].runs[0].bold = True

for i, row_letter in enumerate(plate_rows):
    plate_table.rows[i+1].cells[0].text = row_letter
    plate_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
    plate_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    for j in range(12):
        c = plate_table.rows[i+1].cells[j+1]
        is_edge = (i == 0 or i == 7 or j == 0 or j == 11)
        if is_edge:
            c.text = 'PBS'
            c.paragraphs[0].runs[0].font.size = Pt(7)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(c, 'BDD7EE')  # light blue for PBS
        else:
            c.text = '●'
            c.paragraphs[0].runs[0].font.size = Pt(8)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(c, 'E2EFDA')  # light green for cells
        for col_obj in plate_table.columns:
            col_obj.width = Cm(1.2)

p_legend = doc.add_paragraph()
r1 = p_legend.add_run('■ 蓝色PBS封边孔（36孔）：')
r1.font.color.rgb = RGBColor(0x1F, 0x7A, 0xC0)
p_legend.add_run('加200μL无菌PBS，不种细胞，不加药；')
r2 = p_legend.add_run('  ■ 绿色实验孔（60孔）：')
r2.font.color.rgb = RGBColor(0x37, 0x8A, 0x38)
p_legend.add_run('B2-G11区域，用于分组实验（空白组/对照组/给药组）')
p_legend.add_run().font.size = Pt(9)
set_para_spacing(p_legend, before=2, after=6)

p_density = doc.add_paragraph()
p_density.add_run('细胞接种密度刚性规范：').bold = True
set_para_spacing(p_density, before=4, after=2)

density_items = [
    '每孔细胞接种数量：8,000个CNE-2细胞/孔（参考范围：6,000-10,000个/孔，根据预实验优化）；',
    '细胞悬液浓度：调整为8×10⁴个/mL（即每100μL含8,000个细胞）；',
    '每孔接种体积：100μL细胞悬液（含8,000个CNE-2细胞）；',
    '计算所需细胞总数（以3块96孔板为例）：每块板使用60孔，其中对照组8孔，空白组5孔，给药组47孔。'
    '3块板共需细胞=60×3×8,000=1,440,000个细胞（约144万个）。建议准备约180万个细胞备用（冗余20%）。',
]
add_numbered_list(doc, density_items)

doc.add_heading('3.1.3  铺板分步操作SOP（精确到分钟级）', level=3)

p_pre = doc.add_paragraph()
p_pre.add_run('铺板前准备（D1当天早上，铺板操作开始前60分钟）：').bold = True
set_para_spacing(p_pre, before=6, after=2)

pre_items = [
    '提前60min：打开生物安全柜电源，开启送风，用70%乙醇（75%乙醇效果更佳）全面擦拭超净台内壁、托盘表面；',
    '提前45min：开启紫外灯照射超净台内部30分钟；',
    '提前30min：将所需试剂（PBS、胰蛋白酶-EDTA、完全培养基、台盼蓝）从冰箱取出，放入超净台内平衡至室温（或37℃水浴预热培养基）；',
    '提前15min：关闭紫外灯，开启送风，所有操作人员穿戴好实验服、手套（双层）、口罩；',
    '正式操作前：再次用70%乙醇喷洒手套，在超净台中稳定操作。',
]
add_numbered_list(doc, pre_items)

add_step_header(doc, '0-15 min', '准备阶段：超净台消毒 + 试剂平衡')
step0 = [
    '确认生物安全柜已预热（送风运行≥15分钟），超净台内壁无明显乙醇残留；',
    '将完全培养基（37℃预热）、PBS（室温）、0.25%胰蛋白酶-EDTA（室温）、台盼蓝、96孔板（3块）放入超净台内；',
    '从CO₂培养箱取出CNE-2细胞培养瓶（汇合度80%-90%），显微镜观察细胞形态：贴壁良好、形态正常、无漂浮死细胞。',
]
add_numbered_list(doc, step0)

add_step_header(doc, '15-30 min', '细胞消化收集（胰蛋白酶消化法）')
step1 = [
    '（第1步）弃培养基：倾斜培养瓶，用无菌吸管吸去培养瓶中的旧培养基，尽量去除干净；',
    '（第2步）PBS洗涤：向培养瓶中加入5mL（T-25瓶）或10mL（T-75瓶）无菌PBS，轻轻摇晃润洗细胞表面，吸去PBS。重复洗涤2次，彻底洗去残余血清（血清会抑制胰蛋白酶活性）；',
    '（第3步）胰蛋白酶消化：向培养瓶中加入1mL（T-25）或2mL（T-75）0.25%胰蛋白酶-EDTA，轻轻摇晃使胰蛋白酶均匀覆盖细胞层；',
    '（第4步）消化观察：将培养瓶放回37℃培养箱中消化2-4分钟（CNE-2消化时间参考2-3分钟）。每1分钟取出用倒置显微镜观察：当细胞变圆、细胞间隙增大、轻拍培养瓶有少量细胞脱落时，立即终止消化；',
    '（第5步）终止消化：向培养瓶中加入5mL（T-25）或10mL（T-75）含10%FBS的完全培养基（FBS中和胰蛋白酶），用移液管反复吹打细胞层，使细胞充分脱离瓶壁形成单细胞悬液；',
    '（第6步）离心收集：将细胞悬液转移至15mL离心管，以1000rpm（约200×g）、室温离心5分钟。离心完成后，弃去上清，保留细胞沉淀（pellet）；',
    '（第7步）重悬：向细胞沉淀中加入2mL完全培养基，用移液枪轻柔吹打重悬，打散细胞团块，形成均匀的单细胞悬液。',
]
add_numbered_list(doc, step1)

add_warning_box(doc,
    '【避坑提示1——消化时间】 胰蛋白酶消化时间过长会损伤细胞，过短细胞脱落不完全。'
    'CNE-2的消化时间通常为2-3分钟，具体时间因批次和细胞代次而异。'
    '建议每隔1分钟观察一次。绝对禁止消化超过5分钟。',
    color='FFF3CD')

add_step_header(doc, '30-45 min', '细胞计数与密度调整')
step2 = [
    '取50μL细胞悬液加入50μL 0.4%台盼蓝，混匀后滴入血球计数板，按3.1.1节方法计数；',
    '记录活细胞数、死细胞数、活率（要求≥95%）、细胞浓度（个/mL）；',
    '计算稀释比例：目标浓度=8×10⁴个/mL。例如计数得细胞浓度=5.2×10⁵个/mL，'
    '则稀释倍数=5.2×10⁵÷8×10⁴=6.5倍。取1mL细胞悬液加入5.5mL完全培养基，'
    '混匀后即为8×10⁴个/mL的细胞接种悬液；',
    '稀释后再次取样计数（选做，建议新手执行），确认稀释后浓度在(7-9)×10⁴个/mL范围内，误差≤10%。',
]
add_numbered_list(doc, step2)

add_step_header(doc, '45-60 min', '96孔板铺板操作')
step3 = [
    '（第1步）PBS封边：用1000μL移液枪，向第1列（A1-H1）、第12列（A12-H12）、第A行（A2-A11）、第H行（H2-H11）共36孔，每孔加入200μL无菌PBS。操作时保持移液枪垂直，沿孔壁缓慢加入，避免产生气泡；',
    '（第2步）加细胞悬液：用200μL移液枪，向中间实验区（B2至G11，共60孔）每孔加入100μL细胞悬液（含8,000个CNE-2细胞）。建议使用多道移液枪提高效率，按列顺序依次加样（B2→G2，B3→G3，...，B11→G11）；',
    '（第3步）混匀：加样完毕后，将96孔板水平放置，用双手拇指和食指轻轻夹持孔板两侧，进行"水平十字晃动"（前后左右各轻摇5次，幅度约2cm），使细胞在孔中均匀分布，避免细胞聚集于孔中央；',
    '（第4步）放入培养箱：将混匀的96孔板放入37℃、5%CO₂培养箱内，水平放置，静置培养18-24小时，让细胞充分贴壁；',
    '（第5步）重复以上步骤完成3块96孔板的铺板（分别对应24h、48h、72h检测时间点）。每块孔板需单独标注检测时间点（如"24h板"、"48h板"、"72h板"）。',
]
add_numbered_list(doc, step3)

add_warning_box(doc,
    '【避坑提示2——气泡问题】 加样时若产生气泡，应立即用细针头或吸头轻轻挑破气泡，'
    '或用酒精灯火焰快速掠过孔板（不建议用于96孔板，仅作应急）。'
    '有气泡的孔CCK-8检测时OD值会异常偏低，必须处理后再检测。',
    color='FFF3CD')

doc.add_heading('3.1.4  铺板后质控检查', level=3)
qc_post = [
    '铺板后12小时（通常在铺板当天晚上或次日早晨）：从培养箱取出96孔板，在倒置显微镜下观察B2至G11所有实验孔中的细胞状态；',
    '合格标准：细胞贴壁良好，分布较均匀（允许轻微不均），细胞密度约占孔底面积15-25%（汇合度），无漂浮死细胞聚集，无肉眼可见污染（培养基颜色正常）；',
    '给药前（Day 2上午）再次观察：细胞汇合度应达到30%-40%。汇合度过低（＜20%）提示铺板密度不足，需调整下次实验的铺板密度；汇合度过高（＞60%）提示铺板密度过大，药物处理时细胞可能已进入接触抑制期，影响实验结果；',
    '若观察到任意孔有污染迹象（培养基浑浊、颜色变黄、有颗粒悬浮），该孔数据必须废弃，污染孔不得纳入分析。',
]
add_numbered_list(doc, qc_post)

# ── DAY 2 ──
doc.add_heading('3.2  Day 2：药物配制与给药干预', level=2)

p = doc.add_paragraph(
    'Day 2核心任务为：（1）苏木工作液的梯度稀释配制；（2）给药干预操作。'
    '本步骤的最关键质控点为DMSO终浓度控制（≤0.1%）和给药操作的无菌性保证。'
)
set_para_spacing(p, before=4, after=4)

doc.add_heading('3.2.1  实验分组设置（刚性规定，无任何例外）', level=3)

group_table = add_three_line_table(doc,
    ['分组名称', '处理内容', '复孔数量', '功能说明'],
    [
        ['空白组（Blank）', '无细胞，加100μL含0.04%DMSO的完全培养基', '≥5孔', '用于扣除培养基与CCK-8试剂背景OD值'],
        ['对照组（Vehicle Control）', 'CNE-2细胞+含0.04%DMSO的完全培养基（无药物）', '≥8孔', '用于计算细胞存活率（分母），定义为100%存活'],
        ['给药组C1（400μg/mL）', 'CNE-2细胞+400μg/mL苏木工作液', '≥5孔/浓度', '最高浓度'],
        ['给药组C2（40μg/mL）', 'CNE-2细胞+40μg/mL苏木工作液', '≥5孔/浓度', '10倍梯度稀释'],
        ['给药组C3（4μg/mL）', 'CNE-2细胞+4μg/mL苏木工作液', '≥5孔/浓度', '10倍梯度稀释'],
        ['给药组C4（0.4μg/mL）', 'CNE-2细胞+0.4μg/mL苏木工作液', '≥5孔/浓度', '10倍梯度稀释'],
        ['给药组C5（0.04μg/mL）', 'CNE-2细胞+0.04μg/mL苏木工作液', '≥5孔/浓度', '10倍梯度稀释'],
        ['给药组C6（0.004μg/mL）', 'CNE-2细胞+0.004μg/mL苏木工作液', '≥5孔/浓度', '最低浓度'],
    ],
    col_widths=[3.5, 5, 2.5, 4]
)

p_note = doc.add_paragraph(
    '注：以上分组方案假设使用整版孔板（B2-G11，60孔）。实际孔位分配方案：'
    '空白组=B2-B6（5孔）；对照组=B7-C6（8孔）；C1-C6各5孔（共30孔）；剩余空孔保留备用。'
    '每个时间点（24h/48h/72h）独立使用一块96孔板，避免反复开盖导致污染。'
)
p_note.runs[0].font.size = Pt(9.5)
set_para_spacing(p_note, before=2, after=6)

doc.add_heading('3.2.2  药物工作液配制分步SOP（避光全程操作）', level=3)

p_imp = doc.add_paragraph()
p_imp.add_run('重要提示：').bold = True
r = p_imp.add_run('苏木提取物/巴西苏木素对光敏感，配制工作液的全程操作必须在避光条件下进行（超净台内关灯，用铝箔包裹离心管）。工作液现配现用，不得提前配制后储存。')
r.font.size = Pt(10.5)
set_para_spacing(p_imp, before=4, after=4)

add_step_header(doc, '第一步', '配制1mg/mL中间稀释液（在超净台内操作）')
inter_steps = [
    '从-20℃冰箱取出100mg/mL苏木DMSO母液（分装管，每管10μL），置于冰浴中缓慢解冻；',
    '用铝箔包裹15mL离心管，加入990μL完全培养基；',
    '向上述培养基中加入10μL的100mg/mL母液，用200μL移液枪上下吹打20次充分混匀；',
    '此时即得1mg/mL（1000μg/mL）的含1%DMSO中间液。中间液在冰浴中暂存，30分钟内使用完毕。',
]
add_numbered_list(doc, inter_steps)

add_step_header(doc, '第二步', '10倍梯度稀释配制6个浓度工作液')
dilution_steps = [
    '（C1=400μg/mL）：取6支1.5mL无菌EP管，标记C1-C6，各加入900μL完全培养基。向C1管中加入400μL 1mg/mL中间液，混匀（最终体积1300μL，浓度=1000×400/1300≈308μg/mL）——注意，此为2倍工作液，加入等体积培养基后终浓度为400μg/mL？'
    '\n    【正确方案】为直接在含培养基的体系中稀释至目标浓度：\n'
    '    C1工作液（400μg/mL）：取400μL 1mg/mL中间液+600μL完全培养基=1000μL，DMSO终浓度0.04%；',
    '（C2=40μg/mL）：取100μL C1（400μg/mL）+900μL完全培养基=1000μL；',
    '（C3=4μg/mL）：取100μL C2（40μg/mL）+900μL完全培养基=1000μL；',
    '（C4=0.4μg/mL）：取100μL C3（4μg/mL）+900μL完全培养基=1000μL；',
    '（C5=0.04μg/mL）：取100μL C4（0.4μg/mL）+900μL完全培养基=1000μL；',
    '（C6=0.004μg/mL）：取100μL C5（0.04μg/mL）+900μL完全培养基=1000μL；',
    '对照组培养基配制：在1.5mL EP管中，取400μL 1mg/mL DMSO/培养基稀释液+600μL培养基=1000μL含0.04%DMSO的对照培养基（与C1 DMSO浓度一致）。',
]
add_numbered_list(doc, dilution_steps)

add_warning_box(doc,
    '【DMSO浓度验证】 C1工作液（400μg/mL）中DMSO浓度 = '
    '(400μL × 1%) / 1000μL = 0.4%/1000 × 400 = 0.4%。\n'
    '等体积加入细胞孔（弃去旧培养基后加入100μL工作液），孔内DMSO终浓度 = 0.04%，满足≤0.1%要求。\n'
    '对照组同样加入含0.04%DMSO的培养基，溶剂浓度完全一致。',
    color='E2EFDA')

doc.add_heading('3.2.3  给药操作分步SOP', level=3)

add_step_header(doc, '给药前（Day 2上午）', '细胞状态确认')
pre_drug = [
    '从培养箱取出3块96孔板，放置于倒置显微镜下依次观察B2-G11区域所有孔的细胞状态；',
    '确认：①细胞贴壁良好，汇合度约30%-40%（正常状态，可以给药）；②无肉眼可见污染；③孔间细胞密度无明显差异（允许±10%以内的变异）；',
    '任何孔有污染迹象：在孔板示意图上标记该孔，后续数据处理时剔除该孔数据；',
    '若整板超过10孔有污染：该板报废，重新铺板，实验延期至下一批次CNE-2细胞状态良好时进行。',
]
add_numbered_list(doc, pre_drug)

add_step_header(doc, '给药操作（Day 2下午）', '弃培养基 + 加工作液')
drug_steps = [
    '在生物安全柜内，使用无菌吸管或移液枪，从96孔板实验孔（B2-G11）中小心吸去旧培养基（100μL），注意避免碰触孔底细胞层；',
    '按分组方案（见3.2.1），依次向各孔加入对应的工作液/培养基：',
    '   → 空白孔（无细胞孔）：加入100μL含0.04%DMSO的完全培养基；',
    '   → 对照孔：加入100μL含0.04%DMSO的完全培养基（无药物）；',
    '   → 给药孔C1-C6：加入100μL对应浓度的苏木工作液；',
    '加样顺序建议：从低浓度（C6）到高浓度（C1）依次加样，降低高浓度样品对吸头的污染风险；',
    '加样完毕后，将96孔板水平十字轻摇混匀（前后左右各5次，幅度约2cm）；',
    '分别将3块孔板放入37℃、5%CO₂培养箱中，水平放置，按预定检测时间点分批取出检测：24h板→Day 3上午取出；48h板→Day 4上午取出；72h板→Day 5上午取出；',
    '在实验记录本上详细记录给药操作的时间（精确到分钟），作为后续计时的基准。',
]
add_numbered_list(doc, drug_steps)

# ── DAY 3/4/5 ──
doc.add_heading('3.3  Day 3（24h检测）：CCK-8检测操作SOP', level=2)

p = doc.add_paragraph(
    'Day 3核心任务为24h时间点的CCK-8（Cell Counting Kit-8，细胞计数试剂盒-8）检测。'
    'CCK-8中的WST-8（2-(2-methoxy-4-nitrophenyl)-3-(4-nitrophenyl)-5-(2,4-disulfophenyl)-2H-tetrazolium）'
    '可被线粒体脱氢酶还原为橙黄色的甲瓒（formazan），其在450nm处的吸光度（OD₄₅₀）与活细胞数成正比[18,19]。'
)
set_para_spacing(p, before=4, after=4)

doc.add_heading('3.3.1  检测前准备（提前60-90分钟）', level=3)
prep_det = [
    '打开酶标仪电源，开机预热30分钟（必须预热，确保读数稳定）；',
    '酶标仪参数设置：检测波长450nm，参比波长600nm（或630nm，根据仪器型号）；完成自校准程序；',
    '从4℃冰箱中取出CCK-8试剂盒，置于室温（避光位置）平衡30分钟；',
    '全程注意避光：CCK-8工作液见光后反应加速，操作时应尽量减少光照，可用黑色盖布遮盖或在较暗条件下操作。',
]
add_numbered_list(doc, prep_det)

doc.add_heading('3.3.2  CCK-8加样与孵育SOP', level=3)
cck8_steps = [
    '从培养箱取出24h检测板（Day 3），放置于超净台内；',
    '用倒置显微镜快速观察孔板整体状态：确认各孔无明显污染，培养基颜色正常；',
    '使用10μL多道移液枪（或单道，从低浓度到高浓度顺序），向所有实验孔（B2-G11，共60孔）'
    '及空白孔（B2-B6）加入10μL CCK-8试剂，每个孔均加，无例外；',
    '加样过程全程避光操作，吸头垂直插入孔内液面下方约3-4mm，缓慢排出CCK-8液，避免产生气泡；',
    '加样完毕，肉眼检查确认所有孔的CCK-8已均匀加入（可轻轻倾斜孔板在光线下观察）；',
    '用黑色铝箔或不透光盖板遮住孔板，放入37℃、5%CO₂培养箱中，避光孵育1.5小时（90分钟）；',
    '孵育期间严禁频繁开启培养箱（最多2次），避免温度和CO₂波动影响显色反应；',
    '孵育时间优化方法（新手预实验建议）：若出现所有孔颜色极深（OD＞2.5）或极浅（OD＜0.1）的情况，'
    '下次实验可相应缩短或延长孵育时间，建议在正式实验前用对照组细胞预实验确定最佳孵育时间（通常1-2h）。',
]
add_numbered_list(doc, cck8_steps)

doc.add_heading('3.3.3  酶标仪检测分步操作（点击级教程）', level=3)

p = doc.add_paragraph('以下以BioTek Epoch2酶标仪为例，其他品牌操作界面略有差异，参照说明书相应设置。')
set_para_spacing(p, before=4, after=4)

reader_steps = [
    '孵育完成后立即取出96孔板：肉眼确认颜色变化正常（空白孔：浅黄色；对照孔：橙黄色；高浓度给药孔：颜色较浅，甚至接近空白）；',
    '将孔板轻放置于酶标仪板仓前，再次水平轻摇孔板5-10秒使颜色均匀；',
    '打开酶标仪操作软件（如Gen5或SoftMax Pro），新建检测任务（New Task）；',
    '选择96孔板板型（96 Well Plate），取消勾选"Lid"（无盖检测）；',
    '设置检测参数：读取模式（Read Mode）选择"Absorbance"；检测波长（Primary Wavelength）填入"450"nm；参比波长（Reference Wavelength）填入"600"nm；',
    '将96孔板正确放入板仓（注意孔板的A1方向与仪器标记一致），关闭板仓门；',
    '点击"Read Plate"（读取孔板）按钮，等待检测完成（通常15-30秒）；',
    '检测完成后，界面显示96孔板的OD值数据矩阵；',
    '数据导出：点击"File"→"Export"→选择"Microsoft Excel（.xlsx）"格式，保存至电脑桌面，命名格式为"苏木-CNE2-CCK8-24h-日期（如20260308）-操作者姓名缩写.xlsx"；',
    '同步备份：将数据文件复制至U盘或云存储（课题组共享盘），防止数据丢失；',
    '检测完毕的96孔板按照实验室生物安全规程处理：向孔内加入适量消毒液（84消毒液或高压灭菌），封口膜密封后放入生物废弃物垃圾袋，按感染性废物处理，不得随意丢弃。',
]
add_numbered_list(doc, reader_steps)

p_raw = doc.add_paragraph()
p_raw.add_run('原始数据导出格式示例（部分）：').bold = True
set_para_spacing(p_raw, before=6, after=2)

example_table = add_three_line_table(doc,
    ['孔位', '分组', '浓度（μg/mL）', 'OD₄₅₀（原始）', 'OD₄₅₀（扣除空白）'],
    [
        ['B2', '空白组', '—', '0.082', '—（空白对照）'],
        ['B3', '空白组', '—', '0.079', '—'],
        ['B7', '对照组', '0（DMSO）', '1.245', '1.163'],
        ['B8', '对照组', '0（DMSO）', '1.231', '1.149'],
        ['C2', '给药组C1', '400', '0.198', '0.116'],
        ['C3', '给药组C1', '400', '0.205', '0.123'],
        ['D2', '给药组C3', '4', '0.812', '0.730'],
        ['...', '...', '...', '...', '...'],
    ],
    col_widths=[2, 3, 3, 3.5, 3.5]
)

doc.add_heading('3.4  Day 4（48h检测）：CCK-8检测操作SOP', level=2)
p = doc.add_paragraph(
    'Day 4的48h检测操作步骤与Day 3（24h检测）完全相同，请严格按照3.3节所有步骤执行。'
    '以下几点需要特别注意：'
)
set_para_spacing(p, before=4, after=4)

d4_notes = [
    '确认取出的是"48h检测板"（Day 2给药后48小时的孔板），而非24h检测板；',
    '由于48h细胞增殖更充分，对照组OD值通常高于24h检测结果，孵育时间可根据颜色深浅适当调整（±15分钟）；',
    '数据文件命名时修改为"苏木-CNE2-CCK8-48h-日期-操作者缩写.xlsx"；',
    '检测完成后，48h检测板按生物安全规程处理，72h检测板继续放回培养箱培养至Day 5。',
]
add_numbered_list(doc, d4_notes)

doc.add_heading('3.5  Day 5（72h检测）：CCK-8检测操作SOP', level=2)
p = doc.add_paragraph(
    'Day 5的72h检测操作与3.3节相同，另需注意：72h孔板中细胞密度最高，'
    '若高浓度给药孔中细胞死亡完全，OD值将接近空白孔；CCK-8孵育时间可能需要缩短至1小时。'
    '72h检测完成后，应在当天进行初步数据整理，确认三个时间点数据完整性，进入第4章数据分析流程。'
)
set_para_spacing(p, before=4, after=4)

d5_final = [
    'Day 5实验结束后，在实验记录本上记录全部实验操作完成情况，签字确认；',
    '将24h/48h/72h三个时间点的原始数据文件整理为统一文件夹，命名为"苏木-CNE2-CCK8实验-独立重复1（或2/3）"；',
    '检查三个时间点原始数据的完整性：每块板需有所有分组的OD读数，无缺失孔（污染孔除外）；',
    '按照第4章操作指南进行数据处理和IC₅₀计算。',
]
add_numbered_list(doc, d5_final)

doc.add_page_break()

doc.save('/app/sandbox/session_20260308_142335_c3580fd63071/writing_outputs/final/苏木对CNE-2细胞增殖抑制率及IC50测定全流程操作手册_v2.0.docx')
print("Chapter 3 complete and saved.")
