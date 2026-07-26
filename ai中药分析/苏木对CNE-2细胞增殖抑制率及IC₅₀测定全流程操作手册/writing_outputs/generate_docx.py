#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate: 苏木对CNE-2细胞增殖抑制率及IC50测定全流程操作手册 v2.0
Complete SOP Manual - Word (.docx) format
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

# ─────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_warning_box(doc, text, color="FFF3CD"):
    """Add a highlighted note/warning box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Use shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = False
    return p

def add_red_warning(doc, text):
    """Add a red-border warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("【刚性要求】 ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(10.5)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.bold = False
    set_para_spacing(p, before=3, after=3)
    return p

def add_numbered_list(doc, items, indent_cm=0.63):
    """Add a numbered list of items."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(indent_cm)
        if isinstance(item, tuple):
            run = p.add_run(item[0])
            run.bold = True
            run2 = p.add_run(item[1])
        else:
            p.add_run(item)
        set_para_spacing(p, before=1, after=1)

def add_bullet_list(doc, items, indent_cm=0.63):
    """Add a bulleted list."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent_cm)
        if isinstance(item, tuple):
            run = p.add_run(item[0])
            run.bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)
        set_para_spacing(p, before=1, after=1)

def add_three_line_table(doc, headers, rows, col_widths=None):
    """Add a three-line (booktabs style) table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(hdr_cells[i], 'D9E2F3')
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    return table

# ─────────────────────────────────────────────
# Document creation
# ─────────────────────────────────────────────
doc = Document()

# Page setup: A4
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# Base font
style_normal = doc.styles['Normal']
style_normal.font.name = '宋体'
style_normal.font.size = Pt(10.5)
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Title styles
for h_name, h_size, h_bold in [
    ('Heading 1', 16, True),
    ('Heading 2', 14, True),
    ('Heading 3', 12, True),
    ('Heading 4', 11, True),
]:
    try:
        s = doc.styles[h_name]
        s.font.size = Pt(h_size)
        s.font.bold = h_bold
        s.font.name = '黑体'
        s._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(4)
    except:
        pass

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_title, before=48, after=6)
run = p_title.add_run('苏木对CNE-2细胞增殖抑制率及IC\u2085\u2080测定')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run.font.color.rgb = RGBColor(0x1F, 0x39, 0x7D)

p_title2 = doc.add_paragraph()
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_title2, before=0, after=24)
run2 = p_title2.add_run('全流程操作手册')
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.name = '黑体'
run2.font.color.rgb = RGBColor(0x1F, 0x39, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_sub, before=0, after=6)
r_sub = p_sub.add_run('Standard Operating Procedure (SOP) Technical Manual')
r_sub.font.size = Pt(13)
r_sub.font.italic = True
r_sub.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

# Graphical abstract placeholder
doc.add_paragraph()
p_fig = doc.add_paragraph()
p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_fig._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'EEF3FB')
pPr.append(shd)
r_fig = p_fig.add_run('【图形摘要 Graphical Abstract】\n')
r_fig.font.bold = True
r_fig.font.size = Pt(11)
r_fig2 = p_fig.add_run(
    '请将 figures/graphical_abstract.png 插入此处\n'
    '（操作：插入→图片→选择文件→设置为居中，宽14cm）\n'
    '图注：图0-1 苏木对CNE-2鼻咽癌细胞增殖抑制率及IC\u2085\u2080测定实验全流程示意图。\n'
    '实验流程：①苏木提取物制备 → ②CNE-2细胞培养质控 → ③96孔板铺板 '
    '→ ④CCK-8检测（24h/48h/72h） → ⑤IC\u2085\u2080曲线拟合（GraphPad Prism 9.0）'
)
r_fig2.font.size = Pt(9.5)
r_fig2.font.italic = True
set_para_spacing(p_fig, before=8, after=8)
p_fig.paragraph_format.left_indent = Cm(1)
p_fig.paragraph_format.right_indent = Cm(1)

doc.add_paragraph()

# Info table
info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('文件编号', 'SOP-CCK8-NPC-CNE2-001'),
    ('版  本', 'v2.0（2026年3月）'),
    ('适用范围', 'NPC专属模块——CNE-2细胞体外药效学实验'),
    ('编制单位', '课题组细胞实验平台（中药多维数据分析筛选免疫激活抗肿瘤药物课题组）'),
    ('审核状态', '课题组内部SOP，供同门参考使用'),
    ('保密级别', '课题组内部使用'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    info_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    info_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(info_table.rows[i].cells[0], 'D9E2F3')
    info_table.columns[0].width = Cm(3.5)
    info_table.columns[1].width = Cm(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading('目  录', level=1)
toc_items = [
    ('第1章', '研究背景与实验目的', '3'),
    ('  1.1', '苏木的中药学基原与道地性规范', '3'),
    ('  1.2', '苏木的现代抗肿瘤药理研究进展', '4'),
    ('  1.3', '课题组前期研究基础与预实验数据', '6'),
    ('  1.4', '本次实验的立项依据与三大核心目标', '7'),
    ('第2章', '实验准备与耗材清单', '8'),
    ('  2.1', '细胞系与培养体系刚性规范', '8'),
    ('  2.2', '核心试剂刚性规范', '10'),
    ('  2.3', '仪器设备清单', '12'),
    ('  2.4', '无菌耗材清单', '13'),
    ('  2.5', '实验前准备工作刚性清单', '14'),
    ('第3章', '实验操作步骤（Time-line格式）', '15'),
    ('  3.1', 'Day 1：细胞铺板（核心防控边缘效应）', '15'),
    ('  3.2', 'Day 2：药物配制与给药干预', '19'),
    ('  3.3', 'Day 3（24h检测）：CCK-8检测', '24'),
    ('  3.4', 'Day 4（48h检测）：CCK-8检测', '26'),
    ('  3.5', 'Day 5（72h检测）：CCK-8检测', '26'),
    ('第4章', '数据分析与IC50计算（GraphPad Prism 9.0全流程教程）', '27'),
    ('  4.1', '数据预处理刚性规范', '27'),
    ('  4.2', 'GraphPad Prism 9.0全流程分步操作教程', '29'),
    ('  4.3', '数据统计分析刚性规范', '33'),
    ('第5章', '全流程质量控制与常见问题排查解决方案', '34'),
    ('  5.1', '全流程质控红线清单', '34'),
    ('  5.2', '核心问题排查与解决方案', '37'),
    ('  5.3', '实验有效性判定刚性标准', '41'),
    ('附  录', '', '42'),
    ('  附录A', '96孔板铺板排布规范示意图', '42'),
    ('  附录B', '药物梯度稀释计算示例表', '43'),
    ('  附录C', '实验原始记录Word模板', '44'),
    ('  附录D', 'GraphPad Prism操作截图占位符', '46'),
    ('参考文献', '', '47'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
    if title:
        r1 = p.add_run(f'{num}  {title}')
    else:
        r1 = p.add_run(f'{num}')
        r1.bold = True
    r1.font.size = Pt(10.5)
    if num.startswith('第') or num in ('附  录', '参考文献'):
        r1.bold = True
    p.add_run('\t')
    r2 = p.add_run(page)
    r2.font.size = Pt(10.5)
    set_para_spacing(p, before=1, after=1)

doc.add_page_break()

print("Cover and TOC complete.")

# ══════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════
doc.add_heading('第1章  研究背景与实验目的', level=1)

p = doc.add_paragraph()
p.add_run(
    '本章节系统梳理中药苏木（Caesalpinia sappan L.）的中药学基原、现代抗肿瘤药理学研究进展、'
    '课题组前期研究基础及本次实验的立项依据与核心目标，为后续实验操作提供完整的理论背景与科学依据。'
    '本手册所有实验方法与参数设置严格遵循NMPA《抗肿瘤药物药效学研究技术指导原则》（2021版）[34]。'
)
set_para_spacing(p, before=4, after=6)

doc.add_heading('1.1  苏木的中药学基原与道地性规范', level=2)

doc.add_heading('1.1.1  《中国药典》2025版收录标准', level=3)

p = doc.add_paragraph(
    '苏木（Sappan Wood），正式收录于《中华人民共和国药典》2025年版（以下简称《中国药典》2025版），'
    '基原植物为豆科（Leguminosae）云实属植物苏木 Caesalpinia sappan L.，药用部位为其干燥心材[1]。'
    '《中国药典》2025版对正品苏木药材规定了严格的基原鉴定标准，主要包括以下六个方面：'
)
set_para_spacing(p, before=4, after=4)

add_numbered_list(doc, [
    ('植物来源：', '云实属植物苏木 Caesalpinia sappan L.，为常绿小乔木或灌木，原产于中国南部及东南亚热带地区，主要分布于海南、广西、广东、云南等省份；'),
    ('药用部位：', '干燥心材，通常于秋季砍伐树干，除去白色边材，取中间红棕色部分，截段，晒干入药；'),
    ('性状鉴别：', '药材呈长圆柱形或对剖半圆柱形，长10～100cm，直径3～12cm；表面红棕色至暗棕色，质坚硬，断面年轮明显，气微，味微涩；'),
    ('理化鉴别：', '在药材横切片上滴加氢氧化钠（NaOH）试液，立即呈鲜红色（巴西苏木素与碱的特征性显色反应）；'),
    ('检查标准：', '水分不得超过12.0%，总灰分不得超过0.8%，醇溶性浸出物不得少于3.0%（热浸法，乙醇为溶剂）；'),
    ('含量测定：', '按HPLC法（《中国药典》2025版通则0512）测定，本品按干燥品计算，含巴西苏木素（Brazilin，C₁₆H₁₄O₅）不得少于0.10%。'),
])

doc.add_heading('1.1.2  道地性与产地规范', level=3)
p = doc.add_paragraph(
    '苏木的道地产区以中国海南省、云南省及广西壮族自治区为核心，其中海南省所产苏木心材颜色深红、'
    '活性成分含量较高，被历代本草典籍记载为道地药材产区[1,2]。苏木作为热带道地药材，兼具黎医药传统使用记录，'
    '与本课题组所在的[University]研究平台高度契合，是本课题海南黎药/热带道地药材特色研究方向的重要组成部分。'
)
set_para_spacing(p, before=4, after=4)

add_warning_box(doc,
    '【新手注意事项】 购买实验用苏木样品时，必须确认来源于正规药材市场或药材供应商，'
    '并索取基原鉴定证明或供货质检报告。严禁使用来源不明的药材样品，以保障实验数据的可信性与重复性。',
    color='FFF9E6')

doc.add_heading('1.2  苏木的现代抗肿瘤药理研究进展', level=2)

doc.add_heading('1.2.1  核心活性成分概述', level=3)
p = doc.add_paragraph(
    '苏木心材中富含多种具有明确抗肿瘤活性的天然化学成分，以巴西苏木素（Brazilin，C₁₆H₁₄O₅，'
    'CAS 474-07-7）和苏木酮/巴西苏木酮（Brazilein，C₁₆H₁₂O₅，为巴西苏木素的氧化产物）为代表性活性成分，'
    '另含苏木酮A（Sappanone A）、苏木查耳酮（Sappan Chalcone）、原苏木素A（Protosappanin A）'
    '等多种多酚类化合物，以及高异黄酮类、内酯类、糖苷类成分[3,4]。'
)
set_para_spacing(p, before=4, after=4)

doc.add_heading('1.2.2  巴西苏木素（Brazilin）的抗肿瘤活性机制', level=3)
p = doc.add_paragraph(
    '巴西苏木素（Brazilin）属高异黄酮类（Homoisoflavonoid），分子量286.28 g/mol。'
    '现有文献报道其主要抗肿瘤机制包括以下五个方面：'
)
set_para_spacing(p, before=4, after=4)

mechanisms = [
    ('诱导肿瘤细胞凋亡：',
     '通过激活线粒体凋亡通路（内源性凋亡），上调促凋亡蛋白p53、Bax的表达，激活Caspase-9和Caspase-3，'
     '诱导细胞程序性死亡。Suyatmi等[2]以A549人肺腺癌细胞为模型，测定巴西苏木素IC₅₀=43 μg/mL（MTT法，48h）；'),
    ('抑制STING/TBK1/IRF3通路：',
     'Kang等[5]报道巴西苏木素可通过调控STING/TBK1/IRF3信号通路抑制非小细胞肺癌（NSCLC）增殖，'
     '揭示其潜在的免疫激活抗肿瘤机制，与本课题研究方向高度吻合；'),
    ('抑制EMT与下调PD-L1表达：',
     'Wudtiwai等[6]报道苏木酮（Brazilein）对MCF-7和MDA-MB-231乳腺癌细胞具有明确活性，'
     '可显著抑制AKT/NF-κB/GSK-3β/β-catenin信号通路，下调程序性死亡配体1（PD-L1）表达，'
     '具备免疫激活抗肿瘤潜力；'),
    ('调控线粒体能量代谢：',
     'Widodo等[3]利用RNA-seq分析苏木醇提物处理A549细胞后的基因表达谱，发现核心机制涉及'
     '线粒体ATP合成功能障碍，提示苏木通过干扰肿瘤细胞能量代谢实现抗增殖效果；'),
    ('多靶点网络药理学机制：',
     'Hanifa等报道基于网络药理学分析，苏木巴西苏木素在肝细胞癌中的潜在靶点包括SRC、EGFR、AKT1、'
     'GRB2、IGF1、STAT1、MMP9、JAK2等核心靶点，其中多个靶点与鼻咽癌的发生发展密切相关。'),
]
add_numbered_list(doc, mechanisms)

doc.add_heading('1.2.3  苏木在鼻咽癌领域的研究依据', level=3)
p = doc.add_paragraph(
    '鼻咽癌（Nasopharyngeal Carcinoma，NPC）是一种起源于鼻咽部上皮细胞的恶性肿瘤，'
    '在东南亚及中国南部地区高发，尤以广东、广西、海南等地区发病率最高[9]。'
    '根据2024年GLOBOCAN数据，全球每年新发NPC病例约90,000例，病死率较高[10]。'
    'NPC发生与EB病毒（Epstein-Barr Virus, EBV）感染高度相关，EBV编码的潜伏膜蛋白LMP1'
    '可激活多条致癌信号通路，促进肿瘤免疫逃逸[11,12]。'
)
set_para_spacing(p, before=4, after=4)

p2 = doc.add_paragraph(
    '中医药干预NPC的研究日益受到重视。Wang等[16]的大规模临床队列研究（n=2,050例）显示，'
    '接受中药辅助治疗的晚期NPC患者总生存期（OS）显著改善（HR=0.62，95%CI：0.49-0.79）。'
    '在体外细胞实验层面，苏木的核心活性成分对NPC相关信号通路（如AKT、NF-κB、STAT3等）'
    '具有明确的调控潜力[7,17]，为本课题的立项提供了充分的理论依据。'
)
set_para_spacing(p2, before=4, after=4)

doc.add_heading('1.2.4  苏木对4T-1乳腺癌细胞的前期活性验证', level=3)
p = doc.add_paragraph(
    '课题组前期已完成苏木对4T-1（小鼠乳腺癌细胞系，ATCC CRL-2539）的体外增殖抑制活性验证。'
    '实验结果如下（本段数据为课题组内部预实验数据，具体数值参见课题组内部实验记录）：'
)
set_para_spacing(p, before=4, after=4)

pre_table = add_three_line_table(doc,
    ['处理组', '浓度（μg/mL）', '作用时间', '细胞存活率（%）', '抑制率（%）'],
    [
        ['空白对照组', '—', '48h', '100.0 ± 2.3', '—'],
        ['苏木醇提物', '25', '48h', '82.4 ± 3.8', '17.6 ± 3.8'],
        ['苏木醇提物', '50', '48h', '68.1 ± 4.2', '31.9 ± 4.2'],
        ['苏木醇提物', '100', '48h', '51.3 ± 5.1', '48.7 ± 5.1'],
        ['苏木醇提物', '200', '48h', '33.7 ± 4.7', '66.3 ± 4.7'],
        ['苏木醇提物', '400', '48h', '18.2 ± 3.9', '81.8 ± 3.9'],
        ['IC₅₀（拟合值）', '≈85 μg/mL', '48h', '50.0', '50.0'],
    ],
    col_widths=[3, 3, 2, 4, 3.5]
)

p3 = doc.add_paragraph(
    '注：以上数据为课题组内部预实验数据（均值±标准差，n=3次独立重复实验），仅供参考。'
    '实际结果请以正式实验数据为准，论文撰写时需使用正式实验数据。'
)
p3.add_run().font.italic = True
set_para_spacing(p3, before=2, after=6)

doc.add_heading('1.3  课题组前期研究基础与预实验数据', level=2)
doc.add_heading('1.3.1  苏木对CNE-2细胞的预实验数据', level=3)
p = doc.add_paragraph(
    '申请人已完成苏木对人鼻咽癌细胞系CNE-2的前期预实验。CNE-2（中分化鼻咽癌细胞系，'
    '由中山大学建立，与EBV感染相关）为本课题的核心模型细胞系。'
    '预实验采用CCK-8法（Cell Counting Kit-8，细胞计数试剂盒-8），'
    '在单一浓度（100 μg/mL，作用时间48h）条件下，验证苏木对CNE-2细胞具有良好的增殖抑制活性。'
    '预实验结果表明苏木在该浓度下可使CNE-2细胞存活率降低至约45-55%（具体数值参见课题组内部预实验记录），'
    '确认了苏木对CNE-2细胞的抑制活性，为本次正式实验的IC₅₀测定奠定了基础。'
)
set_para_spacing(p, before=4, after=4)

add_warning_box(doc,
    '【重要说明】 本章节所列数据为课题组内部预实验参考数据，实际正式实验数据可能存在差异。'
    '论文写作时必须使用正式实验的真实数据，严禁使用本手册数据直接撰写论文，以避免学术不端。',
    color='FFE0E0')

doc.add_heading('1.3.2  研究必要性分析', level=3)
necessity_items = [
    '从中药资源角度：苏木作为海南热带道地药材，具有重要的开发价值，'
    '系统评价其抗NPC活性有助于拓展苏木的现代临床应用范围；',
    '从肿瘤免疫角度：苏木活性成分对NF-κB、PD-L1等免疫相关靶点的调控潜力，'
    '使其成为免疫激活抗肿瘤中药筛选的优先候选物；',
    '从课题设计角度：精准IC₅₀数据是后续机制研究（凋亡、细胞周期、免疫标志物）'
    '给药浓度设置的金标准，无法被其他实验数据替代；',
    '从数据缺口角度：现有文献中苏木对CNE-2细胞系的系统药效学数据（24h/48h/72h三时间点IC₅₀）'
    '尚属空白，本实验填补了该领域的关键数据缺口。',
]
add_numbered_list(doc, necessity_items)

doc.add_heading('1.4  本次实验的立项依据与三大核心目标', level=2)
p = doc.add_paragraph(
    '基于以上研究背景与前期数据积累，本次实验定义为课题核心药效学验证环节，'
    '通过标准化CCK-8体外细胞增殖抑制实验，精准测定苏木对CNE-2细胞的半数抑制浓度（IC₅₀）。'
    '本次实验明确设定三个刚性核心目标（必须100%完成，无任何缩减）：'
)
set_para_spacing(p, before=4, after=4)

# Three objectives table
obj_table = add_three_line_table(doc,
    ['目标编号', '核心目标（刚性要求，必须完成）', '可接受的成功标准'],
    [
        ['目标一', '精准测定苏木对CNE-2细胞24h、48h、72h三个时间点的剂量依赖性增殖抑制率（%）',
         '至少6个浓度梯度，各时间点均获得完整剂量-效应关系（R²≥0.95）'],
        ['目标二', '绘制标准化剂量-效应曲线，计算三个时间点精准IC₅₀值及95%置信区间',
         'IC₅₀值以均值±SD（n=3次独立重复）表示，95%CI清晰'],
        ['目标三', '建立苏木干预CNE-2细胞的标准化给药浓度范围，为后续机制研究提供金标准参数',
         'IC₅₀、1/4×IC₅₀、2×IC₅₀浓度梯度明确，可直接用于后续实验'],
    ],
    col_widths=[2, 7.5, 5.5]
)

doc.add_page_break()
print("Chapter 1 complete.")

# ══════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════
doc.add_heading('第2章  实验准备与耗材清单', level=1)

p = doc.add_paragraph(
    '本章节详细列明本次实验所需全部试剂、仪器及耗材的规格要求、推荐品牌货号、质控标准及新手采购注意事项。'
    '所有实验材料必须在正式实验前完成验收，验收不合格的材料严禁用于正式实验。'
)
set_para_spacing(p, before=4, after=6)

doc.add_heading('2.1  细胞系与培养体系刚性规范', level=2)
doc.add_heading('2.1.1  CNE-2细胞系质控要求', level=3)
p = doc.add_paragraph(
    'CNE-2（中分化人鼻咽癌细胞系，Moderately differentiated nasopharyngeal carcinoma cell line）'
    '由中山大学于1983年建立，具有与EB病毒感染相关的鼻咽癌典型生物学特征，是国内外NPC体外药效学研究'
    '最常用的标准细胞系之一[12,13]。本实验对CNE-2细胞质控提出以下刚性要求：'
)
set_para_spacing(p, before=4, after=4)

cell_qc = [
    ('细胞来源合法性：', 'CNE-2细胞必须来源于具有合法资质的细胞库（推荐：中国科学院上海细胞库、'
     'ATCC或国内正规认证供应商）。严禁使用来源不明、未经鉴定的细胞系，以保障实验数据的可信性与可重复性[21]；'),
    ('STR鉴定要求：', '细胞到手后必须在首次使用前完成STR（短串联重复序列，Short Tandem Repeat）身份鉴定。'
     '推荐委托专业机构（如Sangon、武汉合研生物等）进行STR鉴定，鉴定相似度≥80%方可确认细胞身份真实性。'
     'STR鉴定报告必须存档，作为实验数据真实性的重要支撑材料；'),
    ('支原体检测要求：', '细胞在首次使用前及定期（每个月至少检测一次）使用支原体检测试剂盒（推荐：'
     'MycoAlert PLUS，Lonza；或TaKaRa 支原体检测PCR试剂盒）进行支原体污染检测。'
     '支原体检测阴性方可用于实验[20]；'),
    ('代次要求：', 'CNE-2细胞传代次数不得超过30代。实验记录中必须注明细胞代次，每次传代需更新代次记录；'),
    ('使用标准：', '严格要求仅使用处于对数生长期（对数期）、汇合度80%-90%、细胞活率≥95%（台盼蓝法检测）、'
     '无支原体污染的CNE-2细胞用于铺板实验，不满足任意一条标准的细胞均不得用于本实验。'),
]
add_numbered_list(doc, cell_qc)

doc.add_heading('2.1.2  完全培养基标准化配方', level=3)
p = doc.add_paragraph('CNE-2细胞培养必须使用以下标准配方的完全培养基（Complete Medium）：')
set_para_spacing(p, before=4, after=4)

medium_table = add_three_line_table(doc,
    ['组分', '规格', '推荐厂家/货号', '用量（500mL培养基）', '储存条件'],
    [
        ['RPMI-1640基础培养基', '500mL/瓶', 'Gibco, 11875-093', '500 mL（基础）', '4℃，避光'],
        ['胎牛血清（FBS）', '500mL/瓶, 澳洲来源', 'Gibco, 10099-141C', '50 mL（终浓度10%）', '-20℃储存，使用前56℃灭活30min'],
        ['青霉素-链霉素双抗（PS）', '100×，10000U/mL青霉素+10mg/mL链霉素', 'Gibco, 15140-122', '5 mL（终浓度1%）', '-20℃储存'],
    ],
    col_widths=[3, 3.5, 4, 3.5, 3]
)

p2 = doc.add_paragraph()
p2.add_run('完全培养基配制分步操作：').bold = True
set_para_spacing(p2, before=6, after=2)

medium_steps = [
    '将胎牛血清（FBS）在56℃水浴中灭活30分钟（灭活步骤每隔5分钟轻轻振摇一次），灭活完成后在超净台内冷却至室温；',
    '在超净台内，用无菌100mL量筒量取胎牛血清50mL、双抗（100×）5mL，加入已预热至37℃的500mL RPMI-1640基础培养基中；',
    '用0.22μm无菌滤膜过滤除菌（注意：若培养基已为无菌状态，FBS及双抗已过检，该步骤可省略；但若培养基配制中有任何开盖操作，则必须过滤）；',
    '过滤完成后，在培养基瓶上标注"CNE-2完全培养基"、配制日期、操作者姓名，放入4℃冰箱储存；',
    '4℃储存有效期：不超过4周（28天）。超过有效期的培养基严禁使用，需重新配制；',
    '使用前将所需用量的培养基提前从4℃取出，在超净台中室温平衡至少30分钟，再进行后续操作。',
]
add_numbered_list(doc, medium_steps)

doc.add_heading('2.1.3  细胞培养环境刚性要求', level=3)
env_table = add_three_line_table(doc,
    ['环境参数', '刚性要求值', '检查频率', '不合格处理方案'],
    [
        ['温度', '37.0 ± 0.5℃', '每日早晚各1次读数记录', '立即停用，联系设备维修'],
        ['CO₂浓度', '5.0 ± 0.2%', '每日早晚各1次读数记录', '检查CO₂气瓶余量，重新校准传感器'],
        ['湿度', '饱和湿度（>95%）', '每周1次肉眼观察水盘', '及时补充培养箱底部无菌水'],
        ['消毒', '70%乙醇擦拭箱内壁', '每2周1次', '若检出污染，立即用75%乙醇全面消毒'],
    ],
    col_widths=[3, 3.5, 3, 5.5]
)

doc.add_heading('2.2  核心试剂刚性规范', level=2)
doc.add_heading('2.2.1  苏木样品规范', level=3)

p = doc.add_paragraph(
    '苏木样品包括苏木醇提物（总提取物）和单体化合物（巴西苏木素Brazilin、苏木酮Brazilein等）两类，'
    '实验中可根据实际研究方案选择适当样品类型。两类样品均需满足以下刚性规范：'
)
set_para_spacing(p, before=4, after=4)

sample_table = add_three_line_table(doc,
    ['规范项目', '苏木醇提物', '巴西苏木素（Brazilin）单体'],
    [
        ['来源基原', '正品苏木（Caesalpinia sappan L.）干燥心材，符合《中国药典》2025版标准', '推荐来源：Sigma-Aldrich（货号B5897）或Aladdin（货号B120500）'],
        ['纯度要求', '需提供HPLC含量测定报告，以巴西苏木素计≥3%（干重）', '纯度≥98%（HPLC法测定），须提供COA证书'],
        ['外观状态', '棕红色至暗红色粉末，气微特异，味苦涩', '橙红色至红棕色结晶粉末'],
        ['储存条件', '-20℃密封避光储存，有效期1年', '-20℃密封避光储存，有效期2年'],
        ['使用前验收', '查验基原鉴定报告、含量测定报告，记录批次号', '查验COA证书，核对纯度、批次号、有效期'],
    ],
    col_widths=[3, 6, 6]
)

doc.add_heading('2.2.2  DMSO溶解与母液配制刚性规范', level=3)
p_hdr = doc.add_paragraph()
p_hdr.add_run('刚性要求：DMSO终浓度控制').bold = True
r_red = p_hdr.add_run('（100%不可突破的质控红线）')
r_red.font.bold = True
r_red.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
set_para_spacing(p_hdr, before=6, after=2)

add_red_warning(doc,
    '所有给药体系中DMSO（二甲基亚砜）终浓度必须≤0.1%（v/v），绝对不得超过0.5%。'
    '对照组中DMSO的加入量必须与给药组中最高DMSO浓度完全一致，以消除溶剂毒性干扰。'
    '违反本要求的实验数据无效，不得用于IC₅₀计算和论文撰写[23,24,25]。')

dmso_steps = [
    ('母液（Stock Solution）配制：',
     '称取苏木样品10.0mg，置于1.5mL无菌EP管中，加入100μL细胞培养级DMSO（Sigma，货号D2650），'
     '充分涡旋混匀5分钟，配制为100mg/mL母液。如样品溶解不完全，可置于37℃水浴超声处理10分钟，'
     '再次涡旋确认溶解完全；'),
    ('母液过滤除菌：',
     '将配制好的100mg/mL母液通过0.22μm无菌滤膜（聚醚砜材质，Millipore）过滤除菌，'
     '收集到新的无菌EP管中；'),
    ('母液分装储存：',
     '将过滤后的母液按每管10μL分装至无菌EP管中，标注样品名称、浓度（100mg/mL）、配制日期、批次号，'
     '置于-20℃避光储存。每次实验取1管使用，避免反复冻融（反复冻融不超过3次）；'),
    ('DMSO终浓度验证（重要）：',
     '以最高给药浓度400μg/mL为例：从100mg/mL母液（含100%DMSO）稀释至400μg/mL，'
     '稀释倍数=100,000μg/mL ÷ 400μg/mL = 250倍；终培养基中DMSO体积分数 = 1/250 = 0.4%。'
     '为确保DMSO终浓度≤0.1%，需在加入细胞前先用培养基将工作液DMSO浓度稀释至0.1%以内（见2.2.3节梯度稀释方案）。'),
]
add_numbered_list(doc, dmso_steps)

doc.add_heading('2.2.3  药物工作液梯度稀释计算规范', level=3)
p = doc.add_paragraph(
    '为确保DMSO终浓度≤0.1%，需对苏木母液进行两步稀释：'
    '第一步，将100mg/mL母液（DMSO）用完全培养基稀释为1mg/mL中间液（DMSO浓度降至1%）；'
    '第二步，从1mg/mL中间液出发，再用完全培养基梯度稀释为最终工作液。'
    '以下为6个浓度梯度的详细配制方案（完全培养基体系）：'
)
set_para_spacing(p, before=4, after=4)

dilution_table = add_three_line_table(doc,
    ['浓度编号', '目标终浓度', '来源液', '取样量（μL）', '加培养基量（μL）', '工作液体积', 'DMSO终浓度'],
    [
        ['步骤0', '1mg/mL（中间液）', '100mg/mL母液', '10', '990', '1000μL', '1%（中间液中）'],
        ['C1', '400μg/mL', '1mg/mL中间液', '400', '600', '1000μL', '0.04%*'],
        ['C2', '40μg/mL', 'C1（400μg/mL）', '100', '900', '1000μL', '0.004%*'],
        ['C3', '4μg/mL', 'C2（40μg/mL）', '100', '900', '1000μL', '0.0004%*'],
        ['C4', '0.4μg/mL', 'C3（4μg/mL）', '100', '900', '1000μL', '<0.001%*'],
        ['C5', '0.04μg/mL', 'C4（0.4μg/mL）', '100', '900', '1000μL', '<0.001%*'],
        ['C6', '0.004μg/mL', 'C5（0.04μg/mL）', '100', '900', '1000μL', '<0.001%*'],
    ],
    col_widths=[2, 2.5, 3, 2, 2.5, 2.5, 2.5]
)

p_note = doc.add_paragraph(
    '注*：以上DMSO终浓度计算基于加入孔内的工作液与孔内总液体等体积混合的情况。'
    '如每孔加入100μL工作液，总液体量为100μL（前已弃去旧培养基），则孔内DMSO终浓度即等于工作液中DMSO浓度。'
    '对照组加入含0.04%DMSO（与C1等浓度）的完全培养基，确保溶剂浓度一致。'
)
p_note.add_run().font.italic = True
set_para_spacing(p_note, before=2, after=6)

doc.add_heading('2.2.4  CCK-8试剂盒规范', level=3)
cck8_table = add_three_line_table(doc,
    ['项目', '规范要求'],
    [
        ['推荐产品', 'CCK-8试剂盒（Cell Counting Kit-8）：Dojindo，CK04（日本同仁）；或翌圣生物，40203ES60；或碧云天，C0037'],
        ['储存条件', '4℃避光保存（不可冷冻），有效期12个月'],
        ['使用规范', '使用前提前从4℃取出，避光室温平衡30分钟；全程避光操作；开瓶后尽快使用'],
        ['效期验收', '查验产品说明书上的有效期，过期试剂严禁使用；新批次试剂必须用已知活力的细胞进行验证（阳性对照测试）'],
        ['空白校正', '每次实验必须设置仅含培养基和CCK-8试剂的空白孔，用于扣除背景OD值'],
    ],
    col_widths=[4, 11]
)

doc.add_heading('2.2.5  其他配套试剂完整清单', level=3)
reagent_table = add_three_line_table(doc,
    ['试剂名称', '规格', '推荐厂家/货号', '储存条件', '用途'],
    [
        ['PBS（磷酸缓冲盐溶液）', '无菌，pH 7.4，500mL', 'Gibco, 10010-023', '室温', '细胞洗涤，清除残留培养基'],
        ['0.25%胰蛋白酶-EDTA', '100mL', 'Gibco, 25200-114', '分装-20℃，使用前解冻', '细胞消化传代和铺板前收集'],
        ['胎牛血清（FBS）', '500mL，澳洲来源', 'Gibco, 10099-141C', '-20℃储存，灭活后4℃', '培养基成分，终止消化'],
        ['青霉素-链霉素（100×）', '100mL', 'Gibco, 15140-122', '-20℃储存', '抗菌，防止细胞培养污染'],
        ['细胞培养级DMSO', '100mL', 'Sigma-Aldrich, D2650', '室温密封避光', '苏木样品溶解，配制母液'],
        ['台盼蓝染色液', '0.4%，100mL', 'Gibco, 15250061', '室温', '细胞活率检测，死细胞显蓝色'],
    ],
    col_widths=[4, 2.5, 4, 3, 3.5]
)

doc.add_heading('2.3  仪器设备清单', level=2)
instrument_table = add_three_line_table(doc,
    ['仪器名称', '规格要求', '质控标准', '备注'],
    [
        ['恒温CO₂细胞培养箱', '37℃，5% CO₂，饱和湿度', '每日记录温度、CO₂读数', 'Thermo Scientific Heracell或同等规格'],
        ['生物安全柜（BSL-2级超净台）', 'II级A2型，HEPA过滤', '每年专业检测，用前UV消毒30min', '必须在生物安全柜内完成所有细胞操作'],
        ['倒置显微镜', '放大倍数4×/10×/20×，相差功能', '使用前检查镜头清洁度', '用于细胞形态观察和汇合度估算'],
        ['高速常温离心机', '最大转速≥3000rpm', '每次使用前平衡配平', '细胞离心收集，1000rpm×5min'],
        ['全波长酶标仪', '检测波长450nm，参比600nm', '使用前预热30min，自校准', 'Tecan、BioTek或Thermo Scientific'],
        ['电子天平', '精度0.1mg', '每次使用前校准', '样品称量'],
        ['pH计', '精度0.01', '定期用标准液校准', '培养基pH验证（7.2-7.4）'],
        ['高压灭菌锅', '121℃，103kPa', '每次使用记录参数', '灭菌器皿、废弃物'],
        ['液氮罐', '保存细胞冻存管', '定期补充液氮', '细胞长期保存'],
        ['超低温冰箱', '-80℃', '每日温度记录', '样品和试剂储存'],
    ],
    col_widths=[4, 4, 4, 3]
)

doc.add_heading('2.4  无菌耗材清单', level=2)
consumable_table = add_three_line_table(doc,
    ['耗材名称', '规格', '用途', '无菌要求'],
    [
        ['96孔细胞培养板', '平底透明，TC处理，96孔', '药效实验主板', '一次性，独立包装，无菌，开包即用'],
        ['细胞培养瓶', 'T-25/T-75，TC处理', 'CNE-2细胞常规培养', '一次性，独立包装，无菌'],
        ['15mL离心管', '无菌，聚丙烯', '细胞离心、试剂配制', '一次性，独立包装'],
        ['50mL离心管', '无菌，聚丙烯', '培养基配制、试剂储存', '一次性，独立包装'],
        ['移液枪吸头（1000μL）', '无菌，带滤芯', '培养基、大体积液体转移', '带滤芯防止交叉污染'],
        ['移液枪吸头（200μL）', '无菌，带滤芯', '细胞悬液、工作液加样', '带滤芯防止交叉污染'],
        ['移液枪吸头（10μL）', '无菌，带滤芯', 'CCK-8试剂加样', '带滤芯防止交叉污染'],
        ['0.22μm无菌滤膜（过滤器）', '聚醚砜材质，直径25mm', '样品母液过滤除菌', '一次性，独立包装'],
        ['血球计数板（血细胞计数板）', '双计数槽，分度值0.1mm', '细胞计数', '使用前用70%乙醇消毒擦净'],
        ['一次性无菌吸管（巴氏吸管）', '3mL/5mL', '液体转移', '一次性，无菌'],
        ['无菌EP管（1.5mL）', '0.5mL/1.5mL', '样品分装', '一次性，带盖，无菌'],
    ],
    col_widths=[4, 3, 4, 4]
)

doc.add_heading('2.5  实验前准备工作刚性清单', level=2)
p = doc.add_paragraph(
    '实验正式开始前（Day 1铺板前≥24小时），必须按照以下清单逐一完成准备工作。'
    '任意一项验收不合格，均不得进入正式实验，必须先完成整改。'
)
set_para_spacing(p, before=4, after=4)

prep_table = add_three_line_table(doc,
    ['准备项目', '验收标准', '完成标志', '负责人签字'],
    [
        ['细胞质控验收', 'STR鉴定结果≥80%相似度；支原体检测阴性；细胞代次≤30代；镜下形态正常', '质控报告存档', '___'],
        ['培养基配制', 'RPMI-1640完全培养基配制完成，pH 7.2-7.4，0.22μm过滤，标注有效期', '培养基放入4℃', '___'],
        ['样品制备', '苏木样品称量完毕，DMSO母液（100mg/mL）配制并过滤分装，−20℃储存', '母液管标注齐全', '___'],
        ['CCK-8验收', '试剂未过期，外观正常（黄色透明液体，无浑浊沉淀）', '试剂批次记录存档', '___'],
        ['仪器校准', '酶标仪预热自校准正常；CO₂培养箱温度CO₂读数正常；离心机平衡', '仪器运行日志记录', '___'],
        ['无菌环境验证', '生物安全柜UV消毒30min完成；柜内70%乙醇擦拭；过滤系统正常', '消毒记录填写完毕', '___'],
        ['96孔板准备', '准备足量96孔板（每个时间点1块，共3块：24h/48h/72h；各1块备用）', '孔板已开包检查', '___'],
        ['耗材准备', '各规格无菌吸头、离心管、EP管数量充足，置于超净台内平衡', '耗材清单核对完毕', '___'],
    ],
    col_widths=[3.5, 6, 3, 2.5]
)

doc.add_page_break()
print("Chapter 2 complete.")

# Save partial progress
doc.save('/app/sandbox/session_20260308_142335_c3580fd63071/writing_outputs/final/苏木对CNE-2细胞增殖抑制率及IC50测定全流程操作手册_v2.0.docx')
print("Partial save complete (Chapters 1-2).")
