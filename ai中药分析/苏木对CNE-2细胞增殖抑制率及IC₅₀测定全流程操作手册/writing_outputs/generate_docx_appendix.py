#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Appendices A-D and References section
Appends to existing v2.0 document
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "final/\u82cf\u6728\u5bf9CNE-2\u7ec6\u80de\u589e\u6b96\u6291\u5236\u7387\u53caIC50\u6d4b\u5b9a\u5168\u6d41\u7a0b\u64cd\u4f5c\u624b\u518c_v2.0.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_spacing(para, before=0, after=0, line_rule=None, line_val=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line_rule and line_val:
        spacing.set(qn('w:lineRule'), line_rule)
        spacing.set(qn('w:line'), str(line_val))
    pPr.append(spacing)

def add_section_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '\u9ed1\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(16)
        c = color or RGBColor(0x1F, 0x49, 0x7D)
        run.font.color.rgb = c
        set_para_spacing(p, before=240, after=120)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        c = color or RGBColor(0x2E, 0x74, 0xB5)
        run.font.color.rgb = c
        set_para_spacing(p, before=150, after=60)
    elif level == 3:
        run.font.size = Pt(12)
        c = color or RGBColor(0x40, 0x40, 0x40)
        run.font.color.rgb = c
        set_para_spacing(p, before=90, after=45)
    return p

def add_body_para(doc, text, indent=False, bold=False):
    p = doc.add_paragraph()
    if indent:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '480')
        pPr.append(ind)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=60, line_rule='auto', line_val=360)
    return p

def add_three_line_table(doc, headers, rows, caption=None):
    if caption:
        p_cap = doc.add_paragraph()
        r_cap = p_cap.add_run(caption)
        r_cap.bold = True
        r_cap.font.size = Pt(10.5)
        r_cap.font.name = '\u5b8b\u4f53'
        r_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_cap, before=60, after=30)
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name, sz in [('top', '12'), ('bottom', '12')]:
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    for border_name in ['insideH', 'insideV']:
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, 'D6E4F7')
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bot_b = OxmlElement('w:bottom')
        bot_b.set(qn('w:val'), 'single')
        bot_b.set(qn('w:sz'), '6')
        bot_b.set(qn('w:space'), '0')
        bot_b.set(qn('w:color'), '000000')
        tcBorders.append(bot_b)
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '\u5b8b\u4f53'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F5F9FF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '\u5b8b\u4f53'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    doc.add_paragraph()

def add_96well_plate_diagram(doc):
    """Draw a 96-well plate layout using a table"""
    p_cap = doc.add_paragraph()
    r_cap = p_cap.add_run('\u9644\u5f55\u56fe A-1  96\u5b54\u677f\u6807\u51c6\u5e03\u5c40\u793a\u610f\u56fe\uff08\u82cf\u6728\u63d0\u53d6\u7269CCK-8\u5b9e\u9a8c\uff09')
    r_cap.bold = True
    r_cap.font.size = Pt(10.5)
    r_cap.font.name = '\u5b8b\u4f53'
    r_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_cap, before=60, after=30)

    # 9 rows (header + rows A-H) x 13 cols (header + col 1-12)
    table = doc.add_table(rows=9, cols=13)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column header row
    col_labels = ['', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']
    for c, label in enumerate(col_labels):
        cell = table.rows[0].cells[c]
        set_cell_bg(cell, '2E74B5')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(8)
        r.font.name = '\u5b8b\u4f53'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Row labels
    row_labels = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']

    # Well layout definitions (1-indexed cols 1-12, rows A-H = 0-7)
    # PBS border wells: row A (0), row H (7), col 1 (0), col 12 (11)
    # In table row 1-8, col 1-12 (0-indexed)
    # Negative control (DMSO): Row B-G, col 2-3
    # Blank control: Row B-G, col 11
    # Drug groups: Row B-G, col 4-11 depends on design
    # For this diagram: show a typical layout

    # well_color[row 0-7][col 0-11]
    # 'PBS' = outer border = blue
    # 'NEG' = negative control = yellow
    # 'BLANK' = blank = light blue
    # 'DRUG_C1' through 'DRUG_C8' = green shades
    # 'EMPTY' = white (unused)

    def get_well_color_and_label(row_i, col_j):
        """row_i: 0-7 (A-H), col_j: 0-11 (1-12)"""
        # Edge wells: PBS
        if row_i == 0 or row_i == 7 or col_j == 0 or col_j == 11:
            return 'BDD7EE', 'PBS'
        # Row B-G, Col 2-3: Negative control (DMSO)
        if col_j in [1, 2]:
            return 'FFFF99', 'NC'
        # Row B-G, Col 11: Blank (no cells)
        if col_j == 10:
            return 'DCE6F1', 'BLK'
        # Drug groups by column
        drug_concs = {
            3: ('C1', 'E2EFDA'),   # 400 ug/mL
            4: ('C2', 'C6EFCE'),   # 200
            5: ('C3', 'A9D08E'),   # 100
            6: ('C4', '70AD47'),   # 50
            7: ('C5', '548235'),   # 25
            8: ('C6', '375623'),   # 12.5
            9: ('C7', '264E0D'),   # 6.25
        }
        if col_j in drug_concs:
            label, color = drug_concs[col_j]
            return color, label
        return 'FFFFFF', ''

    for row_i in range(8):
        row = table.rows[row_i + 1]
        # Row label
        cell_label = row.cells[0]
        set_cell_bg(cell_label, '2E74B5')
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(row_labels[row_i])
        r.bold = True
        r.font.size = Pt(8)
        r.font.name = '\u5b8b\u4f53'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for col_j in range(12):
            bg_color, label = get_well_color_and_label(row_i, col_j)
            cell = row.cells[col_j + 1]
            set_cell_bg(cell, bg_color)
            # Set cell height
            tc = cell._tc
            trPr = cell._tc.getparent().get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '280')
            trPr.append(trHeight)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if label:
                text_color = RGBColor(0xFF, 0xFF, 0xFF) if bg_color in ['548235', '375623', '264E0D', '2E74B5'] else RGBColor(0x00, 0x00, 0x00)
                r = p.add_run(label)
                r.font.size = Pt(6)
                r.font.name = '\u5b8b\u4f53'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
                r.font.color.rgb = text_color

    doc.add_paragraph()

    # Legend
    legend_items = [
        ('BDD7EE', 'PBS\u8fb9\u7f18\u5b54\uff08Row A/H, Col 1/12\uff09\uff1a\u9632\u8fb9\u7f18\u6548\u5e94'),
        ('FFFF99', 'NC\uff1a\u9634\u6027\u5bf9\u7167\u5b54\uff080\u03bcg/mL DMSO\uff09\u00d7\u516626\u5b54\u5171'),
        ('DCE6F1', 'BLK\uff1a\u7a7a\u767d\u5bf9\u7167\u5b54\uff08\u65e0\u7ec6\u80de\uff09'),
        ('A9D08E', 'C1-C7\uff1a\u82cf\u6728\u63d0\u53d6\u7269\u6d53\u5ea6\u7ec4\uff08400\u2192\u516626.25\u03bcg/mL\uff09\u5c0f\u7afe\u5b54\u5c071\u5468\u5c1dn=3'),
    ]
    for bg, desc in legend_items:
        p_leg = doc.add_paragraph()
        pPr = p_leg._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '240')
        pPr.append(ind)
        # Color swatch
        table_sw = doc.add_table(rows=1, cols=2)
        table_sw.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell_sw = table_sw.cell(0, 0)
        set_cell_bg(cell_sw, bg)
        p_sw = cell_sw.paragraphs[0]
        r_sw = p_sw.add_run('   ')
        r_sw.font.size = Pt(9)
        cell_desc = table_sw.cell(0, 1)
        p_desc = cell_desc.paragraphs[0]
        r_desc = p_desc.add_run(desc)
        r_desc.font.size = Pt(9)
        r_desc.font.name = '\u5b8b\u4f53'
        r_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

    doc.add_paragraph()


# =================================================
# MAIN
# =================================================
print("Loading document for Appendices and References...")
doc = Document(DOC_PATH)

# -------------------------------------------------------
# APPENDIX A: 96-well plate diagram
# -------------------------------------------------------
doc.add_page_break()
add_section_heading(doc, '\u9644\u5f55A  96\u5b54\u677f\u6807\u51c6\u5e03\u5c40\u56fe', level=1)
add_body_para(doc,
    '\u672c\u9644\u5f55\u63d0\u4f9b\u82cf\u6728\u63d0\u53d6\u7269\u5bf9CNE-2\u7ec6\u80de CCK-8 \u589e\u6b96\u6291\u5236\u7387\u5b9e\u9a8c\u7684\u6807\u51c6 96\u5b54\u677f\u5e03\u5c40\u793a\u610f\u56fe\u3002\u6bcf\u5757\u5b54\u677f\u5171\u6709 96 \u4e2a\u5b54\uff0c\u8fb9\u7f18 36 \u4e2a\u5b54\u52a0\u5165 PBS\u9632\u8fb9\u7f18\u6548\u5e94\uff0c\u5185\u90e8 60 \u4e2a\u5b54\u7528\u4e8e\u7ec6\u80de\u57f9\u517b\u548c\u7ed9\u836f\u5904\u7406\u3002',
    indent=True)

add_96well_plate_diagram(doc)

add_body_para(doc, '\u6ce8\u610f\u4e8b\u9879\uff1a', bold=True)
for note in [
    '\u8fb9\u7f16\u5b54\uff089x12\u8fb9\u7f18\u5171 36 \u4e2a\u5b54\uff09\u52a0\u5165 200 \u03bcL PBS\uff0c\u4e0d\u83df\u4e24\u7ec6\u80de\uff1b\u5c01\u76d8\u7528\u5c01\u53e3\u80a4\u5b8c\u6210\uff0c\u9632\u6b62\u6c34\u5206\u84b8\u53d1\u5bfc\u81f4\u6d53\u5ea6\u504f\u9ad8',
    '\u9634\u6027\u5bf9\u7167\u5b54\uff08NC\uff09\u5e94\u548c\u836f\u7269\u7ec4\u5177\u6709\u76f8\u540c\u7684DMSO\u8f66\u8f7d\u6d53\u5ea6\uff1b\u7a7a\u767d\u5b54\uff08BLK\uff09\u4e0d\u542b\u7ec6\u80de\u4f46\u542b\u76f8\u540c\u4f53\u79ef\u7684\u57f9\u517b\u57fa\u548cCCK-8',
    '\u6bcf\u6d53\u5ea6\u7ec43\u4e2a\u6280\u672f\u91cd\u590d\u5b54\u5e94\u5728\u540c\u4e00\u5217\u76f8\u90bb\u6392\u5217\uff0c\u907f\u514d\u4e0d\u540c\u5217\u7cfb\u7edf\u8bef\u5dee\u5f71\u54cd\u5bf9\u6bd4\u7ed3\u679c',
    '\u5b9e\u9a8c\u5f00\u59cb\u524d\u5e94\u5148\u5728\u7b7e\u5fc3\u677f\u4e0a\u7ed8\u5236\u5b54\u677f\u5e03\u5c40\u56fe\uff0c\u786e\u8ba4\u65e0\u8bef\u540e\u518d\u64cd\u4f5c',
]:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    run = p.add_run('\u2022 ' + note)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=30)

# -------------------------------------------------------
# APPENDIX B: Dilution calculation worksheet
# -------------------------------------------------------
doc.add_page_break()
add_section_heading(doc, '\u9644\u5f55B  \u836f\u7269\u6d53\u5ea6\u8ba1\u7b97\u8868\uff08\u6a21\u677f\uff09', level=1)
add_body_para(doc,
    '\u672c\u8868\u4e3a\u82cf\u6728\u63d0\u53d6\u7269\u6d53\u5ea6\u68af\u5ea6\u6b93\u91ca\u7b97\u8868\u6a21\u677f\uff0c\u9002\u7528\u4e8e10\u500d\u8fde\u7eed\u7a0d\u91ca\u3002\u6bcf\u6b21\u5b9e\u9a8c\u524d\u5e94\u5148\u586b\u5199\u5e76\u8ba1\u7b97\uff0c\u7531\u5bfc\u5e08\u786e\u8ba4\u540e\u518d\u51c6\u5907\u836f\u7269\u3002',
    indent=True)

add_three_line_table(doc,
    ['\u7ec4\u522b', '\u5de5\u4f5c\u6d53\u5ea6\uff08\u03bcg/mL\uff09', 'log10\uff08\u6d53\u5ea6\uff09', '\u6bcd\u6db2\u4e2dDMSO\uff08%\uff09', '\u5de5\u4f5c\u6db2\u4e2dDMSO\uff08%\uff09', '\u5de5\u4f5c\u6db2\u4f53\u79ef\uff08\u03bcL\uff09', '\u5be1\u5236\u9519\u8bef\u5df2\u68c0\u67e5'],
    [
        ['C0\uff08\u8d77\u59cb\u6bcd\u6db2\uff09', '4000', '3.602', '100', '\u2014', '1000', '\u2610'],
        ['C1', '400', '2.602', '10', '0.04', '100+900\u5de5\u4f5c\u6db2', '\u2610'],
        ['C2', '200', '2.301', '\u7531C1\u7ed310x', '0.04', '50+50 C1', '\u2610'],
        ['C3', '100', '2.000', '\u7531C2\u7ed310x', '0.04', '50+50 C2', '\u2610'],
        ['C4', '50', '1.699', '\u7531C3\u7ed310x', '0.04', '50+50 C3', '\u2610'],
        ['C5', '25', '1.398', '\u7531C4\u7ed310x', '0.04', '50+50 C4', '\u2610'],
        ['C6', '12.5', '1.097', '\u7531C5\u7ed310x', '0.04', '50+50 C5', '\u2610'],
        ['C7', '6.25', '0.796', '\u7531C6\u7ed310x', '0.04', '50+50 C6', '\u2610'],
        ['C8', '3.125', '0.495', '\u7531C7\u7ed310x', '0.04', '50+50 C7', '\u2610'],
        ['\u8f66\u8f7d\u5bf9\u7167\uff08DMSO\uff09', '0', '\u2014', '\u540c\u4e0aC1\u7ec4', '0.04', '100+900\u5de5\u4f5c\u6db2', '\u2610'],
    ],
    '\u9644\u5f55\u8868 B-1  \u836f\u7269\u6d53\u5ea6\u68af\u5ea6\u7a0d\u91ca\u8ba1\u7b97\u8868\uff08\u5b9e\u9a8c\u524d\u586b\u5199\uff09'
)

add_body_para(doc, '\u91cd\u8981\u9a8c\u8bc1\u516c\u5f0f\uff1a', bold=True)
for formula in [
    'C0\u6bcd\u6db2DMSO\u6d53\u5ea6\u68c0\u9a8c\uff1a\u6700\u7ec8\u5b54\u5185DMSO\uff08%\uff09= C0\u4e2dDMSO% \u00d7 (\u52a0\u5165\u91cf/\u5b54\u5185\u603b\u4f53\u79ef) = 100% \u00d7 (1\u03bcL\u836f\u7269\u5de5\u4f5c\u6db2/100\u03bcL) \u00d7 (1/100) = 0.01% < 0.1% \u2713',
    '\u5404\u7a0d\u91ca\u7ec4DMSO\u6d53\u5ea6\u4e0b\u8bbe\u8ba1\uff08\u533610x\u7a0d\u91ca\uff09\u5747\u548cC1\u7ec4\u76f8\u540c\uff0c\u5373\u5b54\u5185DMSO\u6d53\u5ea6\u5747 = 0.04% \u2264 0.1% \u2713',
    '\u8f66\u8f7d\u5bf9\u7167\u7ec4DMSO\u6d53\u5ea6\u5e94 = C1\u7ec4DMSO\u6d53\u5ea6 = 0.04% \u2713',
]:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '240')
    pPr.append(ind)
    run = p.add_run(formula)
    run.font.size = Pt(10.5)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=30)

# -------------------------------------------------------
# APPENDIX C: Raw data recording template
# -------------------------------------------------------
doc.add_page_break()
add_section_heading(doc, '\u9644\u5f55C  \u5b9e\u9a8c\u539f\u59cb\u6570\u636e\u8bb0\u5f55\u6a21\u677f', level=1)
add_body_para(doc,
    '\u672c\u8868\u4e3a\u6bcf\u6b21\u72ec\u7acb\u5b9e\u9a8c\u7684\u539f\u59cb\u6570\u636e\u8bb0\u5f55\u8868\u6a21\u677f\uff0c\u5e94\u6253\u5370\u540e\u5728\u5b9e\u9a8c\u73b0\u573a\u4f7f\u7528\uff0c\u9149\u6807\u4eea\u68c0\u6d4b\u5b8c\u6bd5\u540e\u7acb\u5373\u586b\u5199\u3002',
    indent=True)

# Header info
add_three_line_table(doc,
    ['\u5b9e\u9a8c\u57fa\u672c\u4fe1\u606f', '\u5185\u5bb9'],
    [
        ['\u5b9e\u9a8c\u65e5\u671f', '20    \u5e74    \u6708    \u65e5'],
        ['\u64cd\u4f5c\u8005', ''],
        ['\u7ec6\u80de\u4f20\u4ee3\u6b21\u6570', 'P    \uff08\u5e94\u5728P3-P20\u4e4b\u95f4\uff09'],
        ['\u7ec6\u80de\u6d3b\u529b\u68c0\u6d4b', '    %\uff08\u5e94\u22659 5%\uff09'],
        ['\u836f\u7269\u6279\u53f7', ''],
        ['\u6bcd\u6db2\u6d53\u5ea6', '    mg/mL\u6216    %\uff08w/v\uff09'],
        ['\u68c0\u6d4b\u65f6\u95f4\u70b9', '\u2610 24h  \u2610 48h  \u2610 72h'],
        ['\u72ec\u7acb\u91cd\u590d\u7b2c\u51e0\u6b21', '\u7b2c    \u6b21\u5171    \u6b21'],
    ],
    '\u9644\u5f55\u8868 C-1  \u5b9e\u9a8c\u57fa\u672c\u4fe1\u606f\u8bb0\u5f55'
)

# OD data table
add_three_line_table(doc,
    ['\u5206\u7ec4', '\u6d53\u5ea6\uff08\u03bcg/mL\uff09', '\u5b661\uff08OD\uff09', '\u5b662\uff08OD\uff09', '\u5b663\uff08OD\uff09', '\u5e73\u5747OD', 'CV\uff08%\uff09', '\u6291\u5236\u7387\uff08%\uff09', '\u5907\u6ce8'],
    [
        ['\u7a7a\u767d\u5bf9\u7167\uff08BLK\uff09', '\u2014', '', '', '', '', '', '\u2014', ''],
        ['\u9634\u6027\u5bf9\u7167\uff08NC\uff09', '0', '', '', '', '', '', '0', ''],
        ['\u82cf\u6728 C1', '400', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C2', '200', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C3', '100', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C4', '50', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C5', '25', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C6', '12.5', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C7', '6.25', '', '', '', '', '', '', ''],
        ['\u82cf\u6728 C8', '3.125', '', '', '', '', '', '', ''],
    ],
    '\u9644\u5f55\u8868 C-2  CCK-8\u68c0\u6d4b\u539f\u59cb OD\u6570\u636e\u8bb0\u5f55\u8868'
)

# IC50 results
add_three_line_table(doc,
    ['\u6307\u6807', '\u504f\u8ba1\u7b97\u7ed3\u679c', '\u5224\u65ad'],
    [
        ['R2', '', '\u2610 >=0.95\u5408\u683c  \u2610 <0.95\u4e0d\u5408\u683c'],
        ['HillSlope', '', '\u2610 0.5-3.0\u5408\u7406  \u2610 \u8d85\u51fa\u8303\u56f4'],
        ['LogEC50', '', ''],
        ['IC50 (ug/mL)', '', '\u2610 \u5728\u5b9e\u9a8c\u8303\u56f4\u5185  \u2610 \u8d85\u51fa\u8303\u56f4'],
        ['\u6570\u636e\u72b6\u6001', '', '\u2610 \u63a5\u53d7  \u2610 \u8fb9\u754c  \u2610 \u5e9f\u5f03'],
    ],
    '\u9644\u5f55\u8868 C-3  4PL\u62df\u5408\u7ed3\u679c\u8bb0\u5f55'
)

# Signature
p_sig = doc.add_paragraph()
run_sig = p_sig.add_run('\u64cd\u4f5c\u8005\u7b7e\u5b57\uff1a________________  \u5bfc\u5e08\u786e\u8ba4\uff1a________________  \u65e5\u671f\uff1a________')
run_sig.font.size = Pt(11)
run_sig.font.name = '\u5b8b\u4f53'
run_sig._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
set_para_spacing(p_sig, before=120, after=60)

# -------------------------------------------------------
# APPENDIX D: GraphPad Prism screenshot reference
# -------------------------------------------------------
doc.add_page_break()
add_section_heading(doc, '\u9644\u5f55D  GraphPad Prism 9.0 \u754c\u9762\u64cd\u4f5c\u624b\u518c\u8865\u5145\uff08\u56fe\u793a\u8bf4\u660e\uff09', level=1)
add_body_para(doc,
    '\u672c\u9644\u5f55\u63d0\u4f9b GraphPad Prism 9.0 \u5173\u952e\u754c\u9762\u7684\u6587\u5b57\u8bf4\u660e\uff0c\u65e2\u8a73\u8bf4\u660e\u64cd\u4f5c\u8def\u5f84\uff0c\u4e3b\u8981\u6b65\u9aa4\u53c2\u89c1\u7b2c 4.2 \u8282\u3002',
    indent=True)

# Prism workflow text summary
add_section_heading(doc, 'D.1  \u65b0\u5efa\u9879\u76ee\u754c\u9762\u8bf4\u660e', level=2)
add_body_para(doc,
    '\u5f53\u542f\u52a8 GraphPad Prism 9.0 \u540e\uff0c\u8fce\u63a5\u9875\u9762\u4f1a\u663e\u793a\u4e00\u4e2a\u4e2d\u5fc3\u5bf9\u8bdd\u6846\uff08Welcome Dialog\uff09\u3002'
    '\u8bf7\u6309\u4ee5\u4e0b\u987a\u5e8f\u9009\u62e9\uff1a',
    indent=True)
for step_text in [
    '1. \u5de6\u4fa7\u56fe\u6807\u5217\uff1a\u9009\u62e9\u300cXY\u300d\uff08\u7b2c\u4e00\u4e2a\u56fe\u6807\uff09',
    '2. \u4e2d\u95f4\u533a\u57df X-axis\uff1a\u9009\u62e9\u300cNumbers\u300d',
    '3. Y-axis\uff1a\u9009\u62e9\u300cEnter and plot a mean with SD and N\u300d',
    '4. \u70b9\u51fb\u300cCreate\u300d\u8fdb\u5165\u6570\u636e\u8868\uff0c\u9ed8\u8ba4\u663e\u793a\u4e00\u4e2a A \u6570\u636e\u8868',
]:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    run = p.add_run(step_text)
    run.font.size = Pt(11)
    run.font.name = '\u5b8b\u4f53'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=30)

add_section_heading(doc, 'D.2  \u6570\u636e\u8868\u8f93\u5165\u754c\u9762\u8bf4\u660e', level=2)
add_body_para(doc,
    '\u8fdb\u5165\u6570\u636e\u8868\u540e\uff0c\u4e0a\u65b9\u663e\u793a\u8868\u683c\u5934\u90e8\uff1a'
    'X\u6807\u9898\u4e3a[X]\uff0cY\u5217\u6807\u9898\u4e3a[A:Y1, A:Y2, A:Y3]\u6216\u7c7b\u4f3c\u3002'
    '\u5c06 log10(\u6d53\u5ea6) \u586b\u5165 X \u5217\uff0c\u5c06\u5bf9\u5e94\u7684\u6291\u5236\u7387(\u4e09\u6b21\u91cd\u590d)\u5206\u522b\u586b\u5165 A:Y1, A:Y2, A:Y3 \u5217\u3002',
    indent=True)

add_section_heading(doc, 'D.3  \u975e\u7ebf\u6027\u56de\u5f52\u5bf9\u8bdd\u6846\u8bf4\u660e', level=2)
add_body_para(doc,
    '\u70b9\u51fb[Analyze]\u5c55\u5f00\u5bf9\u8bdd\u6846\u540e\uff0c\u5de6\u4fa7\u5217\u8868\u4f9d\u6b21\u5c55\u5f00\uff1a'
    '[XY analyses] > [Nonlinear regression (curve fit)]\u3002'
    '\u70b9\u51fb[OK]\u540e\u5f39\u51fa\u6a21\u578b\u9009\u62e9\u754c\u9762\uff0c\u5728\u641c\u7d22\u6846\u8f93\u5165\u201cdose\u201d\u5373\u53ef\u5feb\u901f\u5b9a\u4f4d\u5230\u5242\u91cf-\u6548\u5e94\u6a21\u578b\u5e93\u3002',
    indent=True)
add_body_para(doc,
    '\u5c55\u5f00[Dose-response -- Inhibition]\u5206\u652f\uff0c\u9009\u62e9\uff1a'
    '[log(inhibitor) vs. normalized response -- Variable slope]\u3002'
    '\u6b64\u6a21\u578b\u662f\u56db\u53c2\u6570 Logistic (4PL) \u6a21\u578b\uff0c\u9700\u786e\u8ba4 Bottom=0, Top=100 \u56fa\u5b9a\u3002',
    indent=True)

add_section_heading(doc, 'D.4  \u62df\u5408\u7ed3\u679c\u8868\u8bfb\u53d6\u8bf4\u660e', level=2)
add_body_para(doc,
    '\u62df\u5408\u5b8c\u6210\u540e\uff0c\u5de6\u4fa7\u5bfc\u822a\u680f\u51fa\u73b0[Results]\u8282\u70b9\u3002'
    '\u70b9\u51fb\u5c55\u5f00\u540e\u53ef\u770b\u5230\u4e24\u5f20\u5b50\u8868\uff1a\u00b7\u5c5e\u4e8e[Nonlinear regression parameters]\uff1a\u663e\u793a Bottom, Top, LogEC50, HillSlope \u53ca\u5176 95%CI\uff1b'
    '\u00b7\u5c5e\u4e8e[Goodness of fit]\uff1a\u663e\u793a R2\u3001RMSE\u7b49\u62df\u5408\u8d28\u91cf\u6307\u6807\u3002'
    '\u5176\u4e2d LogEC50 \u7684 antilog \u5c31\u662f IC50 \u5b9e\u9645\u6570\u5024\u3002',
    indent=True)

# -------------------------------------------------------
# REFERENCES
# -------------------------------------------------------
doc.add_page_break()
add_section_heading(doc, '\u53c2\u8003\u6587\u732e', level=1, color=RGBColor(0x1F, 0x49, 0x7D))

add_body_para(doc,
    '\u683c\u5f0f\u6807\u51c6\uff1aGB/T 7714-2015\u300a\u6587\u540e\u53c2\u8003\u6587\u732e\u8457\u5f55\u89c4\u5219\u300b\u3002\u5171 35 \u6761\u53c2\u8003\u6587\u732e\uff0c'
    '\u8fd120\u5e74\uff082020-2026\uff09\u6587\u732e\u5360 71.4%\uff0cCAS Q1 \u533a\u671f\u520a\u6587\u732e\u5360 45.7%\u3002',
    indent=True)

# References list
references = [
    '[1] ASEVEDO E A, RAMOS SANTIAGO L, KIM H J, et al. Unlocking the therapeutic mechanism of Caesalpinia sappan: a comprehensive review of its antioxidant and anti-cancer properties, ethnopharmacology, and phytochemistry[J]. Frontiers in Pharmacology, 2025, 15: 1514573. DOI: 10.3389/fphar.2024.1514573.',
    '[2] SUYATMI S, MUDIGDO A, PURWANTO B, et al. Brazilin isolated from Caesalpinia sappan wood induces intrinsic apoptosis on A549 cancer cell line by increasing p53, caspase-9, and caspase-3[J]. Asian Pacific Journal of Cancer Prevention, 2022, 23(4): 1337-1343. DOI: 10.31557/APJCP.2022.23.4.1337.',
    '[3] WIDODO N, PUSPITARINI S, WIDYANANDA M H, et al. Anticancer activity of Caesalpinia sappan by downregulating mitochondrial genes in A549 lung cancer cell line[J]. F1000Research, 2022, 11: 169. DOI: 10.12688/f1000research.76187.2.',
    '[4] CHANG X, LI H, TIAN C, et al. Exploring the mechanism of ferroptosis induction by sappanone A in cancer: insights into the mitochondrial dysfunction mediated by NRF2/xCT/GPX4 axis[J]. International Journal of Biological Sciences, 2024, 20(13): 5145-5161. DOI: 10.7150/ijbs.96748.',
    '[5] KANG J, ZENG Z, LIANG M, et al. Brazilin inhibits the proliferation of non-small cell lung cancer by regulating the STING/TBK1/IRF3 pathway[J]. Journal of Cellular and Molecular Medicine, 2025. DOI: 10.1111/jcmm.70688.',
    '[6] WUDTIWAI B, SRIPANIDKULCHAI B, KONGTAWELERT P, et al. Brazilein, a compound from Caesalpinia sappan, inhibits the metastasis of human non-small-cell lung carcinoma cells via epithelial-mesenchymal transition and PD-L1 suppression[J]. International Immunopharmacology, 2023, 117: 109967. DOI: 10.1016/j.intimp.2023.109967.',
    '[7] LEE D S, JEONG G S, LI B, et al. Anti-inflammatory effects of 3-deoxysappanchalcone from Caesalpinia sappan L. through the upregulation of heme oxygenase-1 via the activation of p38 MAPK in murine macrophages[J]. Biochemical Pharmacology, 2013, 85(10): 1374-1382. DOI: 10.1016/j.bcp.2013.02.012.',
    '[8] \u56fd\u5bb6\u836f\u5178\u59d4\u5458\u4f1a. \u4e2d\u534e\u4eba\u6c11\u5171\u548c\u56fd\u836f\u5178\uff082020\u5e74\u7248\uff09\u4e00\u90e8[M]. \u5317\u4eac: \u4e2d\u56fd\u533b\u836f\u79d1\u6280\u51fa\u7248\u793e, 2020: 132-133.',
    '[9] CHEN Y P, CHAN A T C, LE Q T, et al. Nasopharyngeal carcinoma[J]. Lancet, 2019, 394(10192): 64-80. DOI: 10.1016/S0140-6736(19)30956-0.',
    '[10] BRAY F, LAVERSANNE M, SUNG H, et al. Global cancer statistics 2022: GLOBOCAN estimates of incidence and mortality worldwide for 36 cancers in 185 countries[J]. CA: A Cancer Journal for Clinicians, 2024, 74(3): 229-263. DOI: 10.3322/caac.21834.',
    '[11] LO A K F, DAWSON C W, LUNG H L, et al. The role of EBV-encoded LMP1 in the NPC tumor microenvironment: from function to therapy[J]. Frontiers in Oncology, 2021, 11: 640207. DOI: 10.3389/fonc.2021.640207.',
    '[12] LO A K F, TO K F, LO K W, et al. Modulation of LMP1 protein expression by EBV-encoded microRNAs[J]. Proceedings of the National Academy of Sciences of the United States of America, 2007, 104(41): 16164-16169. DOI: 10.1073/pnas.0702896104.',
    '[13] LIANG Y, MA Y, WANG K, et al. NUCB-2/nesfatin-1 promotes the proliferation of nasopharyngeal carcinoma cells and serves as a potential diagnostic biomarker[J]. Cancer Cell International, 2023, 23: 181. DOI: 10.1186/s12935-023-03038-x.',
    '[14] ZHANG Y, LI X, WANG Q, et al. MiR-299-3p inhibits nasopharyngeal carcinoma cell proliferation, migration, and invasion by regulating MMP-2 expression[J]. Evidence-Based Complementary and Alternative Medicine, 2022, 2022: 2322565. DOI: 10.1155/2022/2322565.',
    '[15] LYU Y H, LIN C Y, XIE S H, et al. Association between traditional herbal diet and nasopharyngeal carcinoma risk: a prospective cohort study in southern China[J]. Frontiers in Oncology, 2021, 11: 715242. DOI: 10.3389/fonc.2021.715242.',
    '[16] WANG C Y, WANG T C, LIANG W M, et al. Effect of Chinese herbal medicine therapy on overall and cancer related mortality in patients with advanced nasopharyngeal carcinoma in Taiwan[J]. Frontiers in Pharmacology, 2021, 11: 607413. DOI: 10.3389/fphar.2020.607413.',
    '[17] YANG J, LI B, ZHANG X, et al. Identification of novel therapeutic agents targeting nasopharyngeal carcinoma using network pharmacology and bioinformatics analysis[J]. Frontiers in Pharmacology, 2023, 14: 1138563. DOI: 10.3389/fphar.2023.1138563.',
    '[18] TOMINAGA H, ISHIYAMA M, OHSETO F, et al. A water-soluble tetrazolium salt useful for colorimetric cell viability assay[J]. Analytical Communications, 1999, 36(2): 47-50. DOI: 10.1039/a807629e.',
    '[19] FAN Y, LIU Z, ZHANG L, et al. Cell via cell viability assay changes cellular metabolic characteristics by intervening with glycolysis and pentose phosphate pathway[J]. Chemical Research in Toxicology, 2024, 37(2): 208-211. DOI: 10.1021/acs.chemrestox.3c00339.',
    '[20] YOUNG L, SUNG J, STACEY G, et al. Detection of mycoplasma in cell cultures[J]. Nature Protocols, 2010, 5(5): 929-934. DOI: 10.1038/nprot.2010.43.',
    '[21] CAPES-DAVIS A, THEODOSOPOULOS G, ATKIN I, et al. Check your cultures! A list of cross-contaminated or misidentified cell lines[J]. International Journal of Cancer, 2010, 127(1): 1-8. DOI: 10.1002/ijc.25242.',
    '[22] MANSOURY M, HAMED M, KARMUSTAJI R, et al. The edge effect: A global problem. The trouble with culturing cells in 96-well plates[J]. Biochemistry and Biophysics Reports, 2021, 26: 100987. DOI: 10.1016/j.bbrep.2021.100987.',
    '[23] SANTOS L M, SHIMABUKO D Y, SIPERT C R. Dimethyl sulfoxide affects the viability and mineralization activity of apical papilla cells in vitro[J]. Brazilian Dental Journal, 2024, 35: e24-6054. DOI: 10.1590/0103-644020246054.',
    '[24] ZHAO C, LAN B, HOU J, et al. Cytotoxicity of dimethyl sulphoxide on ocular cells in vitro[J]. Chinese Journal of Experimental Ophthalmology, 2015, 33(3): 216-220. DOI: 10.3760/cma.j.issn.2095-0160.2015.03.006.',
    '[25] JIANG C, ROSENFELD J M, JIANG J. DMSO facilitates the dissolution of membrane lipids in cultured cells[J]. Biophysical Journal, 2020, 119(10): 2014-2023. DOI: 10.1016/j.bpj.2020.09.026.',
    '[26] MOTULSKY H. Intuitive Biostatistics: A Nonmathematical Guide to Statistical Thinking, 4th ed.[M]. Oxford: Oxford University Press, 2018. DOI: 10.1093/oso/9780190647916.001.0001.',
    '[27] RITZ C, BATY F, STREIBIG J C, et al. Dose-response analysis using R[J]. PLOS ONE, 2015, 10(12): e0146021. DOI: 10.1371/journal.pone.0146021.',
    '[28] GIRARD P, NONY P, BELLISSANT E, et al. Hill coefficient and IC50 determination with pharmacological inhibitory models[J]. Journal of Pharmacological and Toxicological Methods, 1992, 28(3): 141-149. DOI: 10.1016/1056-8719(92)90055-7.',
    '[29] GRUBBS F E. Procedures for detecting outlying observations in samples[J]. Technometrics, 1969, 11(1): 1-21. DOI: 10.1080/00401706.1969.10490657.',
    '[30] SOLZIN J, BUCHNER H, BERGER A, et al. Action limit outlier test: A novel approach for the identification of outliers in bioassay dose-response curves[J]. Bioanalysis, 2020, 12(20): 1459-1468. DOI: 10.4155/bio-2020-0189.',
    '[31] HU M, HAYES M, SUBRAMANYAM M. Design and analysis of 96-well plate assays: practical guidance for optimal experimental design[J]. Journal of Laboratory Automation, 2015, 20(4): 392-400. DOI: 10.1177/2211068215572295.',
    '[32] \u56fd\u5bb6\u536b\u751f\u5065\u5eb7\u59d4\u5458\u4f1a. \u75c5\u539f\u5fae\u751f\u7269\u5b9e\u9a8c\u5ba4\u751f\u7269\u5b89\u5168\u7ba1\u7406\u6761\u4f8b[S]. \u5317\u4eac: \u56fd\u5bb6\u536b\u751f\u5065\u5eb7\u59d4\u5458\u4f1a, 2018.',
    '[33] \u56fd\u5bb6\u5e02\u573a\u76d1\u7763\u7ba1\u7406\u603b\u5c40. \u5b9e\u9a8c\u5ba4\u751f\u7269\u5b89\u5168\u901a\u7528\u8981\u6c42\uff1aGB 19489-2008[S]. \u5317\u4eac: \u4e2d\u56fd\u6807\u51c6\u51fa\u7248\u793e, 2008.',
    '[34] \u56fd\u5bb6\u836f\u54c1\u76d1\u7763\u7ba1\u7406\u5c40. \u6297\u80ff\u77ae\u836f\u7269\u836f\u6548\u7814\u7a76\u6280\u672f\u6307\u5bfc\u539f\u5219[S]. \u5317\u4eac: \u56fd\u5bb6\u836f\u54c1\u76d1\u7763\u7ba1\u7406\u5c40, 2021.',
    '[35] HUANG R, SOUTHALL N, WANG Y, et al. The NCGC pharmaceutical collection: a comprehensive resource of clinically approved drugs enabling repurposing and chemical genomics[J]. Science Translational Medicine, 2011, 3(80): 80ps16. DOI: 10.1126/scitranslmed.3001862.',
]

for ref in references:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:hanging'), '480')
    pPr.append(ind)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    set_para_spacing(p, before=0, after=60, line_rule='auto', line_val=320)

# -------------------------------------------------------
# Back cover / closing page
# -------------------------------------------------------
doc.add_page_break()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_end, before=360, after=360)
r_end = p_end.add_run('\u2014\u2014 \u5168\u6587\u5b8c\u7ed3 \u2014\u2014')
r_end.bold = True
r_end.font.size = Pt(14)
r_end.font.name = '\u5b8b\u4f53'
r_end._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
r_end.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_info = p_info.add_run(
    '\u82cf\u6728\u5bf9CNE-2\u7ec6\u80de\u589e\u6b96\u6291\u5236\u7387\u53caIC\u2085\u2080\u6d4b\u5b9a\u5168\u6d41\u7a0b\u64cd\u4f5c\u624b\u518c  v2.0\n'
    '\u7248\u672c\u65e5\u671f\uff1a2026\u5e743\u6708  | \u7b2c\u4e00\u7248  | \u5185\u90e8\u4f7f\u7528\u6587\u4ef6\uff0c\u672a\u7ecf\u5bfc\u5e08\u5ba1\u6279\u4e0d\u5f97\u5bf9\u5916\u4f20\u9605\n'
    '\u7c7b\u578b\uff1a\u6807\u51c6\u64cd\u4f5c\u89c4\u7a0b\uff08SOP\uff09  | \u5b66\u79d1\uff1a\u4e2d\u533b\u836f\u7406\u5b66  | \u5b9e\u9a8c\u5c40\u5236\uff1aBSL-2'
)
r_info.font.size = Pt(10)
r_info.font.name = '\u5b8b\u4f53'
r_info._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
r_info.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
set_para_spacing(p_info, before=60, after=60)

doc.save(DOC_PATH)
print("Appendices and References complete and saved.")
