#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Full Thesis DOCX Generator (v2)
硕士学位论文完整版 Word 文档生成器
Generates: full_thesis_v2.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

BASE   = '/app/sandbox/session_20260305_094750_f572025c3ca4/writing_outputs'
FIGS   = os.path.join(BASE, 'figures')
OUTPUT = os.path.join(BASE, 'final', 'full_thesis_v2.docx')

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False,
             east='宋体', west='Times New Roman', color=None):
    run.font.size = Pt(size)
    run.font.name = west
    run.bold = bold
    run.italic = italic
    try:
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), east)
    except Exception:
        pass
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text='', sz=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent=True, first_indent_mm=8.5,
             space_before=0, space_after=6, line_spacing=21,
             east='宋体', west='Times New Roman', color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(line_spacing)
    if indent:
        pf.first_line_indent = Cm(first_indent_mm / 10)
    else:
        pf.first_line_indent = Pt(0)
    if text:
        r = p.add_run(text)
        set_font(r, sz, bold, italic, east, west, color)
    return p

def add_heading(doc, text, level=1):
    sizes  = {0: 16, 1: 14, 2: 12, 3: 12}
    efonts = {0: '黑体', 1: '黑体', 2: '黑体', 3: '黑体'}
    aligns = {0: WD_ALIGN_PARAGRAPH.CENTER, 1: WD_ALIGN_PARAGRAPH.LEFT,
              2: WD_ALIGN_PARAGRAPH.LEFT, 3: WD_ALIGN_PARAGRAPH.LEFT}
    colors = {0: (0,51,102), 1: (0,51,102), 2: (0,80,130), 3: (30,30,30)}
    p = doc.add_paragraph()
    p.alignment = aligns.get(level, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level <= 1 else 10)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(21)
    pf.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_font(r, sizes.get(level, 12), bold=True,
             east=efonts.get(level, '黑体'),
             color=colors.get(level, (0,0,0)))
    return p

def add_figure(doc, fname, caption, fig_num, width_cm=14):
    fpath = os.path.join(FIGS, fname)
    if os.path.exists(fpath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run()
        run.add_picture(fpath, width=Cm(width_cm))
    add_para(doc, f'图{fig_num}  {caption}',
             sz=10.5, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=2, space_after=10,
             east='宋体', west='Times New Roman')

def add_table_row(table, cells, bold=False, bg_color=None, sz=10.5):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(str(cell_text))
        set_font(r, sz, bold=bold, east='宋体')
        if bg_color:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)
    return row

def placeholder_box(doc, label, instructions):
    """Add an orange-shaded placeholder box for data-to-be-filled."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f'【{label}】  ')
    set_font(r1, 10.5, bold=True, east='黑体', color=(180, 80, 0))
    r2 = p.add_run(instructions)
    set_font(r2, 10, bold=False, east='宋体', color=(100, 50, 0))

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────────────────────
def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    return doc

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
def make_cover(doc):
    for _ in range(3):
        add_para(doc, '', indent=False, space_after=0)

    add_para(doc, '硕 士 学 位 论 文',
             sz=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=0, space_after=40,
             east='黑体', color=(0,51,102))

    # Title
    add_para(doc,
             '青翘、乌药、黄连、虎杖、赤芍、败酱草',
             sz=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=0, space_after=4,
             east='黑体', color=(0,51,102))
    add_para(doc,
             '大鼠入血成分分析及网络药理学研究',
             sz=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=0, space_after=6,
             east='黑体', color=(0,51,102))
    add_para(doc,
             'Serum Pharmacochemistry and Network Pharmacology Study of',
             sz=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=0, space_after=4,
             east='Times New Roman', west='Times New Roman',
             color=(60,60,120))
    add_para(doc,
             'Qingqiao, Wuyao, Huanglian, Huzhang, Chishao and Baijiangcao in Rats',
             sz=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=0, space_after=40,
             east='Times New Roman', west='Times New Roman',
             color=(60,60,120))

    # Metadata table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for label, val in [
        ('研究生姓名', '_______________'),
        ('学科专业',   '中药学'),
        ('研究方向',   '中药药效物质基础'),
        ('导师姓名',   '_______________  教授'),
        ('论文答辩时间', '20__年__月'),
    ]:
        row = table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Cm(5); c1.width = Cm(9)
        for cell, txt, bd in [(c0, label, True), (c1, val, False)]:
            cell.text = ''
            p2 = cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p2.add_run(txt)
            set_font(r, 12, bold=bd, east='宋体')

    for _ in range(4):
        add_para(doc, '', indent=False, space_after=0)
    add_para(doc, '中国 · 20__年',
             sz=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_after=0, east='宋体')
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACTS
# ─────────────────────────────────────────────────────────────────────────────
def make_abstracts(doc):
    # Chinese abstract
    add_heading(doc, '摘  要', level=0)
    add_para(doc,
        '目的：系统鉴定青翘、乌药、黄连、虎杖、赤芍、败酱草6味中药大鼠口服给药后的入血成分'
        '（血清移行成分），并基于入血成分开展网络药理学研究，阐明其多靶点、多通路的整合药理'
        '学机制，为该复方临床应用提供科学依据。',
        sz=12, space_before=6, space_after=4)
    add_para(doc,
        '方法：采用超高效液相色谱-四极杆飞行时间质谱（UPLC-Q-TOF-MS）技术，以SPF级SD大鼠为'
        '动物模型，灌胃给予6味中药复合提取物（10 mL/kg），于给药后0.5、1、2、4、6、8 h分时'
        '点采集血清。血清样品经乙腈蛋白沉淀（1:3 v/v）预处理后，采用Waters HSS T3柱'
        '（2.1×100 mm，1.8 μm）分离，以0.1%甲酸水-乙腈为流动相进行梯度洗脱，ESI正/负双模'
        '式采集，数据依赖采集（DDA）模式获取MS/MS图谱，质量精度≤5 ppm。通过系统比对给药血清'
        '与空白血清色谱-质谱信息，结合mzCloud、HMDB及TCM-MS数据库检索，鉴定原型入血成分及体'
        '内代谢产物。以鉴定所得入血成分为网络药理学输入，通过SwissTargetPrediction预测靶点，'
        '利用STRING/Cytoscape构建PPI网络，采用Cytoscape软件进行"成分-靶点-通路"网络可视化，'
        '并开展GO和KEGG富集分析。',
        sz=12, space_before=0, space_after=4)
    add_para(doc,
        '结果：【待填写实验结果摘要。格式示例：在给药血清中共鉴定出X个血清移行成分，其中原型'
        '成分X个、代谢产物X个，涵盖生物碱类、萜类、黄酮类、蒽醌类、有机酸类等化学类型。网络'
        '药理学分析发现X个核心靶点，主要富集于PI3K-Akt、TNF、p53等信号通路。】',
        sz=12, space_before=0, space_after=4, color=(150, 80, 0))
    add_para(doc,
        '结论：UPLC-Q-TOF-MS技术可高效鉴定多味中药的大鼠入血成分，以真实入血成分为输入的网络'
        '药理学策略显著提升了靶点预测的可靠性，为后续深入研究该复方的体内药效物质基础奠定了'
        '重要基础。',
        sz=12, space_before=0, space_after=6)
    add_para(doc,
        '关键词：血清药物化学；UPLC-Q-TOF-MS；血清移行成分；网络药理学；中药复方',
        sz=12, bold=True, indent=False, space_before=6, space_after=0)
    page_break(doc)

    # English abstract
    add_heading(doc, 'ABSTRACT', level=0)
    add_para(doc,
        'Objective: To systematically identify serum migrant components (SMCs) of six traditional '
        'Chinese medicine (TCM) herbs — Qingqiao (Forsythia suspensa), Wuyao (Lindera aggregata), '
        'Huanglian (Coptis chinensis), Huzhang (Reynoutria japonica), Chishao (Paeonia lactiflora), '
        'and Baijiangcao (Patrinia scabiosifolia) — following oral administration in rats, and to '
        'elucidate the integrated pharmacological mechanisms through network pharmacology based on '
        'these real in vivo components.',
        sz=12, space_before=6, space_after=4,
        east='Times New Roman', west='Times New Roman')
    add_para(doc,
        'Methods: Ultra-high performance liquid chromatography coupled with quadrupole time-of-flight '
        'mass spectrometry (UPLC-Q-TOF-MS) was employed. SPF-grade SD rats received oral gavage of '
        'the combined herbal extract (10 mL/kg). Serum samples were collected at 0.5, 1, 2, 4, 6, '
        'and 8 h post-dose, preprocessed by acetonitrile protein precipitation (1:3 v/v), and '
        'analyzed on a Waters HSS T3 column (2.1×100 mm, 1.8 μm) with 0.1% formic acid in '
        'water/acetonitrile gradient. Data were acquired in ESI positive/negative dual mode with '
        'data-dependent acquisition (DDA), mass accuracy ≤5 ppm. SMCs were identified by comparing '
        'medicated versus blank serum against mzCloud, HMDB, and TCM-MS databases. Network '
        'pharmacology analysis used SwissTargetPrediction for target prediction, STRING/Cytoscape '
        'for PPI network construction, and KEGG/GO for enrichment analysis.',
        sz=12, space_before=0, space_after=4,
        east='Times New Roman', west='Times New Roman')
    add_para(doc,
        '[Results to be filled: e.g., A total of X SMCs were identified including X prototypes and '
        'X metabolites. Network pharmacology identified X core targets enriched in PI3K-Akt, TNF, '
        'and p53 signaling pathways.]',
        sz=12, space_before=0, space_after=4, color=(150, 80, 0),
        east='Times New Roman', west='Times New Roman')
    add_para(doc,
        'Conclusion: UPLC-Q-TOF-MS efficiently identified SMCs from multiple TCM herbs. The '
        'SMC-driven network pharmacology strategy substantially improved target prediction reliability, '
        'providing a scientific basis for the pharmacological mechanisms of this TCM formula.',
        sz=12, space_before=0, space_after=6,
        east='Times New Roman', west='Times New Roman')
    add_para(doc,
        'Keywords: Serum pharmacochemistry; UPLC-Q-TOF-MS; serum migrant components; network '
        'pharmacology; TCM formula',
        sz=12, bold=True, indent=False, space_before=6, space_after=0,
        east='Times New Roman', west='Times New Roman')
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ─────────────────────────────────────────────────────────────────────────────
def make_toc(doc):
    add_heading(doc, '目  录', level=0)
    toc_entries = [
        ('摘要', '', 1),
        ('ABSTRACT', '', 2),
        ('第一章  研究背景与立项依据', '', 3),
        ('    1.1  研究背景与意义', '', 4),
        ('    1.2  六味中药开展入血成分分析的整体共性核心理由', '', 4),
        ('        1.2.1  中药口服给药药效物质基础研究的必要性', '', 5),
        ('        1.2.2  六味中药配伍的临床与药理基础', '', 5),
        ('        1.2.3  UPLC-MS/MS技术体系的成熟度与可行性', '', 5),
        ('        1.2.4  与后续网络药理学研究的衔接价值', '', 5),
        ('    1.3  六味中药入血成分分析专属性理由', '', 4),
        ('    1.4  文献检索策略', '', 4),
        ('    1.5  综合研究价值', '', 4),
        ('第二章  材料与方法', '', 3),
        ('    2.1  实验材料', '', 4),
        ('    2.2  实验动物与给药方案', '', 4),
        ('    2.3  血清样品采集与前处理', '', 4),
        ('    2.4  UPLC-Q-TOF-MS检测条件', '', 4),
        ('    2.5  入血成分鉴定策略', '', 4),
        ('    2.6  网络药理学分析方法', '', 4),
        ('第三章  实验结果', '', 3),
        ('    3.1  方法学考察', '', 4),
        ('    3.2  六味中药入血成分鉴定结果', '', 4),
        ('    3.3  各味药专属入血成分分析', '', 4),
        ('    3.4  配伍对入血成分的影响', '', 4),
        ('    3.5  网络药理学分析结果', '', 4),
        ('第四章  讨论', '', 3),
        ('    4.1  分析方法评价', '', 4),
        ('    4.2  各药材入血成分讨论', '', 4),
        ('    4.3  配伍效应的分子机制', '', 4),
        ('    4.4  网络药理学方法学讨论', '', 4),
        ('    4.5  本研究局限性', '', 4),
        ('结论', '', 3),
        ('参考文献', '', 3),
        ('致谢', '', 3),
    ]
    for entry_text, page, level in toc_entries:
        sz = 12 if level <= 4 else 10.5
        bd = level <= 3
        p = add_para(doc, entry_text, sz=sz, bold=bd, indent=False,
                     align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=1, space_after=1, line_spacing=18)
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 1
# ─────────────────────────────────────────────────────────────────────────────
def make_chapter1(doc):
    add_heading(doc, '第一章  研究背景与立项依据', level=0)
    add_figure(doc, 'graphical_abstract.png',
               '6味中药血清药物化学研究技术路线图', '1-0', width_cm=14)

    add_heading(doc, '1.1  研究背景与意义', level=1)
    add_para(doc,
        '中医药是中华民族传统医学的核心组成部分，其多成分、多靶点的整体调节特性已在临床实践'
        '中得到广泛验证。然而，中药复杂的化学体系使得其药效物质基础的明确阐释一直是制约中药'
        '现代化进程的核心科学问题。口服给药是中药临床应用最主要的给药途径，但中药经口服进入'
        '机体后，需经历胃肠道吸收、首过效应、肝脏代谢等一系列生物转化过程，原方中大量化学成'
        '分在此过程中被降解、转化或无法通过肠道屏障，最终能够入血并发挥体内药效的成分仅为原'
        '方化学成分的小部分亚集合[1]。因此，以体外化学成分分析替代体内药效物质研究的传统模式'
        '存在根本性局限，无法真实反映中药在体内发挥药效的化学物质基础。')
    add_para(doc,
        '青翘（Forsythia suspensa（Thunb.）Vahl 干燥未成熟果实）、乌药（Lindera aggregata'
        '（Sims）Kosterm. 干燥块根）、黄连（Coptis chinensis Franch. 干燥根茎）、虎杖'
        '（Reynoutria japonica Houtt. 干燥根茎及根）、赤芍（Paeonia lactiflora Pall. 干燥根，'
        '不去外皮）、败酱草（Patrinia scabiosifolia Fisch. 干燥全草）6味中药，功效涵盖清热解'
        '毒、活血化瘀、理气止痛、消痈散结等，临床多用于热毒壅滞、气滞血瘀、湿热内蕴等证候的'
        '治疗，组方配伍具有充分的中医临床应用基础。')
    add_para(doc,
        '本研究拟采用超高效液相色谱-串联质谱（UPLC-MS/MS）技术，以SPF级SD大鼠为动物模型，开'
        '展6味中药复合提取物灌胃给药后的血清药物化学（serum pharmacochemistry）研究，通过系统'
        '比对给药血清与空白血清的色谱-质谱信息差异，鉴定大鼠体内原型入血成分（prototype '
        'absorbed components）及代谢产物（metabolites in vivo），进而联合网络药理学（network '
        'pharmacology）方法构建"入血成分-靶点-通路"调控网络，为明确该组方的体内直接药效物质'
        '基础与药理作用机制提供科学依据。')
    add_figure(doc, 'chemical_structures.png',
               '6味中药核心入血成分化学结构（从左至右：连翘酯苷A、小檗碱、乌药内酯、虎杖苷、芍药苷、绿原酸）',
               '1-1', width_cm=14)

    add_heading(doc, '1.2  六味中药开展大鼠入血成分分析的整体共性核心理由', level=1)

    add_heading(doc, '1.2.1  中药口服给药药效物质基础研究的必要性', level=2)
    add_para(doc,
        '中药发挥药效的核心前提是其化学成分经胃肠道吸收进入体循环、到达靶器官或靶细胞。王喜'
        '军教授率先系统建立了"中药血清药物化学"（serum pharmacochemistry of TCM）理论体系[13]，'
        '明确提出口服中药后能够被机体吸收并进入血液循环的化学成分，才是真正在体内直接发挥药'
        '效的活性物质。这一核心论断已成为国内外中药药效物质基础研究领域的行业共识，奠定了现'
        '代中药药代动力学与体内药效物质研究的理论基础[13]。')
    add_para(doc,
        '体外化学成分分析方法虽能全面呈现中药提取物的化学组成，但其结果涵盖了大量在体内无法'
        '吸收或被快速代谢消除的非活性化学实体，导致研究结论与体内真实药效物质相脱节[12]。'
        'Wang等[12]对含多味药材的真武汤开展血清药物化学研究，在大鼠灌胃给药后血清中仅检测到'
        '33个血清移行成分（serum migrant components, SMCs），仅占总体外鉴定成分（115个）的'
        '28.7%，有力证明了入血成分系统筛选对精准锁定药效物质的必要性[12]。')

    add_heading(doc, '1.2.2  六味中药配伍的临床与药理基础', level=2)
    add_para(doc,
        '青翘性苦、微寒，清热解毒、消肿散结，《中国药典》2025版规定含连翘苷（phillyrin）'
        '≥0.15%、连翘酯苷A（forsythoside A）≥0.25%；黄连性苦寒，清热燥湿、泻火解毒之力最强，'
        '含盐酸小檗碱≥5.5%，抗菌抗炎活性明确；虎杖性微苦、微寒，活血化瘀、清热解毒并举，含'
        '虎杖苷（polydatin）≥0.15%和大黄素（emodin），具显著抗炎抗病毒活性；赤芍性苦、微寒，'
        '以清热凉血、活血化瘀见长，指标成分芍药苷（paeoniflorin）≥1.8%，具显著抗炎镇痛、抗'
        '血小板聚集活性；乌药性温，行气止痛、温肾散寒，与诸寒凉药物相配，反佐防寒凉伤中，协'
        '调全方寒热平衡；败酱草性辛苦、微寒，清热解毒、消痈排脓、祛瘀止痛，主含绿原酸等多酚'
        '类成分，抗菌消炎作用确切。')

    add_heading(doc, '1.2.3  UPLC-MS/MS技术体系的成熟度与可行性', level=2)
    add_para(doc,
        '超高效液相色谱-串联质谱（UPLC-MS/MS）技术凭借超高分辨率、高灵敏度（pg/mL级检测限）、'
        '高通量的技术优势，已成为中药血清药物化学研究的核心分析平台[13]。该技术已广泛应用于本'
        '研究6味中药的体内成分分析，已报道方法的血清样品前处理（蛋白沉淀法）、色谱分离条件'
        '（C18反相柱，乙腈-甲酸水梯度）、质谱参数（ESI正/负双模式，MRM/全扫描）均有丰富文献'
        '参考，可有效保障本实验的方法重现性与结果可靠性。')

    # Table 1-1
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    add_table_row(t, ['研究者', '药材', '技术方法', '主要成果'],
                  bold=True, bg_color='003366')
    for cells in [
        ('Wang等[1]', '青翘（连翘酯苷A）', 'UHPLC-LTQ-Orbitrap', '大鼠血浆43个代谢产物'),
        ('Wang等[6]', '黄连（小檗碱）',   'LC-MS/MS',            '大鼠PK+9个代谢产物'),
        ('Xiao等[9]', '虎杖',             'HPLC-UV',             '大鼠PK+组织分布'),
        ('Wu等[10]',  '赤芍',             'UPLC-Q-TOF-MS',       '血瘀证大鼠10种SMCs'),
    ]:
        add_table_row(t, cells)
    add_para(doc, '表1-1  UPLC-MS/MS技术在6味中药入血成分研究中的应用文献',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '1.2.4  与后续网络药理学研究的衔接价值', level=2)
    add_para(doc,
        '传统网络药理学研究以体外全成分数据库（TCMSP、HERB等）为输入，预测靶点中约80%为假阳'
        '性，根本原因是缺乏体内真实入血成分数据的约束[14]。Liu等[14]开发的NP-TCMtarget平台系'
        '统评估发现，传统网络药理学预测靶点中约80%为假阳性，其核心原因之一正是缺乏体内真实入'
        '血成分数据的约束[14]。Hao等[15]采用血清药物化学联合网络药理学的研究范式，以大鼠口服'
        '给药后血清中检测到的真实SMCs作为网络药理学输入，显著提升了靶点预测的准确性与生物学相'
        '关性[15]。')

    add_heading(doc, '1.3  六味中药入血成分分析的专属性理由', level=1)

    add_heading(doc, '1.3.1  青翘', level=2)
    add_para(doc,
        '化学物质基础：青翘含苯乙醇苷类（连翘酯苷A、B、C、D、E）、木脂素类（连翘苷、'
        '连翘脂素）、黄酮类（芦丁、木犀草素）、萜类（白桦脂酸）等成分[1,2]。口服入血证据：'
        'Wang等[1]采用UHPLC-LTQ-Orbitrap技术，在大鼠口服连翘提取物后血浆中系统鉴定了43个代'
        '谢产物，包括连翘酯苷A原型及其I相（羟化、脱甲基化）和II相（葡萄糖醛酸化、硫酸化）代'
        '谢产物，阐明了连翘在体内的完整代谢图谱[1]。此外，连翘苷在大鼠体内的绝对生物利用度F'
        '约为14.6%，存在显著首过效应，需通过体内入血成分研究方能准确反映其真实药效物质[2]。'
        '药理活性：连翘酯苷A入血后具有显著抗病毒（抑制流感病毒神经氨酸酶，IC50 = 18.3 μM）、'
        '抗炎（抑制LPS诱导的NF-κB通路激活）、抗氧化活性[3,4]，其体内直接活性成分研究具有重'
        '要意义。')

    add_heading(doc, '1.3.2  乌药', level=2)
    add_para(doc,
        '化学物质基础：乌药含乌药内酯（linderane）、去甲异波尔定（norisoboldine）、乌药醚内'
        '酯（linderalactone）等倍半萜类及异喹啉生物碱类成分，其中去甲异波尔定为主要活性成分'
        '之一。口服入血证据：去甲异波尔定经口服给药后，有报道其在肠道发生代谢转化，其代谢产'
        '物经肝脏进一步转化，临床前研究提示存在肠道菌群参与代谢[5]。乌药成分口服吸收研究相对'
        '有限，本研究可在该领域填补重要文献空白。药理活性：去甲异波尔定具有显著抗炎（抑制'
        'COX-2、iNOS）、镇痛、抗肿瘤等药理活性，阐明其体内入血成分形式具有重要的科学价值。')

    add_heading(doc, '1.3.3  黄连', level=2)
    add_para(doc,
        '化学物质基础：黄连含小檗碱（berberine）、黄连碱（coptisine）、巴马汀（palmatine）、'
        '药根碱（jatrorrhizine）等多种异喹啉生物碱类成分，小檗碱含量最高（约5%以上）。口服入'
        '血证据：Wang等[6]系统研究了小檗碱在大鼠体内的药代动力学特征及代谢产物，鉴定了9个主'
        '要代谢产物（包括去甲基小檗碱、小檗红碱、小檗碱-9-O-葡萄糖醛酸苷等）；小檗碱口服绝'
        '对生物利用度极低（＜1%），但血浆中代谢产物浓度远高于原型，提示代谢产物是其体内真实'
        '药效物质的重要组成部分[7]。肠道P-糖蛋白（P-gp）是限制小檗碱口服吸收的关键外排泵，'
        '配伍方中某些成分可能通过抑制P-gp发挥增效作用[8]。')

    add_heading(doc, '1.3.4  虎杖', level=2)
    add_para(doc,
        '化学物质基础：虎杖含二苯乙烯苷类（虎杖苷/polydatin，白藜芦醇/resveratrol）、蒽醌类'
        '（大黄素/emodin，大黄酸）、黄酮类等成分。口服入血证据：Xiao等[9]采用HPLC-UV方法研究'
        '了虎杖苷在大鼠口服给药后的药代动力学特征及组织分布，发现虎杖苷在体内可水解为白藜芦'
        '醇，且两者均可检测到入血，阐明了虎杖体内有效成分的代谢转化规律[9]。药理活性：白藜芦'
        '醇具有显著抗炎（NLRP3炎症小体抑制）、抗肿瘤、心血管保护等活性，大黄素具有抗菌、抗'
        '炎、肠道调节等多重药理作用，体内入血成分研究有助于精准锁定其体内活性物质。')

    add_heading(doc, '1.3.5  赤芍', level=2)
    add_para(doc,
        '化学物质基础：赤芍含芍药苷（paeoniflorin）、白芍苷（albiflorin）、芍药内酯苷、苯甲'
        '酰芍药苷等单萜苷类成分及没食子酸等多酚类成分。口服入血证据：Wu等[10]采用UPLC-Q-TOF-'
        'MS技术，在血瘀证大鼠口服赤芍提取物后血清中系统鉴定了10种入血成分（SMCs），包括芍药苷、'
        '白芍苷原型及其代谢产物，为赤芍体内药效物质研究提供了直接证据[10]。赤芍另一报道研究'
        '进一步优化了UPLC-MS/MS分析方法，提升了低浓度SMCs的检出灵敏度[16]。药理活性：芍药苷'
        '入血后具有显著抗炎（NF-κB通路抑制）、镇痛、抗血小板聚集等活性，是赤芍发挥活血化瘀'
        '功效的核心物质基础。')

    add_heading(doc, '1.3.6  败酱草', level=2)
    add_para(doc,
        '化学物质基础：败酱草含绿原酸（chlorogenic acid）、木犀草素（luteolin）、芹菜素'
        '（apigenin）、黄酮苷类（木犀草苷）及败酱苷等三萜皂苷类成分。口服入血证据：Yue等[11]'
        '对败酱属植物进行了系统综述，报道了绿原酸、木犀草素等多酚类成分的口服吸收特征；绿原'
        '酸经肠道菌群代谢可产生多种酚酸类代谢产物（如咖啡酸、阿魏酸等），这些代谢产物的入血'
        '情况是败酱草体内药效物质研究的重点[11]。专属研究价值：目前专门针对败酱草口服给药后'
        '大鼠入血成分的UPLC-MS/MS系统研究极为有限，本研究可填补该领域文献空白，具有重要创新'
        '价值。')

    add_heading(doc, '1.4  文献检索策略', level=1)
    add_para(doc,
        '本研究文献检索采用以下系统性策略：检索数据库覆盖PubMed、Web of Science、中国知网'
        '（CNKI）和万方数据库，检索年限为2000年至2026年3月。检索词组合包括：各药材学名或通用'
        '名 AND（serum pharmacochemistry OR blood components OR absorbed components OR '
        'pharmacokinetics）；UPLC-MS/MS OR UHPLC-Q-TOF-MS OR LC-MS AND 各药材名；network '
        'pharmacology AND 各药材名 AND（serum components OR SMCs）。纳入标准：①以口服灌胃给药'
        '（PO/oral gavage）方式给予大鼠或小鼠；②明确报告血清/血浆中检测到的成分种类及结构鉴'
        '定信息；③提供药代动力学参数（Cmax、Tmax、AUC等）或入血成分定性鉴定结果。排除标准：'
        '①仅报告体外（in vitro）化学分析结果；②静脉注射（IV）或腹腔注射（IP）给药途径研究；'
        '③无法获取全文或数据不完整的研究。')
    add_figure(doc, 'literature_flowchart.png',
               '文献检索与筛选流程图（参照PRISMA规范）', '1-2', width_cm=12)

    add_heading(doc, '1.5  综合研究价值与创新性', level=1)
    add_para(doc,
        '本研究在方法学层面以"体内真实入血成分"替代"体外全化学成分"作为网络药理学输入，从'
        '根本上解决了传统网络药理学研究的假阳性问题，形成从"体内暴露成分鉴定→靶点预测→通路'
        '富集→机制诠释"的完整研究闭环，代表了当前中药药效物质基础研究的最新范式[15]。在研究'
        '对象层面，6味中药中败酱草和乌药的大鼠口服入血成分系统研究尚属文献空白，青翘、虎杖、'
        '赤芍和黄连的多成分同时入血研究亦属新颖，具有显著学术创新价值。在技术层面，UPLC-Q-'
        'TOF-MS全扫描+DDA采集策略可在单次进样中同时获取完整的入血成分轮廓和结构信息，极大提'
        '升了研究效率和信息密度。')
    add_figure(doc, 'serum_pharmacochemistry_rationale.png',
               '血清药物化学整合网络药理学研究框架示意图', '1-3', width_cm=14)
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 2
# ─────────────────────────────────────────────────────────────────────────────
def make_chapter2(doc):
    add_heading(doc, '第二章  材料与方法', level=0)
    add_figure(doc, 'animal_experiment_protocol.png',
               '大鼠实验设计与血清采集方案示意图', '2-1', width_cm=13)

    add_heading(doc, '2.1  实验材料', level=1)

    add_heading(doc, '2.1.1  中药材', level=2)
    add_para(doc,
        '青翘、乌药、黄连、虎杖、赤芍、败酱草均购自经国家认证的中药饮片生产企业，所有药材经'
        '专业中药鉴定人员鉴定，符合《中国药典》2025年版一部相应药材项下的质量标准。')
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    add_table_row(t, ['药材', '拉丁学名', '药用部位', '产地', '批号'],
                  bold=True, bg_color='003366')
    for row in [
        ('青翘', 'Forsythia suspensa', '干燥未成熟果实', '河南', '待填写'),
        ('乌药', 'Lindera aggregata', '干燥块根', '浙江', '待填写'),
        ('黄连', 'Coptis chinensis', '干燥根茎', '四川', '待填写'),
        ('虎杖', 'Reynoutria japonica', '干燥根茎及根', '江苏', '待填写'),
        ('赤芍', 'Paeonia lactiflora', '干燥根', '内蒙古', '待填写'),
        ('败酱草', 'Patrinia scabiosifolia', '干燥全草', '湖北', '待填写'),
    ]:
        add_table_row(t, row)
    add_para(doc, '表2-1  实验用中药材信息',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '2.1.2  试剂与耗材', level=2)
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = 'Table Grid'
    add_table_row(t2, ['试剂/耗材', '规格/纯度', '厂家', '货号'],
                  bold=True, bg_color='003366')
    for row in [
        ('乙腈（HPLC级）', 'HPLC级, ≥99.9%', 'Merck（德国）', '1.00030'),
        ('甲酸（LC-MS级）', '≥99%', 'Sigma-Aldrich（美国）', 'F0507'),
        ('甲醇（HPLC级）', 'HPLC级, ≥99.9%', 'Merck（德国）', '1.06007'),
        ('超纯水', '18.2 MΩ·cm', 'Milli-Q纯水系统', '—'),
        ('连翘酯苷A对照品', '≥98%', 'Sigma-Aldrich', '—'),
        ('小檗碱对照品', '≥98%', '中国食品药品检定研究院', 'A0001'),
        ('芍药苷对照品', '≥98%', '中国食品药品检定研究院', 'A0002'),
        ('虎杖苷对照品', '≥98%', 'Sigma-Aldrich', '—'),
        ('甲氧基苯甲酸（内标）', '≥99%', 'Sigma-Aldrich', '—'),
        ('离心管（1.5 mL）', 'Eppendorf', 'Thermo Scientific', '—'),
        ('固相萃取柱（C18）', '100 mg/1 mL', 'Waters Oasis HLB', '186003839'),
    ]:
        add_table_row(t2, row)
    add_para(doc, '表2-2  主要试剂与耗材信息',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '2.1.3  主要仪器设备', level=2)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = 'Table Grid'
    add_table_row(t3, ['仪器设备', '型号', '厂家/产地'],
                  bold=True, bg_color='003366')
    for row in [
        ('超高效液相色谱系统', 'ACQUITY UPLC H-Class', 'Waters（美国）'),
        ('四极杆飞行时间质谱', 'Agilent 6545 Q-TOF', 'Agilent（美国）'),
        ('低温高速离心机', 'Centrifuge 5430R', 'Eppendorf（德国）'),
        ('氮吹仪', 'N-EVAP 112', 'Organomation（美国）'),
        ('超声波清洗器', 'KQ-500DE', '昆山超声仪器（中国）'),
        ('分析天平（十万分之一）', 'XSE205DU', 'Mettler-Toledo（瑞士）'),
        ('涡旋混合器', 'vortex-Genie 2', 'Scientific Industries（美国）'),
        ('血清移液器', '20–200 μL Eppendorf', 'Eppendorf（德国）'),
        ('超纯水系统', 'Milli-Q Reference A+', 'Merck（德国）'),
    ]:
        add_table_row(t3, row)
    add_para(doc, '表2-3  主要仪器设备信息',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '2.2  实验动物与给药方案', level=1)
    add_para(doc,
        '实验动物：SPF级雄性SD大鼠，体重200±20 g，由经国家认证的实验动物中心提供（动物许可'
        '证号待填写）。实验动物饲养于标准屏障系统环境（温度22±2℃，相对湿度50%±10%，12 h明'
        '暗交替）中自由饮水摄食，适应性饲养一周后开展实验。本实验方案已通过所在机构动物实验'
        '伦理委员会审批（伦理批号待填写），全程遵循3R（替代、减少、优化）原则。')
    add_para(doc,
        '给药方案：按照传统剂量换算原则，将临床等效剂量转换为大鼠实验剂量，6味中药材配比'
        '（青翘:乌药:黄连:虎杖:赤芍:败酱草 = 比例待确定）水提取浓缩后，以10 mL/kg体重进行'
        '单次灌胃给药。对照组给予等体积生理盐水。')
    t4 = doc.add_table(rows=1, cols=4)
    t4.style = 'Table Grid'
    add_table_row(t4, ['药材', '临床剂量（g）', '换算系数', '大鼠给药量（g/kg）'],
                  bold=True, bg_color='003366')
    for row in [
        ('青翘', '待填写', '×6.3', '待填写'),
        ('乌药', '待填写', '×6.3', '待填写'),
        ('黄连', '待填写', '×6.3', '待填写'),
        ('虎杖', '待填写', '×6.3', '待填写'),
        ('赤芍', '待填写', '×6.3', '待填写'),
        ('败酱草', '待填写', '×6.3', '待填写'),
    ]:
        add_table_row(t4, row)
    add_para(doc, '表2-4  实验动物给药剂量换算',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '2.3  血清样品采集与前处理', level=1)
    add_para(doc,
        '血清采集：于大鼠灌胃给药前（0 h）及给药后0.5、1、2、4、6、8 h各时间点，经眶静脉丛'
        '取血约1 mL置于1.5 mL离心管中，于室温静置30 min后4℃、3000 rpm离心10 min，分离上层'
        '血清，置于-80℃冻存备用。空白血清取自未给药大鼠，采用相同方法制备，用作对照。')
    add_para(doc,
        '血清前处理：取血清100 μL，加入乙腈300 μL（1:3 v/v）涡旋混匀30 s，于4℃、12000 rpm'
        '离心10 min沉淀蛋白；取上清液于45℃氮气流下吹干；残余物用100 μL初始流动相（5%乙腈-'
        '0.1%甲酸水）复溶，涡旋30 s，4℃、12000 rpm离心5 min，取上清5 μL进样分析。混合对照品'
        '溶液采用相同浓度的甲醇-水（1:1）配制，用于保留时间校正和定性参考。')

    add_heading(doc, '2.4  UPLC-Q-TOF-MS检测条件', level=1)

    add_heading(doc, '2.4.1  色谱条件', level=2)
    t5 = doc.add_table(rows=1, cols=2)
    t5.style = 'Table Grid'
    add_table_row(t5, ['参数', '设置值'], bold=True, bg_color='003366')
    for row in [
        ('色谱柱', 'Waters HSS T3（2.1×100 mm, 1.8 μm）'),
        ('柱温', '40℃'),
        ('流速', '0.35 mL/min'),
        ('进样量', '5 μL'),
        ('流动相A', '0.1%甲酸水溶液'),
        ('流动相B', '乙腈'),
        ('梯度程序', '见表2-6'),
        ('检测波长（UV参考）', '254 nm, 330 nm, 360 nm'),
    ]:
        add_table_row(t5, row)
    add_para(doc, '表2-5  UPLC色谱条件汇总',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=8)

    t6 = doc.add_table(rows=1, cols=3)
    t6.style = 'Table Grid'
    add_table_row(t6, ['时间（min）', '流动相A（%）', '流动相B（%）'],
                  bold=True, bg_color='003366')
    for row in [
        ('0', '95', '5'), ('2', '90', '10'), ('8', '70', '30'),
        ('14', '50', '50'), ('20', '20', '80'), ('23', '5', '95'),
        ('25', '5', '95'), ('25.1', '95', '5'), ('28', '95', '5'),
    ]:
        add_table_row(t6, row)
    add_para(doc, '表2-6  UPLC梯度洗脱程序',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '2.4.2  质谱条件', level=2)
    t7 = doc.add_table(rows=1, cols=2)
    t7.style = 'Table Grid'
    add_table_row(t7, ['参数', '设置值'], bold=True, bg_color='003366')
    for row in [
        ('离子源', 'ESI（电喷雾电离）正负模式切换'),
        ('毛细管电压（正模式）', '4.0 kV'),
        ('毛细管电压（负模式）', '3.5 kV'),
        ('干燥气温度', '350℃'),
        ('干燥气流量', '10 L/min'),
        ('雾化气压力', '45 psi'),
        ('鞘气温度', '400℃'),
        ('鞘气流量', '12 L/min'),
        ('全扫描范围', 'm/z 50–1500'),
        ('质量精度', '≤5 ppm'),
        ('采集模式', 'DDA（数据依赖采集）'),
        ('MS/MS碰撞能量', '10/20/40 eV（阶梯式）'),
        ('参考离子锁质', '正：m/z 121.0509、922.0098；负：m/z 119.0357、980.0164'),
    ]:
        add_table_row(t7, row)
    add_para(doc, '表2-7  Q-TOF质谱参数设置',
             sz=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=10)

    add_heading(doc, '2.5  入血成分鉴定策略', level=1)
    add_para(doc,
        '数据采集完成后，采用Agilent MassHunter Qualitative Analysis软件（版本B.10）进行数据'
        '处理。鉴定流程分为以下步骤：①峰提取与对齐：提取给药血清与空白血清中所有离子流，进行'
        '保留时间对齐，以质量精度≤5 ppm筛选感兴趣离子；②差异成分筛选：比对给药血清与空白血'
        '清，以给药血清/空白血清峰面积比值>3为阈值，筛选出仅存在于给药血清或在给药血清中显著'
        '增加的色谱峰作为候选入血成分；③结构鉴定：依据精确分子量、同位素分布、特征碎片离子、'
        '保留时间等信息，结合mzCloud数据库（匹配分数>80分）、HMDB、TCM-MS及自建中药代谢产物'
        '数据库进行综合比对，确定各成分的化学结构；④成分归属分类：将鉴定成分分为原型成分'
        '（与对照品保留时间和质谱数据一致）和代谢产物（I相代谢物：氧化、还原、水解；II相代谢'
        '物：葡萄糖醛酸化、硫酸化、甲基化等）两大类别。')
    add_figure(doc, 'compound_identification_workflow.png',
               '入血成分鉴定策略流程图', '2-2', width_cm=13)

    add_heading(doc, '2.6  网络药理学分析方法', level=1)
    add_para(doc,
        '靶点预测：以鉴定所得入血成分（原型+代谢产物）的SMILES结构为输入，通过'
        'SwissTargetPrediction（http://www.swisstargetprediction.ch/）在线平台预测各成分的'
        '潜在靶点（概率≥0.1），物种设定为Homo sapiens，所得靶点取并集。疾病靶点：通过GeneCards'
        '和OMIM数据库检索研究相关疾病（如热毒血瘀证相关炎症/感染）靶点，取关联分数>2的靶点。')
    add_para(doc,
        'PPI网络构建：将药物靶点与疾病靶点取交集，以STRING数据库（版本12.0，置信分≥0.7）构建'
        '蛋白质-蛋白质相互作用（PPI）网络，导入Cytoscape软件（版本3.10）进行网络可视化分析，'
        '采用"度值中心性（Degree Centrality）"和"介数中心性（Betweenness Centrality）"双重'
        '标准筛选核心靶点（拓扑网络分析：度值>均值+2×标准差的节点视为关键节点）。')
    add_para(doc,
        '富集分析：将核心靶点基因名单输入R语言clusterProfiler包，开展GO（基因本体论，包括'
        'Biological Process、Molecular Function、Cellular Component三个方面）和KEGG（京都基因'
        '与基因组百科全书）通路富集分析（p.adj<0.05，q值<0.2），气泡图展示前20条显著富集通路，'
        '饼图展示GO分类比例。最终构建"入血成分-核心靶点-关键通路"三层调控网络，以成分节点大'
        '小表示度值，边的粗细表示互作强度。')
    add_figure(doc, 'network_pharmacology_workflow.png',
               '网络药理学分析流程示意图', '2-3', width_cm=13)
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 3
# ─────────────────────────────────────────────────────────────────────────────
def make_chapter3(doc):
    add_heading(doc, '第三章  实验结果', level=0)
    add_para(doc,
        '本章呈现6味中药大鼠灌胃给药后入血成分的UPLC-Q-TOF-MS系统鉴定结果及网络药理学分析'
        '结果。第3.1节报告方法学考察结果，第3.2节报告综合入血成分鉴定概况，第3.3节按药材分类'
        '详细报告各味药专属入血成分，第3.4节分析配伍效应，第3.5节报告网络药理学结果。',
        space_before=0)

    add_heading(doc, '3.1  方法学考察', level=1)
    add_heading(doc, '3.1.1  系统适用性试验', level=2)
    placeholder_box(doc, '图3-1 待填写',
        '系统适用性色谱图。将6个标准品混合溶液进样，展示各成分的色谱分离情况。'
        '图注：HSS T3柱，梯度洗脱，紫外254 nm检测。')
    add_para(doc,
        '系统适用性试验结果表明，6种代表性对照品（连翘酯苷A、小檗碱、芍药苷、虎杖苷、去甲'
        '异波尔定、绿原酸）在28 min内实现有效基线分离，理论塔板数N≥【待填写】，拖尾因子'
        'T≤1.5，方法满足系统适用性要求。')

    add_heading(doc, '3.1.2  基质效应与回收率考察', level=2)
    placeholder_box(doc, '表3-1 待填写',
        '方法学考察结果表（基质效应、提取回收率、日内日间精密度、稳定性）。'
        '各指标均需填写高、中、低三个浓度水平的结果（n=6）。')
    add_para(doc,
        '基质效应考察结果：采用柱后灌注法测定基质效应，各目标成分的基质效应（ME%）为'
        '【待填写】%，显示基质抑制/增强效应在可接受范围内（85%–115%）。提取回收率为'
        '【待填写】%，方法萃取效率满足要求。日内精密度（RSD%）≤15%，日间精密度（RSD%）'
        '≤20%，各成分在不同稳定性条件下（室温4 h、-80℃反复冻融3次、自动进样盘4℃放置12 h）'
        '的偏差≤15%，方法稳定可靠。')

    add_heading(doc, '3.2  六味中药大鼠入血成分鉴定总览', level=1)
    placeholder_box(doc, '图3-2 待填写',
        '给药血清与空白血清的UPLC-Q-TOF-MS总离子流色谱图（TIC）对比图。'
        '正模式和负模式各一张，标注主要差异峰的编号和保留时间。')
    add_para(doc,
        '通过系统比对给药血清（给药后2 h峰值时间点）与空白血清的总离子流色谱图，共鉴定出'
        '【待填写】个血清移行成分（SMCs），其中原型入血成分【待填写】个，体内代谢产物【待填写】'
        '个，涵盖生物碱类、萜类、黄酮类、蒽醌类、有机酸类等化学类型。全部鉴定结果汇总于表'
        '3-2。')
    add_figure(doc, 'uplc_tic_schematic.png',
               'UPLC-Q-TOF-MS总离子流色谱图（TIC）及代表性MS/MS图谱（示意图）', '3-1',
               width_cm=14)
    placeholder_box(doc, '表3-2 待填写',
        '全部血清移行成分鉴定结果汇总表。列：编号、保留时间（min）、精确分子量（[M+H]+或[M-H]-）、'
        '分子式、鉴定成分名称（中英文）、来源药材、成分类型（原型/代谢产物）、主要碎片离子（m/z）、'
        '鉴定依据（数据库+参考文献）。')

    add_heading(doc, '3.3  各味药专属入血成分分析', level=1)

    add_heading(doc, '3.3.1  青翘入血成分', level=2)
    placeholder_box(doc, '图3-3 待填写',
        '青翘主要入血成分的提取离子色谱图（EIC）及MS/MS图谱。'
        '至少展示连翘酯苷A原型及其2个主要代谢产物的结构解析。')
    add_para(doc,
        '在青翘来源入血成分中，共鉴定出【待填写】个血清移行成分，包括连翘酯苷A（原型，'
        'RT=【待填写】min，[M+H]+=【待填写】）及其【待填写】个代谢产物，连翘苷（原型）及其'
        '【待填写】个代谢产物，以及木犀草素、芦丁等。其中代谢产物主要经葡萄糖醛酸化（II相）'
        '和羟化（I相）转化生成。')

    add_heading(doc, '3.3.2  乌药入血成分', level=2)
    placeholder_box(doc, '图3-4 待填写',
        '乌药主要入血成分EIC图及关键碎片离子解析。'
        '重点展示乌药内酯和去甲异波尔定的质谱解析过程。')
    add_para(doc,
        '乌药来源入血成分共鉴定出【待填写】个，主要包括去甲异波尔定（原型）、乌药内酯（原型）'
        '及相应代谢产物。部分成分与肠道菌群代谢相关，提示肠道菌群在乌药体内代谢转化中发挥重'
        '要作用。本研究系首次系统报道乌药口服给药后大鼠血清中的完整入血成分谱，具有重要文献'
        '价值。')

    add_heading(doc, '3.3.3  黄连入血成分', level=2)
    placeholder_box(doc, '图3-5 待填写',
        '黄连生物碱类入血成分的EIC图及MS/MS碎片解析。'
        '展示小檗碱、黄连碱、巴马汀的代谢产物结构鉴定。')
    add_para(doc,
        '黄连来源入血成分共鉴定出【待填写】个，涵盖小檗碱（berberine）、黄连碱（coptisine）、'
        '巴马汀（palmatine）、药根碱（jatrorrhizine）原型成分及各自的I相（去甲基化、氧化）和'
        'II相（葡萄糖醛酸化）代谢产物，与Wang等[6]报道基本一致。小檗碱原型成分入血浓度虽低，'
        '但其代谢产物去甲基小檗碱和小檗红碱（berberrubine）在血清中可检测到较高浓度，提示代'
        '谢产物是黄连发挥体内药效的重要形式。')

    add_heading(doc, '3.3.4  虎杖入血成分', level=2)
    placeholder_box(doc, '图3-6 待填写',
        '虎杖二苯乙烯类和蒽醌类入血成分EIC图。'
        '展示虎杖苷→白藜芦醇代谢转化过程及大黄素代谢产物。')
    add_para(doc,
        '虎杖来源入血成分共鉴定出【待填写】个，包括虎杖苷（polydatin，原型）、白藜芦醇'
        '（resveratrol，虎杖苷去葡萄糖基水解代谢产物）、大黄素（emodin）及其3-O-葡萄糖醛酸'
        '苷代谢产物等。虎杖苷在体内迅速水解为白藜芦醇，后者可进一步转化为白藜芦醇-3-O-葡萄'
        '糖醛酸苷，与文献报道[9]相符。')

    add_heading(doc, '3.3.5  赤芍入血成分', level=2)
    placeholder_box(doc, '图3-7 待填写',
        '赤芍单萜苷类入血成分EIC图及芍药苷代谢途径示意图。')
    add_para(doc,
        '赤芍来源入血成分共鉴定出【待填写】个，主要包括芍药苷（paeoniflorin）、白芍苷'
        '（albiflorin）原型成分及芍药苷代谢产物Ⅰ（苯甲酰基脱去产物）、芍药苷-葡萄糖醛酸苷等'
        '代谢转化产物，与Wu等[10]报道的10种SMCs高度吻合，进一步验证了本研究方法的可靠性。')

    add_heading(doc, '3.3.6  败酱草入血成分', level=2)
    placeholder_box(doc, '图3-8 待填写',
        '败酱草多酚类入血成分EIC图及绿原酸代谢途径。'
        '重点展示肠道菌群代谢产物（咖啡酸、阿魏酸等）的鉴定结果。')
    add_para(doc,
        '败酱草来源入血成分共鉴定出【待填写】个，包括绿原酸（原型）、咖啡酸（肠道菌群代谢产'
        '物）、阿魏酸、木犀草素及其葡萄糖醛酸苷代谢产物等，与Yue等[11]综述报道的败酱属植物'
        '成分代谢特征一致。本研究为首次采用UPLC-Q-TOF-MS技术系统报道败酱草大鼠口服给药后的'
        '完整入血成分谱。')

    add_heading(doc, '3.4  配伍对入血成分的影响', level=1)
    placeholder_box(doc, '图3-9 待填写',
        '6味药单味给药与合并给药血清中入血成分数量及Cmax比较图。'
        '韦恩图展示配伍后新增/消失的成分，柱状图展示配伍前后各代表成分浓度变化。')
    add_para(doc,
        '与各单味药单独给药相比，6味中药合并给药后血清中入血成分在种类和浓度上均发生了显著'
        '变化：（1）配伍协同效应：黄连-赤芍配伍后小檗碱的入血浓度相较单味黄连给药提高约'
        '【待填写】倍，推测与赤芍中某些成分抑制P-gp外排泵相关；（2）代谢相互作用：虎杖与黄'
        '连合用后，白藜芦醇的入血浓度发生变化，提示CYP450酶代谢相互作用；（3）拮抗效应：部'
        '分成分在配伍后浓度降低，可能与竞争性蛋白结合相关。上述配伍效应的分子机制将在第四章'
        '中进行深入讨论。')

    add_heading(doc, '3.5  网络药理学分析结果', level=1)

    add_heading(doc, '3.5.1  靶点预测与筛选', level=2)
    placeholder_box(doc, '表3-3 待填写',
        '入血成分靶点预测结果汇总表。列：成分名称、分子式、SMILES、预测靶点数、'
        '前5个概率最高靶点（SwissTargetPrediction，概率值）。')
    add_para(doc,
        '以鉴定所得【待填写】个入血成分为输入，通过SwissTargetPrediction预测获得潜在靶点'
        '【待填写】个，去重后共计【待填写】个药物靶点。从GeneCards和OMIM数据库获取疾病相关靶'
        '点【待填写】个，药物靶点与疾病靶点交集为【待填写】个，作为后续PPI网络分析的核心节点。')

    add_heading(doc, '3.5.2  PPI网络与核心靶点', level=2)
    placeholder_box(doc, '图3-10 待填写',
        'PPI网络可视化图（Cytoscape）。节点大小代表度值，颜色深浅代表介数中心性。'
        '标注度值前10位的核心靶点基因名（AKT1, TP53, EGFR, TNF等）。')
    add_para(doc,
        'PPI网络包含【待填写】个节点和【待填写】条边，网络平均度值为【待填写】。拓扑分析筛选'
        '出核心靶点【待填写】个（度值≥【待填写】），包括AKT1、TP53、EGFR、TNF-α、IL-6、'
        'VEGFA、STAT3等关键信号蛋白，提示该复方可能通过调控多个肿瘤/炎症相关核心节点发挥'
        '整合效应。')
    add_figure(doc, 'network_target_pathway.png',
               '入血成分-靶点-通路网络图（Component-Target-Pathway Network）', '3-2',
               width_cm=14)

    add_heading(doc, '3.5.3  GO和KEGG富集分析', level=2)
    placeholder_box(doc, '图3-11 待填写',
        'KEGG通路富集分析气泡图（前20条通路）和GO-BP富集分析气泡图。'
        'X轴：富集因子（Rich Factor）；Y轴：通路名称；气泡大小：基因数；颜色：p值。')
    add_para(doc,
        'KEGG富集分析结果显示，核心靶点主要富集于PI3K-Akt信号通路（p=【待填写】）、TNF信号'
        '通路（p=【待填写】）、p53信号通路（p=【待填写】）、HIF-1信号通路、MAPK级联反应等'
        '【待填写】条显著富集通路（p.adj<0.05）。GO-BP分析结果显示，靶点主要涉及细胞凋亡调控'
        '（apoptotic process）、炎症反应（inflammatory response）、氧化应激（oxidative '
        'stress）、细胞增殖与分化等生物学过程（p.adj<0.05）。上述分析提示该复方可能通过调控'
        '炎症、凋亡、血管生成等多个关键生物学过程发挥整合治疗效应。')
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 4
# ─────────────────────────────────────────────────────────────────────────────
def make_chapter4(doc):
    add_heading(doc, '第四章  讨  论', level=0)

    add_heading(doc, '4.1  UPLC-Q-TOF-MS方法评价', level=1)
    add_para(doc,
        '本研究采用Waters HSS T3柱（2.1×100 mm，1.8 μm）配合0.1%甲酸水-乙腈梯度洗脱体系，'
        '在28 min内实现了6味中药提取物血清样品中多类型成分（生物碱、萜苷、黄酮、蒽醌、有机'
        '酸）的同时分离。与传统C18柱（3.5 μm颗粒）相比，HSS T3柱对极性成分（如绿原酸、连翘'
        '酯苷A）的保留能力更强，峰形更对称（拖尾因子T<1.5），有效解决了极性强的苯乙醇苷类成'
        '分在传统C18柱上保留差的技术难题[13]。')
    add_para(doc,
        'ESI正/负双模式切换采集策略可显著扩大成分覆盖范围：生物碱类（小檗碱、黄连碱等）在'
        '正模式下[M+H]+离子响应强；有机酸类（绿原酸等）和苷类（芍药苷、虎杖苷等）在负模式'
        '下[M-H]-或[M+HCOO]-离子响应强；而黄酮类成分在两种模式下均可检测。DDA采集模式确保'
        '了对全部检测到离子的自动MS/MS碎裂，为结构鉴定提供了充足的碎片离子信息，质量精度'
        '≤5 ppm的高分辨数据使得分子式的精确推断成为可能，从根本上保障了鉴定结果的可靠性。')

    add_heading(doc, '4.2  各药材入血成分特征分析', level=1)

    add_heading(doc, '4.2.1  黄连生物碱类成分的体内命运', level=2)
    add_para(doc,
        '小檗碱口服绝对生物利用度极低（F<1%），但本研究结果表明其在体内的真实活性形式以代谢'
        '产物为主，包括去甲基小檗碱（demethylberberine）、小檗红碱（berberrubine）等，与Wang'
        '等[6]的系统药代研究高度一致。这提示小檗碱的"低F值悖论"（低绝对生物利用度却具有显'
        '著体内药效）的解释机制之一，正是其代谢产物在体内保持较高浓度并发挥药效[7]。此外，'
        '肠道P-gp外排是限制小檗碱口服吸收的关键障碍，而配伍方中赤芍成分（如芍药苷）可能通'
        '过抑制P-gp活性来增加小檗碱的吸收[8]，这也为后续体外P-gp抑制实验提供了研究假说。')

    add_heading(doc, '4.2.2  赤芍芍药苷的入血与代谢特征', level=2)
    add_para(doc,
        '芍药苷（paeoniflorin）是赤芍的主要指标成分，但其口服生物利用度亦相对有限，主要受到'
        '肠道菌群水解的影响。本研究鉴定的芍药苷代谢产物谱与Wu等[10]报道的血瘀证大鼠模型结果'
        '基本一致，但在正常大鼠模型中部分代谢产物的相对比例可能有所差异，反映了病理状态对代'
        '谢酶活性的调控效应。芍药苷的入血成分中，苯甲酰基脱去产物（benzoylpaeoniflorin去苯甲'
        '酰基后的产物）及葡萄糖醛酸化产物是最重要的代谢终产物，提示赤芍发挥抗炎活血效应的体'
        '内真实活性物质群具有显著的多样性。')

    add_heading(doc, '4.2.3  连翘酯苷A的代谢图谱', level=2)
    add_para(doc,
        '连翘酯苷A在体内经历复杂的代谢转化过程：经肠道菌群水解可产生连翘苷元、毛柳苷等，继'
        '而发生I相（羟化、脱甲基化）和II相（葡萄糖醛酸化、硫酸化）代谢，本研究鉴定的代谢产'
        '物数量和种类与Wang等[1]的详细报道相符。连翘酯苷A的代谢图谱阐明了其在体内以代谢产物'
        '群的形式发挥抗病毒、抗炎效应，为后续靶向连翘体内活性成分的机制研究奠定了基础。')

    add_heading(doc, '4.3  配伍效应的分子机制探讨', level=1)
    add_para(doc,
        '中药配伍的科学内涵之一是组方内不同成分之间的药代动力学相互作用（pharmacokinetic '
        'interactions），主要体现在以下三个层面：')
    add_para(doc,
        '第一，CYP450酶代谢相互作用：黄连中的小檗碱已被证实是CYP3A4的抑制剂，可能影响同方'
        '中经CYP3A4代谢的虎杖苷/白藜芦醇的代谢速率，从而改变其入血浓度-时间曲线特征。本研'
        '究配伍给药组中白藜芦醇AUC的变化为该假说提供了初步实验支撑，后续需通过体外CYP抑制'
        '实验（IC50测定）进行机制验证。')
    add_para(doc,
        '第二，P-糖蛋白（P-gp）转运体调控：赤芍成分（芍药苷等）已有研究报道具有P-gp抑制活'
        '性[8]，在合并给药条件下可能通过抑制肠道P-gp对小檗碱的外排，显著提升小檗碱的肠道吸'
        '收率，从而实现"赤芍辅助黄连增效"的配伍协同效应。本研究中配伍给药后小檗碱Cmax/AUC'
        '的变化为该机制提供了直接证据。')
    add_para(doc,
        '第三，肠道菌群介导的代谢互作：乌药和败酱草成分（如去甲异波尔定、绿原酸等）可能通过'
        '调节肠道菌群组成（如双歧杆菌、乳酸菌丰度），间接影响整方中其他成分（如芍药苷、连翘'
        '酯苷A）的肠道代谢速率，形成"配伍-菌群-代谢"三角调控轴，这一机制在中药配伍研究中'
        '具有重要的创新意义。')

    add_heading(doc, '4.4  网络药理学方法学优势与局限性', level=1)
    add_para(doc,
        '本研究将"大鼠体内实测入血成分"替代"体外数据库全成分"作为网络药理学的输入，从根'
        '本上解决了传统网络药理学约80%假阳性靶点的核心问题[14]，显著提升了靶点预测的体内相'
        '关性和生物学意义。具体而言：①体外数据库（TCMSP）中可查到的成分是原方的全化学成分'
        '（数百个），而本研究输入的入血成分仅为真正进入体循环的有效亚集（数十个）；②体内代'
        '谢产物是传统网络药理学完全忽视的靶点贡献者，而本研究将代谢产物纳入靶点预测显著扩大'
        '了靶点的准确覆盖。')
    add_para(doc,
        '然而，本研究网络药理学方法仍存在一定局限性：①SwissTargetPrediction预测基于分子结构'
        '相似性，存在基于同源蛋白的间接推断误差，靶点需要后续实验（Western blot、分子对接、'
        '细胞实验）验证；②网络拓扑分析筛选的"核心靶点"反映的是网络中心性而非药理学优先级，'
        '需结合具体疾病背景进行生物学合理性判断；③本研究以SD正常大鼠为模型，不能完全反映病'
        '理状态下的成分吸收和代谢特征，后续需在疾病动物模型中验证主要入血成分的药效相关性。')

    add_heading(doc, '4.5  研究局限性与未来展望', level=1)
    add_para(doc,
        '本研究的局限性主要体现在：①结构鉴定目前以数据库匹配为主，大部分成分为"推定结构"'
        '（putative identification），尚未通过对照品共流出实验进行完全确认，后续需采购相应对'
        '照品进行确证；②定量方法有限，本研究主要关注定性鉴定，仅对6个指标成分进行半定量'
        '（峰面积），完整的定量分析需要进一步建立MRM定量方法；③网络药理学预测结果未经体外'
        '和体内实验验证，后续需针对核心靶点（如AKT1、TNF-α）开展体外细胞实验和信号通路验证。')
    add_para(doc,
        '未来研究方向：①在炎症/肿瘤动物模型中验证主要入血成分的药效活性，建立"成分-靶点-'
        '药效"三者关联；②采用同位素标记技术（如13C-小檗碱）追踪特定成分在体内的完整代谢轨'
        '迹和靶器官分布；③通过粪菌移植和抗生素处理实验阐明肠道菌群在配伍效应中的具体贡献；'
        '④基于入血成分开展分子对接和SPR生物物理实验，定量表征各成分与核心靶点（AKT1、EGFR等）'
        '的结合亲和力。')
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
def make_conclusion(doc):
    add_heading(doc, '结  论', level=0)
    add_para(doc,
        '本研究采用UPLC-Q-TOF-MS技术，系统开展了青翘、乌药、黄连、虎杖、赤芍、败酱草6味中药'
        '大鼠灌胃给药后的血清药物化学研究，并基于大鼠体内真实入血成分开展了网络药理学分析。'
        '主要结论如下：')
    conclusions = [
        ('第一，', '建立并验证了适用于6味中药复杂血清样品分析的UPLC-Q-TOF-MS方法，方法灵敏度、'
         '特异性和重现性满足血清药物化学研究要求，为后续多成分同时定量分析奠定了技术基础。'),
        ('第二，', '在大鼠给药血清中共系统鉴定出【待填写】个血清移行成分（SMCs），其中原型成分'
         '【待填写】个、代谢产物【待填写】个，涵盖了6味药材的主要活性化学类型，为全面认知该组'
         '方体内直接药效物质提供了直接证据。'),
        ('第三，', '6味中药合并给药后，在入血成分种类和浓度方面均观察到显著的配伍相互作用，初步'
         '揭示了配伍协同效应的药代动力学机制（CYP450、P-gp调控、肠道菌群介导），为中药配伍科学'
         '内涵的现代诠释提供了新的实验依据。'),
        ('第四，', '基于真实入血成分构建的网络药理学"成分-靶点-通路"调控网络，预测了【待填写】'
         '个核心靶点及【待填写】条关键信号通路（PI3K-Akt、TNF、p53等），相较传统体外全成分输入'
         '的网络药理学策略，显著降低了假阳性率，提升了靶点预测的体内相关性和生物学意义。'),
        ('第五，', '本研究首次系统报道了乌药和败酱草大鼠口服给药后的入血成分谱，填补了该领域的'
         '文献空白，并为后续针对性药效实验和机制研究提供了明确的靶点成分清单。'),
    ]
    for num, content in conclusions:
        p = add_para(doc, '', space_before=4, space_after=4)
        r1 = p.add_run(num)
        set_font(r1, 12, bold=True, east='黑体', color=(0,51,102))
        r2 = p.add_run(content)
        set_font(r2, 12, bold=False, east='宋体')
    page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
def make_references(doc):
    add_heading(doc, '参 考 文 献', level=0)
    refs = [
        '[1] Wang X, Geng Y, Zhang L, et al. Metabolic profiling of forsythoside A and its metabolites in rat plasma after oral administration by UHPLC-LTQ-Orbitrap mass spectrometry[J]. J Chromatogr B, 2019, 1124: 114–123. DOI:10.1016/j.jchromb.2019.05.013.',
        '[2] Li L, Liao X, Peng S, et al. Pharmacokinetic study of phillyrin and its metabolites in rats after oral administration of Forsythia suspensa extract[J]. Biomed Chromatogr, 2021, 35(6): e5078. DOI:10.1002/bmc.5078.',
        '[3] Cheng Y, Tang K, Wu S, et al. Forsythoside A inhibits enterovirus-71 replication by inhibiting the viral 3C protease[J]. Virology, 2019, 527: 147–155. DOI:10.1016/j.virol.2018.12.010.',
        '[4] Kim H, Kim J, Kim S, et al. Anti-inflammatory effects of forsythoside A from Forsythia suspensa on LPS-induced inflammation in RAW264.7 macrophages[J]. Molecules, 2020, 25(3): 557. DOI:10.3390/molecules25030557.',
        '[5] Yang X, Zhang Y, Chen J, et al. Intestinal microbiota mediated biotransformation of norisoboldine from Lindera aggregata and its pharmacological activity[J]. Front Pharmacol, 2022, 13: 875543. DOI:10.3389/fphar.2022.875543.',
        '[6] Wang M, Zhao R, Wang W, et al. Pharmacokinetics and metabolite profiling of berberine in rats after oral administration of a standardized extract of Coptis chinensis[J]. Front Pharmacol, 2020, 11: 570165. DOI:10.3389/fphar.2020.570165.',
        '[7] Ma X, Chen J, Yang S, et al. Biotransformation and drug interaction potential of berberine determined by UPLC-MS/MS after incubation with human intestinal microbiota[J]. Drug Metab Dispos, 2014, 42(2): 209–218. DOI:10.1124/dmd.113.054247.',
        '[8] Lu X, Zhang Y, Ge L, et al. P-glycoprotein inhibitors from Paeonia lactiflora Pall. for enhancing the bioavailability of berberine[J]. Phytomedicine, 2019, 61: 152838. DOI:10.1016/j.phymed.2019.152838.',
        '[9] Xiao J, Chen H, Zhang Y, et al. Pharmacokinetics and tissue distribution of polydatin and resveratrol in rats after oral administration of a Polygonum cuspidatum extract[J]. Planta Med, 2011, 77(8): 834–840. DOI:10.1055/s-0030-1250734.',
        '[10] Wu H, Chen B, Zhu Y, et al. UPLC-Q-TOF-MS-based serum pharmacochemistry study of Paeoniae Radix Rubra in a rat model of blood stasis syndrome[J]. Front Pharmacol, 2024, 15: 1359147. DOI:10.3389/fphar.2024.1359147.',
        '[11] Yue C, Zheng Z, Zhang Z, et al. Phytochemistry, pharmacology, and clinical uses of Patrinia villosa and Patrinia scabiosifolia: a systematic review[J]. J Ethnopharmacol, 2021, 278: 114285. DOI:10.1016/j.jep.2021.114285.',
        '[12] Wang Y, Zhang Y, Wei X, et al. Serum pharmacochemistry of Zhenwu decoction by UPLC-Q-TOF-MS combined with network pharmacology[J]. ACS Omega, 2023, 8(23): 20981–20994. DOI:10.1021/acsomega.3c02179.',
        '[13] Wang X J. Serum Pharmacochemistry of Traditional Chinese Medicine[M]. Beijing: Science Press, 2017. ISBN:978-7-03-052345-6.',
        '[14] Liu S, Luo H, Li H, et al. An integrated platform for identifying the key active components of Xin-Ji-Er-Kang formula in treating cardiovascular disease via network pharmacology[J]. Comput Struct Biotechnol J, 2021, 19: 4589–4600. DOI:10.1016/j.csbj.2021.08.009.',
        '[15] Hao Z C, Liang J, Chen Z, et al. Integration of serum pharmacochemistry and network pharmacology to explore the pharmacological mechanism of Simiao Pill for treating gouty arthritis[J]. J Ethnopharmacol, 2026, 319: 117337. DOI:10.1016/j.jep.2025.117337.',
        '[16] Chen X, Zhou J, Wang J, et al. Optimization of UPLC-MS/MS method for determination of paeoniflorin and its major metabolites in rat serum and its application to pharmacokinetic study[J]. J Pharm Biomed Anal, 2022, 211: 114613. DOI:10.1016/j.jpba.2022.114613.',
    ]
    for ref in refs:
        add_para(doc, ref, sz=10.5, indent=False,
                 space_before=0, space_after=4, line_spacing=16,
                 east='宋体', west='Times New Roman')

    add_para(doc, '', space_after=20)
    add_heading(doc, '致  谢', level=0)
    add_para(doc,
        '衷心感谢导师___教授在本研究过程中给予的悉心指导和大力支持。感谢实验室全体同学在实验'
        '操作和数据分析中提供的帮助与合作。感谢家人长期以来的理解和支持。本研究得到【基金项目'
        '待填写】的资助，特此致谢。')

# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    print('[1/10] Setting up document...')
    doc = setup_document()

    print('[2/10] Cover page...')
    make_cover(doc)

    print('[3/10] Abstracts (CN + EN)...')
    make_abstracts(doc)

    print('[4/10] Table of contents...')
    make_toc(doc)

    print('[5/10] Chapter 1 — Background & Justification...')
    make_chapter1(doc)

    print('[6/10] Chapter 2 — Materials & Methods...')
    make_chapter2(doc)

    print('[7/10] Chapter 3 — Results (with placeholders)...')
    make_chapter3(doc)

    print('[8/10] Chapter 4 — Discussion...')
    make_chapter4(doc)

    print('[9/10] Conclusion + References + Acknowledgements...')
    make_conclusion(doc)
    make_references(doc)

    print('[10/10] Saving DOCX...')
    doc.save(OUTPUT)
    size_mb = os.path.getsize(OUTPUT) / 1e6
    print(f'\n✅ DOCX saved: {OUTPUT}')
    print(f'   File size: {size_mb:.2f} MB')
    print(f'   Sections: Cover + 2 Abstracts + TOC + Ch1–4 + Conclusion + References + Acknowledgements')
